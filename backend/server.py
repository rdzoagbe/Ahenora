import os
import io
import re
import json
import base64
import asyncio
import hashlib
import secrets
import tempfile
import html
import urllib.error
import urllib.request
import urllib.parse
from datetime import datetime, timedelta, timezone
from typing import Any, Optional
from fastapi import FastAPI, HTTPException, Depends, Header, UploadFile, File, Query, Body
from dedupe_core import run as dedupe_run
from fastapi.middleware.cors import CORSMiddleware
from pydantic import BaseModel, Field
try:
    from motor.motor_asyncio import AsyncIOMotorClient
except ImportError:
    AsyncIOMotorClient = None
from google.oauth2 import id_token as google_id_token
from google.auth.transport.requests import Request as GoogleRequest
try:
    from google import genai
    from google.genai import types as genai_types
except ImportError:
    genai = None
    genai_types = None
import PIL.Image

from ai_models import model_candidates, should_try_next_model, summarize_ai_error
from ai_safety import (
    MAX_INGREDIENT_LEN,
    VEGETARIAN,
    RECIPE_SYSTEM_PROMPT,
    SUGGEST_SYSTEM_PROMPT,
    build_suggest_prompt,
    validate_suggestions,
    UnsafeRecipe,
    build_recipe_prompt,
    extract_json,
    sanitize_ingredients,
    sanitize_user_text,
    sanitize_message_text,
    validate_recipe,
    CHEF_SYSTEM_PROMPT,
    MAX_QUESTION_LEN,
    build_chef_prompt,
    validate_chef_answer,
    SHOPPING_SCAN_SYSTEM_PROMPT,
    validate_shopping_scan,
    RECIPE_PHOTO_SYSTEM_PROMPT,
    validate_captured_recipe,
    VAULT_CATEGORIES,
    build_document_scan_prompt,
    validate_document_scan,
)


# -----------------------------------------------------------------------------
# Config
# -----------------------------------------------------------------------------
MONGO_URL = os.environ.get("MONGO_URL", "")
DB_NAME = os.environ.get("DB_NAME", "household_coo")
GOOGLE_WEB_CLIENT_ID = os.environ.get("GOOGLE_WEB_CLIENT_ID", "")
GOOGLE_ANDROID_CLIENT_ID = os.environ.get("GOOGLE_ANDROID_CLIENT_ID", "")
GOOGLE_CLIENT_IDS_EXTRA = os.environ.get("GOOGLE_CLIENT_IDS", "")
# Fallback web client ID — must match the app's hardcoded fallback in
# frontend/app/index.tsx so Google sign-in still verifies if the Railway env
# var is ever unset. The token audience is this web client ID.
GOOGLE_WEB_CLIENT_ID_FALLBACK = "243255248169-cei972lc7kmfig6tmjb6l2nlmgqkjf22.apps.googleusercontent.com"
GOOGLE_CLIENT_IDS = [
    client_id.strip()
    for client_id in [
        GOOGLE_WEB_CLIENT_ID,
        GOOGLE_ANDROID_CLIENT_ID,
        *GOOGLE_CLIENT_IDS_EXTRA.split(","),
        GOOGLE_WEB_CLIENT_ID_FALLBACK,
    ]
    if client_id and client_id.strip()
]
# De-duplicate while preserving order.
GOOGLE_CLIENT_IDS = list(dict.fromkeys(GOOGLE_CLIENT_IDS))
GOOGLE_API_KEY = os.environ.get("GOOGLE_API_KEY", "")
# A household app is opened daily and lived in for years. A seven-day session
# signed everyone out weekly — a tap for a Google user, but a retyped password
# for an email one, which is why coming back felt broken. Ninety days, and the
# clock resets on use (see _touch_session), so an active family is never signed
# out while an abandoned session still lapses.
SESSION_DAYS = int(os.environ.get("SESSION_DAYS", "90"))
INVITE_DAYS = int(os.environ.get("INVITE_DAYS", "14"))
# Invite links must open SOMEWHERE on every device. The old default was the
# native custom scheme, which does nothing on a phone without the app —
# an iPhone tapping it got silence. The web companion handles ?invite=
# end to end, so it is the universal default; the env var can still point
# elsewhere (e.g. a future custom domain).
INVITE_BASE_URL = os.environ.get(
    "INVITE_BASE_URL", "https://ahenora.com/app/"
)

# Email delivery. Resend is used through the standard-library urllib client,
# so no extra Python package is required.
RESEND_API_KEY = os.environ.get("RESEND_API_KEY", "")
INVITE_FROM_EMAIL = os.environ.get("INVITE_FROM_EMAIL", "")
INVITE_REPLY_TO = os.environ.get("INVITE_REPLY_TO", "")
APP_NAME = os.environ.get("APP_NAME", "Ahenora")
MAX_VOICE_AUDIO_BYTES = int(os.environ.get("MAX_VOICE_AUDIO_BYTES", str(12 * 1024 * 1024)))
ADMIN_EMAILS_RAW = os.environ.get("ADMIN_EMAILS", "")
ADMIN_EMAILS = {
    email.strip().lower()
    for email in ADMIN_EMAILS_RAW.split(",")
    if email.strip()
}

ALLOWED_ORIGINS = [
    origin.strip()
    for origin in os.environ.get("ALLOWED_ORIGINS", "").split(",")
    if origin.strip()
] or [
    # The live web app. It moved here from the github.io project page when the
    # domain was set up; the origin changed but this list did not follow, so a
    # browser on ahenora.com had every API call refused by CORS. Both the bare
    # and www hosts, because either can be what the browser actually sends.
    "https://ahenora.com",
    "https://www.ahenora.com",
    "https://household-coo.app",
    "https://www.household-coo.app",
    # The previous home of the web companion app, kept for old links.
    "https://rdzoagbe.github.io",
    "householdcoo://",
    "exp://",
]


# Connection resilience: after the Atlas M0 -> M10 migration the running process
# held connections to servers that no longer existed, so every query waited out
# the full server-selection timeout and the app returned 500s until a manual
# restart. These settings let the driver notice a dead/changed topology quickly,
# recycle stale sockets on its own, and bound how long any single operation can
# hang, so a database blip degrades instead of cascading.
mongo = (
    AsyncIOMotorClient(
        MONGO_URL,
        serverSelectionTimeoutMS=5000,
        connectTimeoutMS=5000,
        socketTimeoutMS=20000,
        # Recycle idle sockets so a silently-dead connection is never reused.
        maxIdleTimeMS=30000,
        heartbeatFrequencyMS=10000,
        # Datetimes come back timezone-aware. Without this, BSON round-trips
        # strip the tzinfo and every stored datetime returns NAIVE, while
        # utcnow() is aware — so `stored < utcnow()` raises TypeError. That is
        # not a theoretical worry: it took out invite acceptance in production
        # (`expires_at < utcnow()` in _resolve_invite) while every local test
        # passed, because the in-memory test double hands back the very objects
        # it was given, tzinfo intact. The fake was kinder than reality.
        tz_aware=True,
        retryWrites=True,
        retryReads=True,
        maxPoolSize=50,
        minPoolSize=0,
        appname="household-coo-backend",
    )
    if MONGO_URL
    else None
)
db: Any = mongo[DB_NAME] if mongo else None

import logging
from collections import defaultdict
from starlette.requests import Request
from starlette.responses import JSONResponse
from starlette.middleware.base import BaseHTTPMiddleware

log = logging.getLogger("household_coo")

app = FastAPI(title="Ahenora Backend")

app.add_middleware(
    CORSMiddleware,
    allow_origins=ALLOWED_ORIGINS,
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)


@app.on_event("startup")
async def reset_testing_window_plans():
    # Cleanup for when real billing goes live: a paid plan stored without any
    # RevenueCat event came from the testing window's self-serve switcher, not
    # from a purchase — reset it so launch gating is honest. Idempotent: real
    # subscribers always carry rc_last_event (set by the webhook).
    if db is None or not os.environ.get("RC_WEBHOOK_SECRET"):
        return
    try:
        result = await db["families"].update_many(
            {"plan": {"$ne": "village"}, "rc_last_event": {"$exists": False}},
            {"$set": {"plan": "village", "updated_at": datetime.now(timezone.utc)}},
        )
        if result.modified_count:
            log.info("Reset %d testing-window plan(s) to village", result.modified_count)
    except Exception as exc:  # pragma: no cover - startup must never crash the app
        log.warning("Testing-window plan reset skipped: %s", exc)


# -----------------------------------------------------------------------------
# Rate limiting (in-memory, per-IP)
# -----------------------------------------------------------------------------
RATE_LIMIT_WINDOW = 60  # seconds
RATE_LIMIT_MAX = int(os.environ.get("RATE_LIMIT_MAX", "120"))
# Auth limit is per IP+path. Kept generous so a cohort of testers sharing one
# egress IP (home/office/school NAT), each with client-side retries, doesn't
# collectively trip a 429 on sign-in. Still low enough to blunt brute force.
RATE_LIMIT_AUTH_MAX = int(os.environ.get("RATE_LIMIT_AUTH_MAX", "60"))

_rate_buckets: dict[str, list[float]] = defaultdict(list)

# Per-identity failed-auth tracker (keyed by email or member, NOT by IP), so an
# attacker cannot sidestep brute-force protection by spoofing X-Forwarded-For or
# rotating IPs. In-process; resets on redeploy, which is acceptable for a
# lockout backstop layered on top of the per-IP limiter.
_auth_fail: dict[str, list[float]] = defaultdict(list)
AUTH_FAIL_MAX = int(os.environ.get("AUTH_FAIL_MAX", "8"))
AUTH_FAIL_WINDOW = int(os.environ.get("AUTH_FAIL_WINDOW", "900"))  # 15 minutes

# Forgotten-password reset: a short-lived numeric code emailed to the account's
# inbox. Kept short so it is easy to type, safe only because it expires fast, is
# tried a handful of times at most, and is rate-limited per email.
PASSWORD_RESET_TTL_MINUTES = int(os.environ.get("PASSWORD_RESET_TTL_MINUTES", "15"))
PASSWORD_RESET_MAX_ATTEMPTS = int(os.environ.get("PASSWORD_RESET_MAX_ATTEMPTS", "5"))


def _auth_locked(identity: str) -> bool:
    import time

    now = time.time()
    hits = [t for t in _auth_fail.get(identity, []) if t > now - AUTH_FAIL_WINDOW]
    if hits:
        _auth_fail[identity] = hits
    else:
        _auth_fail.pop(identity, None)
    return len(hits) >= AUTH_FAIL_MAX


def _auth_record_fail(identity: str) -> None:
    import time

    _auth_fail[identity].append(time.time())
    # Bound memory against identity-spray: drop the oldest keys wholesale.
    if len(_auth_fail) > 20000:
        for k in list(_auth_fail.keys())[:10000]:
            _auth_fail.pop(k, None)


def _auth_clear(identity: str) -> None:
    _auth_fail.pop(identity, None)


def _client_ip(request: Request) -> str:
    # Rightmost X-Forwarded-For hop, not the first: the first entry is
    # client-supplied and freely spoofable, which let one machine rotate fake
    # addresses and sidestep the per-IP limiter entirely. The LAST entry is the
    # one appended by our own fronting proxy (Railway), so it is the only hop
    # we can actually trust.
    forwarded = request.headers.get("x-forwarded-for")
    if forwarded:
        return forwarded.split(",")[-1].strip()
    return request.client.host if request.client else "unknown"


def _prune(bucket: list[float], now: float) -> list[float]:
    cutoff = now - RATE_LIMIT_WINDOW
    return [t for t in bucket if t > cutoff]


class RateLimitMiddleware(BaseHTTPMiddleware):
    async def dispatch(self, request: Request, call_next):  # type: ignore[override]
        import time

        ip = _client_ip(request)
        now = time.time()
        key = f"{ip}:{request.url.path}"

        is_auth = request.url.path.startswith("/api/auth/")
        limit = RATE_LIMIT_AUTH_MAX if is_auth else RATE_LIMIT_MAX

        bucket = _prune(_rate_buckets[key], now)
        if len(bucket) >= limit:
            return JSONResponse(
                status_code=429,
                content={"detail": "Too many requests. Please try again later."},
            )
        bucket.append(now)
        _rate_buckets[key] = bucket

        # Evict empty buckets so path/IP spray can't grow memory unbounded.
        if len(_rate_buckets) > 50000:
            for k, v in list(_rate_buckets.items()):
                if not _prune(v, now):
                    _rate_buckets.pop(k, None)

        return await call_next(request)


app.add_middleware(RateLimitMiddleware)


# -----------------------------------------------------------------------------
# Helpers
# -----------------------------------------------------------------------------
def utcnow() -> datetime:
    return datetime.now(timezone.utc)


def iso(dt: Optional[datetime]) -> Optional[str]:
    if not dt:
        return None
    return dt.astimezone(timezone.utc).isoformat()


def parse_dt(value: Optional[str]) -> Optional[datetime]:
    if not value:
        return None
    value = value.replace("Z", "+00:00")
    try:
        dt = datetime.fromisoformat(value)
    except (ValueError, TypeError):
        # Malformed client input must be a 400, not an unhandled 500.
        raise HTTPException(status_code=400, detail="Invalid date format")
    if dt.tzinfo is None:
        dt = dt.replace(tzinfo=timezone.utc)
    return dt.astimezone(timezone.utc)

def ensure_aware_utc(value):
    if not value:
        return None
    if isinstance(value, str):
        return parse_dt(value)
    if isinstance(value, datetime):
        if value.tzinfo is None:
            return value.replace(tzinfo=timezone.utc)
        return value.astimezone(timezone.utc)
    return None


def advance_due_date(dt: datetime, recurrence: str) -> datetime:
    """Return the next occurrence for a recurring card's due date."""
    if recurrence == "daily":
        return dt + timedelta(days=1)
    if recurrence == "weekly":
        return dt + timedelta(weeks=1)
    if recurrence == "monthly":
        month = dt.month + 1
        year = dt.year + (month - 1) // 12
        month = (month - 1) % 12 + 1
        # Clamp the day to the last valid day of the target month.
        for day in (dt.day, 30, 29, 28):
            try:
                return dt.replace(year=year, month=month, day=day)
            except ValueError:
                continue
    return dt


def new_id(prefix: str) -> str:
    return f"{prefix}_{secrets.token_hex(8)}"


def sha256(value: str) -> str:
    return hashlib.sha256(value.encode("utf-8")).hexdigest()


PASSWORD_ITERATIONS = 200_000


def hash_password(password: str) -> str:
    salt = secrets.token_hex(16)
    dk = hashlib.pbkdf2_hmac(
        "sha256", password.encode("utf-8"), salt.encode("utf-8"), PASSWORD_ITERATIONS
    )
    return f"pbkdf2_sha256${PASSWORD_ITERATIONS}${salt}${dk.hex()}"


def verify_password(password: str, stored: str) -> bool:
    try:
        algo, iterations, salt, hexhash = stored.split("$")
        if algo != "pbkdf2_sha256":
            return False
        dk = hashlib.pbkdf2_hmac(
            "sha256", password.encode("utf-8"), salt.encode("utf-8"), int(iterations)
        )
        return secrets.compare_digest(dk.hex(), hexhash)
    except Exception:
        return False


# A PIN used to be stored as a bare sha256 of the four digits. That is only
# 10,000 possible hashes with no salt, so a copy of the database gave up every
# child's PIN at once, and two families who picked 1234 were visibly the same
# row. Salted PBKDF2 makes a leaked PIN cost work rather than a lookup table.
# Old hashes still verify, and are rewritten the next time they are used
# correctly, so nobody is locked out by the change.
def hash_pin(pin: str) -> str:
    return hash_password(pin)


def verify_pin(pin: str, stored: str) -> bool:
    if not stored:
        return False
    if stored.startswith("pbkdf2_sha256$"):
        return verify_password(pin, stored)
    return secrets.compare_digest(sha256(pin), stored)


def pin_is_legacy(stored: str) -> bool:
    return bool(stored) and not stored.startswith("pbkdf2_sha256$")


async def _verify_pin_and_upgrade(database, member: dict, pin: str) -> bool:
    """Check a PIN, and quietly re-store it salted if it was an old bare hash."""
    stored = member.get("pin_hash") or ""
    if not verify_pin(pin, stored):
        return False
    if pin_is_legacy(stored):
        await database["family_members"].update_one(
            {"member_id": member["member_id"]}, {"$set": {"pin_hash": hash_pin(pin)}})
    return True


async def _touch_session(database, token_hash: str, session: dict) -> None:
    """Push a live session's expiry back out, at most once a day.

    Without this, expiry counts from sign-in rather than from last use, so a
    family that opens the app every morning is still logged out on schedule.
    Renewing only when the session has less than (SESSION_DAYS - 1) left costs
    one write per active day instead of one per request.
    """
    expires = _coerce_dt(session.get("expires_at"))
    if not expires:
        return
    # A one-day floor keeps a deliberately short SESSION_DAYS from renewing on
    # every single request.
    threshold = timedelta(days=max(SESSION_DAYS - 1, 0.5))
    if expires - utcnow() >= threshold:
        return
    try:
        await database["user_sessions"].update_one(
            {"token_hash": token_hash},
            {"$set": {"expires_at": utcnow() + timedelta(days=SESSION_DAYS)}},
        )
    except Exception as exc:  # a renewal failure must never break a request
        log.warning("Session renewal skipped: %s", exc)


def is_admin_email(email: str) -> bool:
    return bool(email) and email.strip().lower() in ADMIN_EMAILS


def is_admin_user(user: dict) -> bool:
    return is_admin_email(user.get("email", ""))


# Tester privileges belong to the HOUSEHOLD, not the email: the founder's
# co-parent joined his family and still saw the Free plan, because the
# admin override was per-user while every real subscription is per-family.
# Cached briefly — build_subscription runs on nearly every request.
_ADMIN_FAMILY_CACHE: dict = {}


async def family_has_admin(database, family_id: str) -> bool:
    if not ADMIN_EMAILS or not family_id:
        return False
    now = utcnow().timestamp()
    cached = _ADMIN_FAMILY_CACHE.get(family_id)
    if cached and now - cached[0] < 60:
        return cached[1]
    has_admin = False
    async for member in database["users"].find({"family_id": family_id}, {"_id": 0, "email": 1}):
        if is_admin_email(member.get("email", "")):
            has_admin = True
            break
    _ADMIN_FAMILY_CACHE[family_id] = (now, has_admin)
    return has_admin


def apply_admin_subscription(subscription: dict) -> dict:
    # Admin/tester accounts keep their own family data, but plan limits are bypassed
    # so the founder can test every feature without changing customer billing rules.
    admin_sub = dict(subscription)
    admin_sub["plan"] = "family_office"
    admin_sub["billing_cycle"] = admin_sub.get("billing_cycle", "yearly")
    admin_sub["grandfathered"] = True
    admin_sub["admin_unlocked"] = True
    admin_sub["limits"] = {
        "max_members": 999,
        "max_children": 999,
        "ai_scans_per_month": 999999,
        "vault_bytes": 50 * 1024 * 1024 * 1024,
        "weekly_brief": True,
        "multi_property": True,
        "meal_planner": True,
        "allowance": True,
        "carpool": True,
        "weekly_report": True,
    }
    admin_sub["price_monthly"] = 0.0
    admin_sub["price_yearly"] = 0.0
    return admin_sub


def get_db() -> Any:
    if db is None:
        raise HTTPException(status_code=500, detail="Database not configured")
    return db


GEMINI_MODEL = os.environ.get("GEMINI_MODEL", "").strip()

# Which model actually answered, and the last per-model failure. Written by
# _gemini_generate, read by /api/health/ai so production can say which model it
# is really on instead of us assuming.
_gemini_state = {"model": None, "last_error": None, "errors": {}, "discovered": None,
                 "client": None}


def _gemini_client():
    """The one client for this process, built on first use.

    The retired google-generativeai package kept the key in module-level global
    state configured at import; google-genai hangs it on a client instead.
    Built lazily rather than at import so a transient failure here cannot
    permanently disable every AI feature for the life of the process — the next
    call simply tries again.
    """
    if _gemini_state["client"] is not None:
        return _gemini_state["client"]
    if not GOOGLE_API_KEY or not genai:
        return None
    try:
        _gemini_state["client"] = genai.Client(api_key=GOOGLE_API_KEY)
    except Exception as exc:  # noqa: BLE001 — surfaced by /api/health/ai
        _gemini_state["last_error"] = f"client: {exc}"[:300]
        log.warning("gemini client could not be created: %s", exc)
        return None
    return _gemini_state["client"]


def _discover_models() -> list:
    """Ask the live key which models it can use for generateContent.

    Model names drift (Google renames and retires them), and a hardcoded list
    goes stale. When every candidate 404s, this is the source of truth: whatever
    the key is actually entitled to. Returns [] if listing itself fails (which
    usually means the Generative Language API is not enabled on the project).
    """
    client = _gemini_client()
    if not client:
        return []
    # Substrings that mark a model as NOT a plain text generator (image, audio,
    # robotics, research, etc.). These are demoted to the very end rather than
    # dropped: a text model is always preferred, but if — on this frozen SDK —
    # none of the text aliases resolve, a proven-working image model is a far
    # better last resort than the feature going dark.
    NON_TEXT = (
        "image", "tts", "audio", "veo", "imagen", "lyria", "nano-banana",
        "robotics", "embedding", "aqa", "computer-use", "deep-research",
        "antigravity", "omni",
    )

    def rank(name: str) -> tuple:
        # Lower sorts first. Non-text last; then clean flash, flash-lite, pro,
        # rest; stable names before -preview; then alphabetical.
        n = name.lower()
        non_text = 1 if any(bad in n for bad in NON_TEXT) else 0
        family = 0 if ("flash" in n and "lite" not in n) else 1 if "flash" in n else 2 if "pro" in n else 3
        preview = 1 if "preview" in n else 0
        latest = 0 if n.endswith("-latest") else 1  # stable aliases first
        return (non_text, family, preview, latest, name)

    try:
        names = []
        # query_base=True asks for the published models rather than tuned ones.
        # It is the default, but the distinction matters enough to say out loud.
        for m in client.models.list(config={"query_base": True}):
            # google-genai renamed this field from supported_generation_methods
            # to supported_actions. Read both, and — importantly — keep a model
            # that reports neither. This discovery exists to survive Google
            # renaming things; a filter that silently empties the list would
            # disable the very self-healing it is here to provide.
            actions = (
                getattr(m, "supported_actions", None)
                or getattr(m, "supported_generation_methods", None)
                or []
            )
            if actions and "generateContent" not in actions:
                continue
            if not m.name:
                continue
            names.append(m.name.split("/")[-1])
        names.sort(key=rank)
        return names
    except Exception as exc:  # noqa: BLE001 — reported by the caller
        _gemini_state["last_error"] = f"models.list: {exc}"[:300]
        return []


# Simple jobs (a chef substitution, reading a shopping list off a photo) go
# to the lite model first: roughly twice as fast, and the safety gates do not
# care which model produced the text they check. Complex jobs (weekly planning,
# structured recipes) stay on the stronger chain.
FAST_MODEL = "gemini-flash-lite-latest"


async def _gemini_generate(contents, system: str = "", temperature: float = None, fast: bool = False) -> str:
    """Generate with model fallback.

    Tries the proven model first, then the operator override, then the built-in
    list. Only a NOT_FOUND-style error moves to the next candidate — anything
    else (quota, safety, network) is a real failure and re-raised immediately,
    because retrying it against three models triples the pain for no benefit.

    `temperature` is optional; when set it controls sampling diversity, which
    the meal planner raises for repeat "different ideas" asks.

    `fast` prepends the lite model for latency-sensitive simple jobs. A fast
    success is deliberately NOT remembered as the proven model — otherwise one
    chef answer would quietly downgrade every recipe that follows it.
    """
    client = _gemini_client()
    if not client:
        raise RuntimeError("Gemini is not configured")

    config = {}
    if system:
        config["system_instruction"] = system
    if temperature is not None:
        config["temperature"] = temperature

    last_error = None
    candidates = model_candidates(GEMINI_MODEL, _gemini_state["model"] or "")
    if fast:
        candidates = [FAST_MODEL] + [c for c in candidates if c != FAST_MODEL]
    # If none of the built-in names exist for this key, ask the key what it has
    # and append those. Self-heals against Google renaming/retiring models.
    discovered = _gemini_state.get("discovered")
    if discovered is None:
        discovered = _discover_models()
        _gemini_state["discovered"] = discovered
    for name in discovered:
        if name not in candidates:
            candidates.append(name)
    for name in candidates:
        try:
            # .aio is natively async, so this no longer needs a worker thread.
            response = await client.aio.models.generate_content(
                model=name,
                contents=contents,
                config=config or None,
            )
            if not fast:
                if _gemini_state["model"] != name:
                    log.info("gemini model resolved to %s", name)
                _gemini_state["model"] = name
            _gemini_state["last_error"] = None
            _gemini_state["errors"].pop(name, None)
            await _bump_metric("ai_call_ok")
            return (response.text or "").strip()
        except Exception as exc:  # noqa: BLE001 — classified below
            _gemini_state["last_error"] = f"{name}: {exc}"[:300]
            # Per-model failure map, so /api/health/ai can show whether one
            # model is out of quota or the whole key is.
            _gemini_state["errors"][name] = summarize_ai_error(str(exc))
            last_error = exc
            if not should_try_next_model(str(exc)):
                await _bump_metric("ai_call_error")
                raise
            log.warning("gemini model %s unavailable (%s), trying next candidate",
                        name, _gemini_state["errors"][name])
    await _bump_metric("ai_call_error")
    raise last_error if last_error else RuntimeError("no gemini model candidates")


async def _gemini_text(prompt: str, system: str = "", temperature: float = None, fast: bool = False) -> str:
    return await _gemini_generate(prompt, system, temperature=temperature, fast=fast)


async def _gemini_vision(prompt: str, image_base64: str, system: str = "", fast: bool = False) -> str:
    if "," in image_base64:
        image_base64 = image_base64.split(",")[-1]
    img_bytes = base64.b64decode(image_base64)
    img = PIL.Image.open(io.BytesIO(img_bytes))
    return await _gemini_generate([prompt, img], system, fast=fast)


# The two launch tiers. Young people are metered (role-aware: parents/caregivers
# never count) — max_children caps kids and teens TOGETHER, so a teen account is
# never a free way around the limit; max_members stays as a generous total
# safety cap. "family_office" is retired — any family doc still carrying it
# resolves to "executive" via plan_catalog_for().
PLAN_CATALOG = {
    "village": {
        "price_monthly": 0.0,
        "price_yearly": 0.0,
        "limits": {
            "max_members": 10,
            "max_children": 2,
            "ai_scans_per_month": 5,
            "vault_bytes": 25 * 1024 * 1024,
            "weekly_brief": False,
            "multi_property": False,
            # Premium feature flags
            "meal_planner": False,
            "allowance": False,
            "carpool": False,
            "weekly_report": False,
        },
    },
    "executive": {
        "price_monthly": 6.99,
        "price_yearly": 49.99,
        "limits": {
            "max_members": 12,
            "max_children": 5,
            "ai_scans_per_month": 100,
            "vault_bytes": 500 * 1024 * 1024,
            "weekly_brief": True,
            "multi_property": False,
            "meal_planner": True,
            "allowance": True,
            "carpool": True,
            "weekly_report": True,
        },
    },
}


def plan_catalog_for(plan: str) -> dict:
    """Resolve a stored plan name to a catalog entry. Unknown/retired paid
    plans (e.g. legacy "family_office") map to executive so no family ever
    loses paid features; anything else falls back to the free tier."""
    if plan in PLAN_CATALOG:
        return PLAN_CATALOG[plan]
    if plan == "family_office":
        return PLAN_CATALOG["executive"]
    return PLAN_CATALOG["village"]


# Features gated behind paid plans. Maps the feature flag to a user-facing
# upgrade message used when a free-tier family hits the gate (HTTP 402).
PREMIUM_FEATURE_MESSAGES = {
    "meal_planner": "Meal Planner is available on Premium.",
    "allowance": "Pocket money tracking is available on Premium.",
    "carpool": "Carpool Coordinator is available on Premium.",
    "weekly_report": "Weekly Report is available on Premium.",
}


def plan_limit_error(feature: str, current_plan: str, message: str, limit=None, used=None):
    raise HTTPException(
        status_code=402,
        detail={
            "error": "plan_limit",
            "feature": feature,
            "current_plan": current_plan,
            "limit": limit,
            "used": used,
            "message": message,
        },
    )


async def require_feature(user: dict, feature: str):
    """Gate a premium feature: raises HTTP 402 if the family's plan doesn't
    include it. Admin/founder accounts always pass. Returns the subscription."""
    sub = await build_subscription(user["family_id"])
    if is_admin_user(user):
        return apply_admin_subscription(sub)
    if not sub["limits"].get(feature, True):
        plan_limit_error(
            feature=feature,
            current_plan=sub["plan"],
            message=PREMIUM_FEATURE_MESSAGES.get(feature, "Upgrade to use this feature."),
        )
    return sub


AI_SCAN_PERIOD_DAYS = 30


def _coerce_dt(value):
    """Best-effort parse of a stored datetime (datetime or ISO string).

    Always returns a timezone-aware UTC datetime so it can be safely compared
    against utcnow(). PyMongo returns naive datetimes by default, so naive
    values are assumed to be UTC.
    """
    parsed = None
    if isinstance(value, datetime):
        parsed = value
    elif isinstance(value, str):
        try:
            parsed = datetime.fromisoformat(value.replace("Z", "+00:00"))
        except ValueError:
            return None
    if parsed is None:
        return None
    if parsed.tzinfo is None:
        parsed = parsed.replace(tzinfo=timezone.utc)
    return parsed


# The founder is the earliest parent. A row with no usable created_at must sort
# as oldest, never "now" — a legacy founder row can lack a timestamp, and
# sorting it last would hand founder status (and its protection) to a later
# co-parent.
_OLDEST_DT = datetime(1970, 1, 1, tzinfo=timezone.utc)


async def get_family_doc(family_id: str):
    database = get_db()
    family = await database["families"].find_one({"family_id": family_id}, {"_id": 0})
    if not family:
        family = {
            "family_id": family_id,
            "plan": "village",
            "billing_cycle": "monthly",
            "grandfathered": False,
            "updated_at": utcnow(),
            "ai_scans_used": 0,
            "ai_scans_period_start": utcnow(),
            "vault_bytes_used": 0,
        }
        await database["families"].insert_one(family)
        return family

    # Roll the monthly AI-scan counter over once the period has elapsed so the
    # "per month" limit actually behaves monthly instead of being a lifetime cap.
    period_start = _coerce_dt(family.get("ai_scans_period_start"))
    now = utcnow()
    if period_start is None or (now - period_start) >= timedelta(days=AI_SCAN_PERIOD_DAYS):
        await database["families"].update_one(
            {"family_id": family_id},
            {"$set": {"ai_scans_used": 0, "ai_scans_period_start": now}},
        )
        family["ai_scans_used"] = 0
        family["ai_scans_period_start"] = now
    return family


async def charge_ai_scan(user) -> None:
    """Charge one AI scan against the family's monthly allowance, atomically.

    The paid AI routes read ai_scans_used, compare, then $inc later — two
    separate operations, so two concurrent scans could both read 9/10, both
    pass, and both increment to 11. This does it in one guarded update: the
    increment only lands while the family is still under the limit at write
    time, the same way the money paths guard their decrements. Over the limit
    (or on a race that took the last slot) it raises the upgrade wall. Admins
    never count.
    """
    database = get_db()
    if is_admin_user(user):
        return
    sub = await build_subscription(user["family_id"])
    limit = sub["limits"]["ai_scans_per_month"]
    result = await database["families"].update_one(
        {"family_id": user["family_id"], "ai_scans_used": {"$lt": limit}},
        {"$inc": {"ai_scans_used": 1}, "$set": {"updated_at": utcnow()}},
    )
    if result.modified_count == 0:
        plan_limit_error(
            feature="ai_scans",
            current_plan=sub["plan"],
            limit=limit,
            used=limit,
            message="AI scan limit reached for this billing period.",
        )


async def ai_scans_remaining(user) -> int:
    """How many scans are left, without charging one.

    For cheap, auto-fired helpers (assign) that should quietly fall back to a
    non-AI answer when a family is out — rather than raise the upgrade wall a
    deliberate document scan does, or burn a document-scan slot on a helper
    the user never explicitly asked for.
    """
    if is_admin_user(user):
        return 1
    sub = await build_subscription(user["family_id"])
    return max(0, sub["limits"]["ai_scans_per_month"] - sub["ai_scans_used"])


async def count_young_people(database, family_id: str) -> int:
    """Kids and teens share one plan limit — both are 'young people'. A teen is
    a separate login, but a family with two teens should get the same allowance
    as a family with two kids, so a teen must never be a free way around the
    cap. Parents, co-parents and carers are never the meter."""
    return await database["family_members"].count_documents(
        {"family_id": family_id, "role": {"$regex": "^(child|teen)$", "$options": "i"}}
    )


CUSTODY_WEEKS = ("even", "odd")
CUSTODY_LABEL_MAX = 40


def public_custody(family: dict) -> dict:
    """Alternating-custody (garde alternée) config for the client.

    French judgments write custody as semaines paires / impaires, so the whole
    schedule is one parity: which ISO weeks the children are in this home. Off by
    default, absent on every family that predates this — a family that never sets
    it up sees nothing, on the Feed or the Calendar.
    """
    weeks = family.get("custody_our_weeks")
    return {
        "enabled": bool(family.get("custody_enabled")),
        "our_weeks": weeks if weeks in CUSTODY_WEEKS else "even",
        "away_label": family.get("custody_away_label") or "",
    }


async def build_subscription(family_id: str):
    database = get_db()
    family = await get_family_doc(family_id)
    members_count = await database["family_members"].count_documents({"family_id": family_id})
    # children_count stays children-only for display; young_people_count is what
    # the plan cap actually meters (kids + teens together).
    children_count = await database["family_members"].count_documents(
        {"family_id": family_id, "role": {"$regex": "^child$", "$options": "i"}}
    )
    young_people_count = await count_young_people(database, family_id)
    catalog = plan_catalog_for(family["plan"])
    limits = catalog["limits"]
    # TESTING WINDOW: until billing is live (RC_WEBHOOK_SECRET set), every
    # family gets Premium limits so closed-test families can exercise the
    # gated features (meal planner, allowance, carpool, weekly report) and
    # aren't blocked by the child cap. Gates enforce automatically the moment
    # billing is configured — same trigger that locks /subscription/change.
    testing_window = not os.environ.get("RC_WEBHOOK_SECRET")
    # A household containing an admin/tester account is a tester household:
    # everyone in it shares the top plan, exactly like a real purchase would
    # be shared. Fixes the co-parent seeing Free next to the founder.
    admin_household = await family_has_admin(database, family_id)
    # Grandfathered families keep Premium after billing goes live — the grace
    # period's exemption for early adopters ("founding families"). Set the flag
    # when the cutover date passes for anyone we're thanking / not charging yet.
    grandfathered = bool(family.get("grandfathered"))
    if testing_window or admin_household or grandfathered:
        limits = PLAN_CATALOG["executive"]["limits"]
    return {
        "plan": "family_office" if admin_household else family["plan"],
        # Lets the app show "you're previewing Premium free" notices so launch
        # gating never feels like a surprise takeaway.
        "testing_window": testing_window,
        # The announced cutover date (ISO, e.g. "2026-10-01") drives the in-app
        # countdown. Empty until we commit a date — the app shows the plain
        # free-preview notice until then, the countdown only once it's set.
        "billing_starts_at": os.environ.get("BILLING_START_DATE") or None,
        "billing_cycle": family["billing_cycle"],
        "grandfathered": grandfathered,
        "updated_at": iso(family.get("updated_at")),
        "ai_scans_used": family.get("ai_scans_used", 0),
        "ai_scans_period_start": iso(family.get("ai_scans_period_start")),
        "vault_bytes_used": family.get("vault_bytes_used", 0),
        "members_count": members_count,
        "children_count": children_count,
        "young_people_count": young_people_count,
        "limits": limits,
        "price_monthly": catalog["price_monthly"],
        "price_yearly": catalog["price_yearly"],
        "custody": public_custody(family),
    }


def public_user(user: dict) -> dict:
    return {
        "user_id": user["user_id"],
        "email": user["email"],
        "name": user["name"],
        "picture": user.get("picture"),
        "family_id": user["family_id"],
        "language": user.get("language", "en"),
        "is_admin": is_admin_email(user.get("email", "")),
        # Absent field (existing users predating onboarding) reads as True so
        # they never get sent through first-run setup. Only accounts explicitly
        # created with False (new sign-ups) see onboarding.
        "onboarding_completed": bool(user.get("onboarding_completed", True)),
        # Whether this account signs in with a password (vs Google). The delete
        # screen asks for the password when there is one, a typed confirmation
        # when there is not.
        "has_password": bool(user.get("password_hash")),
        # A restricted 13-17 account — the frontend routes these straight to the
        # teen view instead of the full app.
        "is_teen": bool(user.get("is_teen")),
        # A helper (grandparent/carer): uses the normal app but the frontend
        # hides the sensitive surfaces (vault, billing, member management),
        # which the backend also denies via require_full_member.
        "is_helper": bool(user.get("is_helper")),
    }


def public_member(member: dict) -> dict:
    return {
        "member_id": member["member_id"],
        "family_id": member["family_id"],
        "name": member["name"],
        "role": member["role"],
        "avatar": member.get("avatar"),
        # `stars` is the saved bank — the child's whole balance, unchanged by
        # the weekly rhythm. `week_earned` is a meter of stars earned since this
        # week began (Monday), clamped at zero; it gates the weekend treats and
        # drives the progress ring, and never holds stars of its own — nothing
        # is ever lost to a reset.
        "stars": member.get("stars", 0),
        "week_earned": max(0, int(member.get("week_earned", 0) or 0)),
        "weekend_goal_reward_id": member.get("weekend_goal_reward_id"),
        # The rule the week is measured against, and whether it has already
        # been cashed in — held here so the client never keeps its own copy
        # of a number the server enforces.
        "weekly_target": WEEKLY_TARGET,
        "week_claimed": _coerce_dt(member.get("week_claimed_for")) == current_week_start(),
        "has_pin": bool(member.get("pin_hash")),
        "has_account": bool(member.get("user_id")),
        # Null until a parent sets it — children added before this existed have
        # no age, and an invented one would be worse than none.
        "age": member.get("age"),
        # A teen's own user_id is their private chat thread's key. Exposing it to
        # the parent-facing member list (this serializer is only ever returned to
        # full members) lets the app open the right teen's thread by id instead
        # of guessing by display name — two teens named the same no longer
        # collide. Null for a managed child with no account.
        "user_id": member.get("user_id"),
    }


# What a full week of the everyday jobs is worth, and the target it buys.
#
# The three quick jobs come to 7 a day — 49 over seven days, one short of the
# target. That gap was deliberate and wrong: a goal a perfect week cannot reach
# reads as "you did everything and it still is not enough", which is the exact
# discouragement the weekly rhythm exists to avoid. The seventh active day pays
# the last star, so a complete week lands exactly on 50 and the final one is
# plainly for keeping it up.
WEEKLY_TARGET = 50
FULL_WEEK_BONUS = 1
FULL_WEEK_DAYS = 7


def current_week_start():
    """Monday 00:00 UTC of the current week.

    The reset boundary for the weekly earned-meter. UTC rather than the family's
    timezone is a deliberate v1 simplification: it means the week turns over at
    the same instant everywhere, which is a few hours off local midnight for
    some families but never loses or double-counts a star.
    """
    now = utcnow()
    monday = now - timedelta(days=now.weekday())
    return monday.replace(hour=0, minute=0, second=0, microsecond=0)


async def roll_week_if_stale(database, member: dict) -> dict:
    """Reset the weekly earned-meter when a new week has begun.

    Lazy, like the AI-scan period: checked whenever a child's stars are read or
    changed, so no scheduled job is needed. Only the meter resets — the bank
    (`stars`) is untouched, because every earned star was banked when it was
    earned. "Leftover" weekly stars were never a separate pile to lose; they
    are already in the bank, and the meter simply starts counting again.
    """
    start = current_week_start()
    member_start = _coerce_dt(member.get("week_start"))
    if member_start is not None and member_start >= start:
        return member
    await database["family_members"].update_one(
        {"member_id": member["member_id"], "family_id": member["family_id"]},
        {"$set": {"week_earned": 0, "week_start": start}},
    )
    member = {**member, "week_earned": 0, "week_start": start}
    return member


# Structured, not prose: the client composes the sentence in the reader's own
# language, so a French co-parent doesn't read an English activity feed.
ACTIVITY_KINDS = {
    "task_created", "task_done", "task_assigned", "stars_awarded",
    "member_joined", "week_planned", "list_cleared", "doc_shared",
}
ACTIVITY_KEEP = 60


async def log_activity(database, user: dict, kind: str, subject: str = "",
                       amount: Optional[int] = None, target: str = "",
                       shared: bool = True, ref: str = "") -> None:
    """Record who did what, for the household feed.

    The app could always say a task was done; it could never say by whom, so
    a co-parent opening the app found chores mysteriously finished. Best
    effort on purpose — an unrecorded line must never fail the action it
    describes.

    `shared` decides who the line is for. Everything you do is yours: a private
    task's completion is logged so YOU see your own history, but marked
    `shared=False` so a co-parent's feed never carries it. Genuinely-common
    household events (a member joining, the shopping list cleared) default to
    shared. `ref` links the line to the thing it describes — a card id — so
    deleting that thing can take its feed line with it.
    """
    if kind not in ACTIVITY_KINDS:
        return
    try:
        family_id = user.get("family_id")
        if not family_id:
            return
        await database["activity"].insert_one({
            "activity_id": new_id("act"),
            "family_id": family_id,
            "actor_user_id": user.get("user_id"),
            "actor_name": user.get("name") or "",
            "kind": kind,
            "subject": (subject or "")[:80],
            "amount": amount,
            # Who it landed on, for events that are about a person as well as
            # a thing — "Roland gave Keigh the school run".
            "target": (target or "")[:60],
            "shared": bool(shared),
            "ref": (ref or "")[:64],
            # Per-person "hide from my view" for shared lines — the record
            # stays for everyone else, it just leaves your feed.
            "hidden_by": [],
            "created_at": utcnow(),
        })
    except Exception as exc:  # noqa: BLE001 — never break the caller
        log.warning("activity not recorded (%s): %s", kind, exc)


def public_activity(row: dict) -> dict:
    return {
        "activity_id": row["activity_id"],
        "actor_name": row.get("actor_name") or "",
        "actor_user_id": row.get("actor_user_id"),
        "kind": row.get("kind"),
        "subject": row.get("subject") or "",
        "amount": row.get("amount"),
        "target": row.get("target") or "",
        # True for common household lines, False for a private one only the
        # actor sees. Lets the client mark "just you" without a second call.
        "shared": row.get("shared", True) is not False,
        "created_at": iso(row.get("created_at")),
    }


def public_card(card: dict) -> dict:
    return {
        "card_id": card["card_id"],
        "family_id": card["family_id"],
        "type": card["type"],
        "title": card["title"],
        "description": card.get("description"),
        "assignee": card.get("assignee"),
        "due_date": iso(card.get("due_date")),
        "status": card["status"],
        "source": card["source"],
        "image_base64": card.get("image_base64"),
        "recurrence": card.get("recurrence", "none"),
        "reminder_minutes": card.get("reminder_minutes", 60),
        "created_at": iso(card["created_at"]),
        "completed_at": iso(card.get("completed_at")),
        "completed_by_name": card.get("completed_by_name"),
        "google_event_id": card.get("google_event_id"),
        "google_ical_uid": card.get("google_ical_uid"),
        "external_source": card.get("external_source"),
        # Legacy cards (created before per-item privacy) have no owner recorded;
        # they were visible to the whole family, so report them as shared.
        "shared": bool(card["shared"]) if card.get("shared") is not None else card.get("created_by_user_id") is None,
        "created_by_user_id": card.get("created_by_user_id"),
    }


def public_reward(reward: dict) -> dict:
    return {
        "reward_id": reward["reward_id"],
        "family_id": reward["family_id"],
        "title": reward["title"],
        "cost_stars": reward["cost_stars"],
        "icon": reward.get("icon"),
        # A weekend treat is bought with this week's earnings — it needs stars
        # earned since Monday, not just a big enough bank. Everything else is a
        # saved-up reward, paid from the bank with no weekly requirement.
        "weekend": bool(reward.get("weekend", False)),
        "created_at": iso(reward["created_at"]),
    }



def public_redemption(r: dict) -> dict:
    return {
        "redemption_id": r["redemption_id"],
        "family_id": r["family_id"],
        "member_id": r["member_id"],
        "reward_id": r.get("reward_id"),
        # The title is copied, not looked up: a reward can be renamed or
        # deleted after a child has earned it, and what was promised should not
        # change or vanish underneath them.
        "reward_title": r.get("reward_title", ""),
        "reward_icon": r.get("reward_icon"),
        "cost_stars": int(r.get("cost_stars", 0) or 0),
        # A weekly claim costs nothing, so cost_stars alone cannot tell it
        # apart from a free reward — and the two are refunded differently.
        "weekly": bool(r.get("weekly")),
        "status": r.get("status", "pending"),
        "created_at": iso(r["created_at"]),
        "fulfilled_at": iso(r.get("fulfilled_at")) if r.get("fulfilled_at") else None,
        "cancelled_at": iso(r.get("cancelled_at")) if r.get("cancelled_at") else None,
    }


def public_star_transaction(transaction: dict) -> dict:
    return {
        "transaction_id": transaction["transaction_id"],
        "family_id": transaction["family_id"],
        "member_id": transaction["member_id"],
        "delta": transaction["delta"],
        "reason": transaction.get("reason"),
        "created_by_user_id": transaction.get("created_by_user_id"),
        "created_by_name": transaction.get("created_by_name"),
        "created_at": iso(transaction.get("created_at")),
        # The day the star was FOR, which is not always the day it was given.
        # The week row on the Kids screen buckets by this; without it a parent
        # filling in a missed Tuesday on Sunday saw the star land on Sunday
        # while the meter above — which already reads awarded_for — counted it
        # on Tuesday. The row and the meter must describe the same week.
        "awarded_for": iso(transaction.get("awarded_for")) if transaction.get("awarded_for") else None,
    }


def public_vault_doc(doc: dict) -> dict:
    return {
        "doc_id": doc["doc_id"],
        "family_id": doc["family_id"],
        "title": doc["title"],
        "category": doc["category"],
        "image_base64": doc["image_base64"],
        "mime_type": doc.get("mime_type") or "image/jpeg",
        "file_name": doc.get("file_name"),
        # Documents are private to their uploader unless explicitly shared.
        # Legacy docs (uploaded before this existed) carry neither field and
        # stay family-visible until someone claims them — see list_vault.
        "visibility": doc.get("visibility") or "shared",
        "owner_user_id": doc.get("owner_user_id"),
        "owner_name": doc.get("owner_name"),
        "created_at": iso(doc["created_at"]),
    }


def build_invite_url(token: str) -> str:
    base = INVITE_BASE_URL.strip() or "householdcoo:///"
    if "{token}" in base:
        return base.replace("{token}", token)
    joiner = "&" if "?" in base else "?"
    return f"{base}{joiner}invite={token}"


async def send_invite_email(to_email: str, invite_url: str, inviter_name: str, inviter_email: str = "", relationship: str = "") -> dict:
    if not RESEND_API_KEY or not INVITE_FROM_EMAIL:
        return {
            "sent": False,
            "error": "Email delivery is not configured. Set RESEND_API_KEY and INVITE_FROM_EMAIL in Railway.",
        }

    safe_app_name = html.escape(APP_NAME)
    safe_inviter = html.escape(inviter_name or "A family member")
    safe_invite_url = html.escape(invite_url)
    safe_to = html.escape(to_email)
    # The role the inviter typed ("Nanny", "Driver", "Grandma") — when present,
    # the invitee is told up front what they are joining AS, so the invite
    # reads personal rather than generic and expectations are set before the
    # first sign-in.
    rel = re.sub(r"\s+", " ", str(relationship or "").strip())
    safe_rel = html.escape(rel)
    html_lead = (
        f"invited you to join their household as their <strong>{safe_rel}</strong> on"
        if rel else "wants to run the household with you on"
    )

    subject = (
        f"{inviter_name or 'A family member'} invited you to join their household as their {rel} on {APP_NAME}"
        if rel else
        f"{inviter_name or 'A family member'} wants to share the load with you on {APP_NAME}"
    )

    text_lead = (
        f"{inviter_name or 'A family member'} invited you to join their household as their {rel} on {APP_NAME}"
        if rel else
        f"{inviter_name or 'A family member'} wants to run the household with you on {APP_NAME}"
    )
    text = (
        f"{text_lead} — "
        "one shared place for schedules, tasks, the kids' stuff and important documents, so it "
        "doesn't all sit in one person's head.\n\n"
        "Get the app on Google Play, then sign in with this email — your invitation will be "
        "waiting for you inside:\n"
        "https://play.google.com/store/apps/details?id=com.householdcoo.app\n\n"
        f"On an iPhone or a computer, open it in your browser instead:\n{invite_url}\n\n"
        "If you were not expecting this invitation, you can ignore this email."
    )

    html_body = f"""
<div style="font-family: -apple-system, 'Segoe UI', Roboto, Arial, sans-serif; background:#f4f5f2; padding:24px;">
  <div style="max-width:520px; margin:0 auto; background:#ffffff; border:1px solid #e6e1da; border-radius:16px; padding:28px;">
    <p style="color:#202323; font-size:16px; line-height:1.55; margin:0 0 14px;">Hi,</p>
    <p style="color:#202323; font-size:16px; line-height:1.55; margin:0 0 20px;">
      <strong>{safe_inviter}</strong> {html_lead} {safe_app_name} —
      one shared place for schedules, tasks, the kids' stuff and important documents, so it
      doesn't all sit in one person's head. Join in and share the load.
    </p>
    <a href="https://play.google.com/store/apps/details?id=com.householdcoo.app" style="display:inline-block; background:#f26a1b; color:#ffffff; text-decoration:none; font-weight:700; padding:12px 22px; border-radius:10px; font-size:15px;">
      Get the app on Google Play
    </a>
    <p style="color:#4a4f50; font-size:14px; line-height:1.55; margin:18px 0 0;">
      Download the app and sign in with <strong>{safe_to}</strong> — your invitation
      will be waiting for you inside. Just accept it to join.
    </p>
    <p style="color:#747b7c; font-size:13px; line-height:1.5; margin:14px 0 0;">
      On an iPhone or a computer? <a href="{safe_invite_url}" style="color:#b8410a;">Open {safe_app_name} in your browser</a> instead.
    </p>
    <p style="color:#a0a6a7; font-size:12px; line-height:1.5; margin:20px 0 0;">
      If you weren't expecting this, you can safely ignore this email.
    </p>
  </div>
</div>
""".strip()

    payload = {
        "from": INVITE_FROM_EMAIL,
        "to": [to_email],
        "subject": subject,
        "text": text,
        "html": html_body,
    }

    reply_to = INVITE_REPLY_TO or inviter_email
    if reply_to:
        payload["reply_to"] = reply_to
        # A List-Unsubscribe header is a strong positive signal for Gmail/Yahoo
        # inbox placement, even for transactional mail.
        payload["headers"] = {
            "List-Unsubscribe": f"<mailto:{reply_to}?subject=unsubscribe>",
            "List-Unsubscribe-Post": "List-Unsubscribe=One-Click",
        }

    def _send():
        req = urllib.request.Request(
            "https://api.resend.com/emails",
            data=json.dumps(payload).encode("utf-8"),
            headers={
                "Authorization": f"Bearer {RESEND_API_KEY}",
                "Content-Type": "application/json",
                # Resend's edge (Cloudflare) rejects requests with no / a
                # default Python-urllib User-Agent with 403 code 1010, before
                # they reach the API. An explicit UA fixes delivery.
                "User-Agent": f"HouseholdCOO/1.0 (+{APP_NAME})",
            },
            method="POST",
        )

        try:
            with urllib.request.urlopen(req, timeout=15) as response:
                raw = response.read().decode("utf-8")
                try:
                    parsed = json.loads(raw) if raw else {}
                except Exception:
                    parsed = {"raw": raw}
                return {"sent": True, "provider": "resend", "response": parsed}
        except urllib.error.HTTPError as e:
            body = e.read().decode("utf-8", errors="replace")
            return {
                "sent": False,
                "provider": "resend",
                "error": f"Resend HTTP {e.code}: {body}",
            }
        except Exception as e:
            return {"sent": False, "provider": "resend", "error": str(e)}

    return await asyncio.to_thread(_send)



def public_invite(invite: dict) -> dict:
    return {
        "invite_id": invite["invite_id"],
        "family_id": invite["family_id"],
        "email": invite.get("email"),
        "relationship": invite.get("relationship"),
        "label": invite.get("label"),
        "status": invite.get("status", "pending"),
        "token": invite.get("token"),
        "invite_url": build_invite_url(invite["token"]),
        "created_at": iso(invite.get("created_at")),
        "expires_at": iso(invite.get("expires_at")),
        "accepted_at": iso(invite.get("accepted_at")),
        "accepted_by_email": invite.get("accepted_by_email"),
        "created_by_name": invite.get("created_by_name"),
    }


def public_calendar_contact(contact: dict) -> dict:
    return {
        "email": contact.get("email"),
        "name": contact.get("name"),
        "event_count": contact.get("event_count", 0),
        "last_seen_at": iso(contact.get("last_seen_at")),
        "first_seen_at": iso(contact.get("first_seen_at")),
        "last_event_title": contact.get("last_event_title"),
    }


def public_handoff_note(note: dict) -> dict:
    return {
        "note_id": note["note_id"],
        "family_id": note["family_id"],
        "member_id": note.get("member_id"),
        "member_name": note.get("member_name"),
        "text": note["text"],
        "author_name": note.get("author_name", ""),
        "created_at": iso(note["created_at"]),
    }


def public_shopping_item(item: dict) -> dict:
    return {
        "item_id": item["item_id"],
        "family_id": item["family_id"],
        "name": item["name"],
        "category": item.get("category", "Other"),
        "checked": item.get("checked", False),
        "added_by": item.get("added_by", ""),
        "created_at": iso(item["created_at"]),
    }


# Till slips do not print clean shop names. A ticket says "CARREFOUR CITY 14EME",
# "AUCHAN RETAIL FRANCE SAS 0472" or "LIDL SARL 4402", and if we store what is
# printed then six months of shopping shows eleven different Carrefours and the
# per-shop view is worthless. Known chains collapse to one canonical spelling;
# anything else is kept close to what the person typed, because guessing at an
# unknown name is worse than leaving it alone.
_MERCHANT_CHAINS = {
    # France / Belgium
    "carrefour": "Carrefour", "auchan": "Auchan", "leclerc": "E.Leclerc",
    "intermarche": "Intermarché", "monoprix": "Monoprix", "franprix": "Franprix",
    "casino": "Casino", "cora": "Cora", "colruyt": "Colruyt", "delhaize": "Delhaize",
    "picard": "Picard", "grand frais": "Grand Frais", "super u": "Super U",
    "hyper u": "Hyper U", "u express": "U Express",
    # Discounters, everywhere
    "lidl": "Lidl", "aldi": "Aldi", "netto": "Netto", "penny": "Penny",
    "dia": "Dia", "action": "Action",
    # Germany / NL / CH
    "rewe": "Rewe", "edeka": "Edeka", "kaufland": "Kaufland", "dm": "dm",
    "rossmann": "Rossmann", "albert heijn": "Albert Heijn", "jumbo": "Jumbo",
    "migros": "Migros", "coop": "Coop",
    # UK
    "tesco": "Tesco", "sainsbury": "Sainsbury's", "asda": "Asda",
    "morrisons": "Morrisons", "waitrose": "Waitrose", "co-op": "Co-op",
    "marks spencer": "Marks & Spencer", "iceland": "Iceland",
    # Spain / Portugal
    "mercadona": "Mercadona", "carrefour express": "Carrefour",
    "continente": "Continente", "pingo doce": "Pingo Doce",
    # US
    "walmart": "Walmart", "target": "Target", "costco": "Costco",
    "kroger": "Kroger", "trader joe": "Trader Joe\'s", "whole foods": "Whole Foods",
    "safeway": "Safeway", "publix": "Publix", "aldi usa": "Aldi",
    # Not a shop, but it turns up on household receipts constantly
    "amazon": "Amazon", "ikea": "IKEA",
}

MAX_MERCHANT_LEN = 60


def tidy_merchant(raw: Optional[str]) -> Optional[str]:
    """Collapse a printed shop name to the brand behind it."""
    text = sanitize_message_text(str(raw or ""), MAX_MERCHANT_LEN).replace("\n", " ").strip()
    if not text:
        return None

    # Normalise for matching only: lowercase, punctuation and digits to spaces.
    norm = " " + re.sub(r"[^a-z]+", " ", text.lower()).strip() + " "
    # Longest chain key first, so "carrefour express" wins over "carrefour" and
    # "super u" is never mistaken for a stray letter.
    for key in sorted(_MERCHANT_CHAINS, key=len, reverse=True):
        needle = " " + re.sub(r"[^a-z]+", " ", key).strip() + " "
        if needle in norm:
            return _MERCHANT_CHAINS[key]

    # Unknown shop: drop the trailing branch/store number, tidy the shouting.
    cleaned = re.sub(r"[\s\-#]*\d{2,}\s*$", "", text).strip() or text
    return cleaned.title() if cleaned.isupper() else cleaned


def parse_spent_on(value: Optional[str], fallback: datetime) -> datetime:
    """A YYYY-MM-DD from the receipt, kept as date-only at midnight UTC.

    Date-only on purpose: a month total must not move because the family flew to
    another timezone, and the day printed on a ticket has no clock attached to it.
    A date in the future, or one absurdly far back, is refused in favour of the
    fallback rather than quietly landing in a month nobody will look at again.
    """
    if isinstance(value, str) and value.strip():
        try:
            parsed = datetime.strptime(value.strip()[:10], "%Y-%m-%d").replace(tzinfo=timezone.utc)
        except ValueError:
            parsed = None
        if parsed is not None:
            today = fallback.replace(hour=0, minute=0, second=0, microsecond=0)
            if today - timedelta(days=3650) <= parsed <= today + timedelta(days=1):
                return parsed
    return fallback.replace(hour=0, minute=0, second=0, microsecond=0)


def public_expense(exp: dict) -> dict:
    return {
        "expense_id": exp["expense_id"],
        "family_id": exp["family_id"],
        "description": exp["description"],
        "amount": exp["amount"],
        "category": exp.get("category", "General"),
        "child_member_id": exp.get("child_member_id"),
        "child_name": exp.get("child_name"),
        "paid_by_name": exp.get("paid_by_name", ""),
        "paid_by_user_id": exp.get("paid_by_user_id"),
        "merchant": exp.get("merchant"),
        # Expenses recorded before shops existed have no spent_on; the day they
        # were added is the best answer available for those, and it is a true one.
        "spent_on": iso(exp.get("spent_on") or exp["created_at"])[:10],
        "created_at": iso(exp["created_at"]),
    }


def public_template(tmpl: dict) -> dict:
    return {
        "template_id": tmpl["template_id"],
        "family_id": tmpl["family_id"],
        "title": tmpl["title"],
        "description": tmpl.get("description"),
        "recurrence": tmpl.get("recurrence", "daily"),
        "time_of_day": tmpl.get("time_of_day"),
        "assignee": tmpl.get("assignee"),
        "enabled": tmpl.get("enabled", True),
        "created_at": iso(tmpl["created_at"]),
    }


def public_routine(r: dict) -> dict:
    return {
        "routine_id": r["routine_id"],
        "family_id": r["family_id"],
        "name": r["name"],
        "steps": r.get("steps", []),
        "member_id": r.get("member_id"),
        "star_reward": int(r.get("star_reward", 2) or 0),
        "created_at": iso(r["created_at"]),
    }


def public_meal(m: dict) -> dict:
    return {
        "meal_id": m["meal_id"],
        "family_id": m["family_id"],
        "day": m["day"],
        "meal_type": m.get("meal_type", "dinner"),
        "title": m["title"],
        "ingredients": m.get("ingredients", []),
        "notes": m.get("notes"),
        "recipe_id": m.get("recipe_id"),
        # Generated methods, cached per language: {"en": {minutes, steps}, ...}
        "ai_recipe": m.get("ai_recipe") or {},
        "created_at": iso(m["created_at"]),
    }


def public_carpool(c: dict) -> dict:
    return {
        "carpool_id": c["carpool_id"],
        "family_id": c["family_id"],
        "title": c["title"],
        "day_of_week": c["day_of_week"],
        "time": c.get("time", ""),
        "driver_name": c.get("driver_name", ""),
        "pickup_kids": c.get("pickup_kids", []),
        "notes": c.get("notes"),
        "created_at": iso(c["created_at"]),
    }


ALLOWANCE_PERIOD_DAYS = {"weekly": 7, "biweekly": 14, "monthly": 30}


def allowance_next_due(a: dict) -> datetime:
    """When this allowance is next payable.

    Counted from the last payment, or from when it was set up if it has never
    been paid — so a freshly configured allowance is due immediately rather
    than making a parent wait a week before the feature does anything.
    """
    days = ALLOWANCE_PERIOD_DAYS.get(a.get("frequency", "weekly"), 7)
    # ensure_aware_utc, not the raw value: Mongo hands back naive datetimes,
    # utcnow() is aware, and comparing the two raises. This 500ed the very
    # first time a real row — as opposed to a test fixture — went through.
    since = ensure_aware_utc(a.get("last_paid_at") or a.get("created_at"))
    return since + timedelta(days=days) if a.get("last_paid_at") else since


def public_allowance_config(a: dict) -> dict:
    # Setting an amount used to be a note to nobody: the config was stored and
    # never read again, so the balance sat at zero until a parent remembered to
    # record every payment by hand. These two fields are what turn it into
    # something the app can prompt about.
    due = allowance_next_due(a)
    return {
        "allowance_id": a["allowance_id"],
        "family_id": a["family_id"],
        "member_id": a["member_id"],
        "amount": a["amount"],
        "frequency": a.get("frequency", "weekly"),
        "created_at": iso(a["created_at"]),
        "last_paid_at": iso(a["last_paid_at"]) if a.get("last_paid_at") else None,
        "next_due_at": iso(due),
        "is_due": due <= utcnow(),
    }


def public_allowance_txn(t: dict) -> dict:
    return {
        "txn_id": t["txn_id"],
        "family_id": t["family_id"],
        "member_id": t["member_id"],
        "amount": t["amount"],
        "description": t.get("description", ""),
        "txn_type": t.get("txn_type", "deposit"),
        "created_at": iso(t["created_at"]),
    }


def public_announcement(a: dict) -> dict:
    return {
        "announcement_id": a["announcement_id"],
        "family_id": a["family_id"],
        "text": a["text"],
        "author_name": a.get("author_name", ""),
        "priority": a.get("priority", "normal"),
        "created_at": iso(a["created_at"]),
    }


def public_chore(c: dict) -> dict:
    return {
        "chore_id": c["chore_id"],
        "family_id": c["family_id"],
        "title": c["title"],
        "frequency": c.get("frequency", "daily"),
        "assigned_members": c.get("assigned_members", []),
        "current_assignee": c.get("current_assignee"),
        "rotate": c.get("rotate", True),
        # Existing chores predate star rewards; default so old rows still pay.
        "star_reward": int(c.get("star_reward", 3) or 0),
        "last_rotated": iso(c.get("last_rotated")),
        "created_at": iso(c["created_at"]),
    }



def invite_member_role(invite: Optional[dict]) -> str:
    """The role a joining member gets: 'teen' for a teen invite (a restricted
    account), otherwise the free-text relationship the inviter wrote ("Grandma",
    "Nanny", "Brother"...) or the historical default."""
    if (invite or {}).get("is_teen"):
        return "teen"
    if (invite or {}).get("is_helper"):
        return "helper"
    raw = re.sub(r"\s+", " ", str((invite or {}).get("relationship") or "").strip())
    return raw[:32] or "Parent"


def _is_parent_role(role: Optional[str]) -> bool:
    """A parent-level adult: the founder and any co-parent. Everyone else —
    a child, or a grandparent/nanny invited as a family member — is not."""
    return str(role or "").strip().lower() in ("parent", "co-parent")


async def _member_for_user(database: Any, family_id: str, user: dict) -> dict:
    """The signed-in user's own member row, resolved resiliently.

    Matching on user_id alone is not enough: the founder's row can predate the
    day member rows started carrying user_id, so a bare user_id lookup finds
    nothing and the founder reads as "not a parent" — which is exactly how the
    household owner got a 403 trying to remove a co-parent. Falling back to a
    case-insensitive email match recognises them as the parent they are.
    """
    row = await database["family_members"].find_one(
        {"family_id": family_id, "user_id": user.get("user_id")}, {"_id": 0})
    if not row and user.get("email"):
        row = await database["family_members"].find_one(
            {"family_id": family_id,
             "email": {"$regex": f"^{re.escape(user['email'])}$", "$options": "i"}},
            {"_id": 0})
    return row or {}


async def _count_parents(database: Any, family_id: str) -> int:
    return await database["family_members"].count_documents(
        {"family_id": family_id,
         "role": {"$regex": "^(parent|co-parent)$", "$options": "i"}})


MAX_PARENTS = 2  # one parent, one co-parent — the rest are family members with roles.


async def _enforce_parent_limit(database: Any, family_id: str, relationship: Optional[str]) -> None:
    """A household has room for two parents, no more.

    Only co-parent invites are capped — an empty relationship becomes the
    "Parent" role. A family-member invite (Grandparent, Nanny, Sibling...) is a
    named role, not a parent, and is never blocked here. The message names the
    way forward rather than just refusing.
    """
    if _clean_invite_text(relationship):
        return  # a named family member, not a co-parent
    if await _count_parents(database, family_id) >= MAX_PARENTS:
        raise HTTPException(
            status_code=409,
            detail="This household already has two parents. To add someone else, "
                   "invite them as a family member and give them a role.",
        )


async def add_user_to_family_if_needed(database: Any, user: dict, family_id: str,
                                       role: Optional[str] = None):
    is_teen_role = (role or "").strip().lower() == "teen"
    is_helper_role = (role or "").strip().lower() == "helper"
    # Set the restriction flag FIRST, before any early return, and on the passed
    # dict too (the auth handlers serialize it right after) — so a helper is
    # locked down from their first request, never routed into the full app.
    if is_helper_role:
        await database["users"].update_one(
            {"user_id": user["user_id"]}, {"$set": {"is_helper": True, "updated_at": utcnow()}})
        user["is_helper"] = True
    # Set the restriction flag FIRST, before any early return: a teen who is
    # already a family member must still be locked down (fail closed, not open).
    if is_teen_role:
        await database["users"].update_one(
            {"user_id": user["user_id"]}, {"$set": {"is_teen": True, "updated_at": utcnow()}})
        # Mutate the in-memory dict too: the auth handlers serialize this same
        # object with public_user() right after calling us, and without this the
        # join response would carry is_teen:false — routing the teen into the
        # full parent app where every request then 403s until a relaunch.
        user["is_teen"] = True
        # Retire the managed child profile this teen is replacing — same name,
        # no login — so they don't appear (or count) twice.
        teen_name = (user.get("name") or "").strip().lower()
        async for m in database["family_members"].find({"family_id": family_id}, {"_id": 0}):
            if (m.get("role") or "").strip().lower() == "child" and not m.get("user_id") \
               and (m.get("name") or "").strip().lower() == teen_name and teen_name:
                # Last line of defence, at the moment the account actually
                # appears. An invite made before this check existed, or one made
                # without naming the child, still cannot turn a profile recorded
                # as under 13 into an account holder.
                recorded = m.get("age")
                if recorded is not None and int(recorded) < 13:
                    raise HTTPException(
                        status_code=400,
                        detail="That person is recorded as under 13. An account of their own "
                               "is for 13 and over.")
                await database["family_members"].delete_one({"member_id": m["member_id"]})

    existing = await database["family_members"].find_one(
        {
            "family_id": family_id,
            "$or": [
                {"user_id": user["user_id"]},
                {"email": user.get("email", "")},
            ],
        },
        {"_id": 0},
    )
    if existing:
        return existing

    # Backstop the young-people cap at accept time, the way the parent cap is
    # backstopped below: an invite sent while there was room can't be accepted
    # once the household has filled up. The friendly gate is at invite time
    # (family_invite); this holds the line if that was bypassed or raced.
    if is_teen_role:
        sub = await build_subscription(family_id)
        max_young = sub["limits"].get("max_children", 2)
        if await count_young_people(database, family_id) >= max_young:
            raise HTTPException(
                status_code=402,
                detail={"error": "plan_limit", "feature": "max_children",
                        "message": "This household is full for its current plan.",
                        "limit": max_young},
            )

    # The real guard: two co-parent invites can be outstanding at once (each was
    # created while only the founder existed), and without this the second to be
    # accepted would quietly become a third parent. The invite-time check is the
    # friendly pop-up; this is the one that actually holds the line.
    if _is_parent_role(role) and await _count_parents(database, family_id) >= MAX_PARENTS:
        raise HTTPException(
            status_code=409,
            detail="This household already has two parents, so this invite can no "
                   "longer be accepted as a co-parent.",
        )

    member = {
        "member_id": new_id("member"),
        "family_id": family_id,
        "user_id": user["user_id"],
        "email": user.get("email", ""),
        "name": user.get("name") or user.get("email") or "Parent",
        "role": role or "Parent",
        "avatar": user.get("picture"),
        "stars": 0,
        "pin_hash": None,
        "created_at": utcnow(),
    }
    await database["family_members"].insert_one(member)
    return member


def alerts_enabled(prefs: Optional[dict], key: str) -> bool:
    """Whether this person wants a given alert. Absent settings mean YES.

    So does a row the user has never edited. Those rows were written by the
    server with everything switched off, which meant a family could install the
    app, grant notification permission, and then hear nothing at all - a
    co-parent assigns the school run and the other parent finds out days later
    by noticing their own name. That is the exact failure this app exists to
    prevent, and it was the default.

    Defaulting to on is not the app being pushy: the operating system still asks
    permission before a single notification is delivered, and a denied prompt
    silences everything regardless of what is stored here. This setting only
    decides what happens once someone has already said yes.

    A row the user HAS edited is obeyed exactly, in both directions - turning an
    alert off stays off. update_notification_settings always stamps updated_at,
    so created_at == updated_at is a reliable "never chosen".
    """
    if not prefs:
        return True
    created, updated = prefs.get("created_at"), prefs.get("updated_at")
    if created and updated and created == updated:
        return True
    return bool(prefs.get(key, True))


def public_notification_settings(settings: Optional[dict]) -> dict:
    # The screen must show what the server will actually do, or a family sees
    # "off" while alerts arrive - or worse, "on" while nothing does.
    return {
        "card_reminders": alerts_enabled(settings, "card_reminders"),
        "new_card_alerts": alerts_enabled(settings, "new_card_alerts"),
        "updated_at": iso((settings or {}).get("updated_at")),
    }


async def get_notification_settings_doc(user_id: str) -> dict:
    database = get_db()
    settings = await database["notification_settings"].find_one(
        {"user_id": user_id},
        {"_id": 0},
    )
    if settings:
        return settings

    settings = {
        "user_id": user_id,
        "card_reminders": True,
        "new_card_alerts": True,
        "created_at": utcnow(),
        "updated_at": utcnow(),
    }
    await database["notification_settings"].insert_one(settings)
    return settings


def _latest_token_per_user(docs: list[dict]) -> list[dict]:
    """One push token per person — the most recently registered.

    A device's Expo token changes across reinstalls and some updates, and the
    old value is left in the table as `active`. Nothing here ever deactivated
    it, so a co-parent who had reinstalled a few times accumulated several
    live-looking tokens, and every family push fanned out to all of them:
    one event, a handful of notifications. Collapsing to the newest token per
    user makes it one person, one push. Dead tokens are separately retired
    when Expo reports them (see send_expo_push_messages).
    """
    EPOCH = datetime(1970, 1, 1, tzinfo=timezone.utc)
    best: dict[str, dict] = {}
    for d in docs:
        uid = d.get("user_id")
        if not uid:
            continue
        cur = best.get(uid)
        if cur is None or (_coerce_dt(d.get("updated_at")) or EPOCH) >= (
                _coerce_dt(cur.get("updated_at")) or EPOCH):
            best[uid] = d
    return list(best.values())


async def _deactivate_dead_tokens(database, tokens: list[str]) -> None:
    """Retire tokens Expo says are gone, so they stop being pushed to."""
    for token in tokens:
        try:
            await database["notification_tokens"].update_one(
                {"token": token}, {"$set": {"active": False, "updated_at": utcnow()}})
        except Exception:
            pass


async def send_expo_push_messages(messages: list[dict], database=None) -> dict:
    if not messages:
        return {"sent": 0, "skipped": True}

    def _send():
        req = urllib.request.Request(
            "https://exp.host/--/api/v2/push/send",
            data=json.dumps(messages).encode("utf-8"),
            headers={
                "Content-Type": "application/json",
                "Accept": "application/json",
            },
            method="POST",
        )
        try:
            with urllib.request.urlopen(req, timeout=12) as response:
                raw = response.read().decode("utf-8")
                try:
                    parsed = json.loads(raw) if raw else {}
                except Exception:
                    parsed = {"raw": raw}
                return {"sent": len(messages), "response": parsed}
        except urllib.error.HTTPError as e:
            body = e.read().decode("utf-8", errors="replace")
            return {"sent": 0, "error": f"Expo push HTTP {e.code}: {body}"}
        except Exception as e:
            return {"sent": 0, "error": str(e)}

    result = await asyncio.to_thread(_send)

    # Expo returns one ticket per message, in order. A DeviceNotRegistered
    # ticket means that token is dead — retire it so it stops fanning out
    # future pushes. Best effort; never let cleanup break the send.
    if database is not None:
        try:
            tickets = (result.get("response") or {}).get("data") or []
            dead = []
            for msg, ticket in zip(messages, tickets):
                if isinstance(ticket, dict) and ticket.get("status") == "error" \
                        and (ticket.get("details") or {}).get("error") == "DeviceNotRegistered":
                    if msg.get("to"):
                        dead.append(msg["to"])
            if dead:
                await _deactivate_dead_tokens(database, dead)
        except Exception:
            pass

    return result


STAR_MILESTONE = 50


async def send_star_milestone_alert(family_id: str, member_name: str, old_total: int, new_total: int):
    """Notify the family's devices when a child crosses a 50-star milestone."""
    try:
        if new_total // STAR_MILESTONE <= old_total // STAR_MILESTONE:
            return
        milestone = (new_total // STAR_MILESTONE) * STAR_MILESTONE
        database = get_db()
        messages = []
        docs = [d async for d in database["notification_tokens"].find(
            {"family_id": family_id, "active": True}, {"_id": 0})]
        for token_doc in _latest_token_per_user(docs):
            token = token_doc.get("token")
            if not token or not token.startswith("ExponentPushToken"):
                continue
            messages.append({
                "to": token,
                "sound": "default",
                "title": f"{member_name} reached {milestone} stars!",
                "body": "Amazing work — time to celebrate with a reward?",
                "data": {"type": "star_milestone", "family_id": family_id},
            })
        if messages:
            await send_expo_push_messages(messages, database)
    except Exception as e:
        log.warning("star milestone alert failed: %s", e)


# Pushes around the invitation loop, in the recipient's own language.
PUSH_I18N = {
    "en": {
        "invited_title": "{name} invited you to their household",
        "invited_body": "Open Ahenora and sign in through the invite link to join.",
        "accepted_title": "{name} accepted your invitation",
        "accepted_body": "They have joined your household.",
        "assigned_title": "{name} handed you something",
        "assigned_body": "{title}",
        "assigned_body_due": "{title} — due {due}",
        "teen_star_title": "Your star is approved ⭐",
        "teen_star_body": "{title}",
        "reminder_title": "Reminder ⏰",
        "reminder_body": "{title} — due {due}",
    },
    "fr": {
        "invited_title": "{name} vous invite dans son foyer",
        "invited_body": "Ouvrez Ahenora et connectez-vous via le lien d'invitation pour le rejoindre.",
        "accepted_title": "{name} a accepté votre invitation",
        "accepted_body": "Cette personne a rejoint votre foyer.",
        "assigned_title": "{name} vous a confié quelque chose",
        "assigned_body": "{title}",
        "assigned_body_due": "{title} — pour le {due}",
        "teen_star_title": "Ton étoile est validée ⭐",
        "teen_star_body": "{title}",
        "reminder_title": "Rappel ⏰",
        "reminder_body": "{title} — pour le {due}",
    },
    "es": {
        "invited_title": "{name} te invitó a su hogar",
        "invited_body": "Abre Ahenora e inicia sesión con el enlace de invitación para unirte.",
        "accepted_title": "{name} aceptó tu invitación",
        "accepted_body": "Ya forma parte de tu hogar.",
        "assigned_title": "{name} te ha encargado algo",
        "assigned_body": "{title}",
        "assigned_body_due": "{title} — para el {due}",
        "teen_star_title": "Tu estrella está aprobada ⭐",
        "teen_star_body": "{title}",
        "reminder_title": "Recordatorio ⏰",
        "reminder_body": "{title} — para el {due}",
    },
    "de": {
        "invited_title": "{name} hat dich in den Haushalt eingeladen",
        "invited_body": "Öffne Ahenora und melde dich über den Einladungslink an, um beizutreten.",
        "accepted_title": "{name} hat deine Einladung angenommen",
        "accepted_body": "Die Person ist deinem Haushalt beigetreten.",
        "assigned_title": "{name} hat dir etwas übergeben",
        "assigned_body": "{title}",
        "assigned_body_due": "{title} — fällig am {due}",
        "teen_star_title": "Dein Stern ist bestätigt ⭐",
        "teen_star_body": "{title}",
        "reminder_title": "Erinnerung ⏰",
        "reminder_body": "{title} — fällig am {due}",
    },
}


async def send_push_to_user(database, user_id: str, title: str, body: str, data: dict):
    """Push to one specific person's devices. Best effort, never raises."""
    try:
        docs = [d async for d in database["notification_tokens"].find(
            {"user_id": user_id, "active": True}, {"_id": 0})]
        messages = []
        # One device per person: newest token only, so accumulated stale
        # tokens from past installs don't each fire their own notification.
        for token_doc in _latest_token_per_user(docs):
            token = token_doc.get("token")
            if token and token.startswith("ExponentPushToken"):
                messages.append({
                    "to": token, "sound": "default",
                    "title": title, "body": body, "data": data,
                })
        if messages:
            await send_expo_push_messages(messages, database)
    except Exception as e:
        log.warning("user push failed: %s", e)


async def notify_invite_accepted(database, invite: dict, acceptor_name: str):
    """Close the loop: the person who sent an invite hears when it lands."""
    inviter_id = invite.get("created_by_user_id")
    if not inviter_id:
        return
    inviter = await database["users"].find_one({"user_id": inviter_id}, {"_id": 0})
    lang = (inviter or {}).get("language") or "en"
    L = PUSH_I18N.get(lang, PUSH_I18N["en"])
    await send_push_to_user(
        database, inviter_id,
        L["accepted_title"].format(name=acceptor_name),
        L["accepted_body"],
        {"type": "invite_accepted", "invite_id": invite.get("invite_id")},
    )


async def resolve_member_user_id(database, family_id: str, name: str) -> Optional[str]:
    """The user behind an assignee name, if that name belongs to a grown-up
    with a login. Children are members but not users, so they resolve to None."""
    wanted = (name or "").strip().lower()
    if not wanted:
        return None
    cursor = database["family_members"].find({"family_id": family_id}, {"_id": 0})
    async for member in cursor:
        if (member.get("name") or "").strip().lower() == wanted:
            return member.get("user_id")
    return None


async def notify_assignment(database, actor: dict, card: dict, assignee_name: str) -> None:
    """Tell someone a job just landed on them.

    Hand-off was the quiet failure in this app: one parent would assign
    something and the other would find out days later, by opening the app and
    noticing their own name — or not noticing. Writing a name into a field is
    not communication.

    Silent in two cases, both deliberate: assigning something to yourself,
    which needs no announcement, and a private item, which is nobody else's
    business until it is shared.
    """
    try:
        if not card.get("shared"):
            return
        target_id = await resolve_member_user_id(database, card["family_id"], assignee_name)
        if not target_id or target_id == actor.get("user_id"):
            return
        target = await database["users"].find_one({"user_id": target_id}, {"_id": 0})
        lang = (target or {}).get("language") or "en"
        L = PUSH_I18N.get(lang, PUSH_I18N["en"])
        title = (card.get("title") or "").strip()
        due = ensure_aware_utc(card.get("due_date"))
        body = (
            L["assigned_body_due"].format(title=title, due=due.strftime("%d/%m"))
            if due else L["assigned_body"].format(title=title)
        )
        await send_push_to_user(
            database, target_id,
            L["assigned_title"].format(name=actor.get("name") or "Someone"),
            body,
            {"type": "task_assigned", "card_id": card.get("card_id")},
        )
    except Exception as e:
        log.warning("assignment notification failed: %s", e)


async def send_new_card_alert(family_id: str, card: dict, created_by_user_id: Optional[str] = None,
                              exclude_user_ids: Optional[set] = None):
    database = get_db()
    messages = []
    # The creator never needs telling, and anyone excluded (the assignee) is
    # getting the more specific hand-off push instead — sending both would be
    # two notifications for one event.
    skip = set(exclude_user_ids or ())
    if created_by_user_id:
        skip.add(created_by_user_id)

    docs = [d async for d in database["notification_tokens"].find(
        {"family_id": family_id, "active": True}, {"_id": 0})]

    for token_doc in _latest_token_per_user(docs):
        uid = token_doc.get("user_id")
        if uid in skip:
            continue

        prefs = await database["notification_settings"].find_one(
            {"user_id": uid}, {"_id": 0})
        if not alerts_enabled(prefs, "new_card_alerts"):
            continue

        token = token_doc.get("token")
        if not token or not token.startswith("ExponentPushToken"):
            continue

        messages.append(
            {
                "to": token,
                "sound": "default",
                "title": "New Ahenora card",
                "body": card.get("title") or "A new card was added.",
                "data": {
                    "type": "new_card",
                    "card_id": card.get("card_id"),
                    "family_id": family_id,
                },
            }
        )

    if messages:
        await send_expo_push_messages(messages, database)


async def send_coparent_alert(family_id: str, title: str, body: str, data_type: str, created_by_user_id: Optional[str] = None):
    """Notify the OTHER family members when a co-parent does something worth
    seeing (a note, an announcement, …) — the two-sided pull that makes a shared
    household app worth opening daily. Respects each user's household-alert
    preference and never notifies the author of their own action."""
    database = get_db()
    messages = []
    preview = (body or "").strip()
    if len(preview) > 120:
        preview = preview[:117].rstrip() + "…"

    docs = [d async for d in database["notification_tokens"].find(
        {"family_id": family_id, "active": True}, {"_id": 0})]
    for token_doc in _latest_token_per_user(docs):
        uid = token_doc.get("user_id")
        if created_by_user_id and uid == created_by_user_id:
            continue
        prefs = await database["notification_settings"].find_one(
            {"user_id": uid}, {"_id": 0})
        if not alerts_enabled(prefs, "new_card_alerts"):
            continue
        token = token_doc.get("token")
        if not token or not token.startswith("ExponentPushToken"):
            continue
        messages.append(
            {
                "to": token,
                "sound": "default",
                "title": title,
                "body": preview or "Open Ahenora to see it.",
                "data": {"type": data_type, "family_id": family_id},
            }
        )

    if messages:
        await send_expo_push_messages(messages, database)


# -----------------------------------------------------------------------------
# Server-side reminder scheduler
# -----------------------------------------------------------------------------
# Until now a card's reminder was scheduled ON THE DEVICE (expo-notifications, a
# local DATE trigger), so it only ever fired if that person had opened the app
# to schedule it — never for an unopened app, after a reboot, or for a co-parent
# who never loaded the card. This is the server half: one process (uvicorn runs a
# single worker) wakes each minute, finds the reminders that have come due, and
# sends the push itself. Idempotent per occurrence, so a restart or a rolling
# deploy's brief container overlap cannot double-send.
REMINDER_SCAN_INTERVAL = int(os.environ.get("REMINDER_SCAN_INTERVAL", "60"))
# Don't blast a backlog: on first deploy (or after a long outage) there are many
# OPEN cards whose reminder time is well in the past. Those are marked handled
# but not actually sent — only reminders that came due within this window fire.
REMINDER_CATCHUP = timedelta(minutes=int(os.environ.get("REMINDER_CATCHUP_MINUTES", "180")))
REMINDER_SCHEDULER_ENABLED = os.environ.get("REMINDER_SCHEDULER", "on").strip().lower() not in (
    "off", "0", "false", "no")


async def _reminder_recipients(database, card: dict) -> list[str]:
    """Who hears a card's due reminder.

    A private card is the creator's alone. A shared card assigned to a grown-up
    goes to that grown-up — the person on the hook. A shared card nobody has been
    given (or one assigned to a child, who has no device) goes to the whole
    household, so an unclaimed 'sign the slip' is not everyone's-and-no-one's.
    """
    if not card.get("shared"):
        uid = card.get("created_by_user_id")
        return [uid] if uid else []
    assignee = (card.get("assignee") or "").strip()
    if assignee:
        uid = await resolve_member_user_id(database, card["family_id"], assignee)
        if uid:
            return [uid]
    docs = [d async for d in database["notification_tokens"].find(
        {"family_id": card["family_id"], "active": True}, {"_id": 0})]
    seen, out = set(), []
    for token_doc in _latest_token_per_user(docs):
        uid = token_doc.get("user_id")
        if uid and uid not in seen:
            seen.add(uid)
            out.append(uid)
    return out


async def send_due_card_reminders(database, now: Optional[datetime] = None) -> int:
    """Fire push reminders for cards that have reached due_date − reminder_minutes.

    Idempotent per occurrence: the ISO due_date the reminder fired for is stored
    on the card as reminder_sent_for, and a card is skipped once it matches. A
    recurring card spawns a fresh doc for its next occurrence, so it earns its
    own reminder; moving a card's due date changes the occurrence key, so it is
    reminded again for the new time. Best effort — one bad card never stops the
    rest. Returns how many pushes were sent.
    """
    now = now or utcnow()
    sent = 0
    cursor = database["cards"].find({"status": "OPEN"}, {"_id": 0})
    async for card in cursor:
        try:
            due = ensure_aware_utc(card.get("due_date"))
            mins = card.get("reminder_minutes") or 0
            if not due or mins <= 0:
                continue
            trigger = due - timedelta(minutes=mins)
            if now < trigger:
                continue  # not due yet
            occurrence = iso(due)
            if card.get("reminder_sent_for") == occurrence:
                continue  # already handled this occurrence
            # Mark handled BEFORE sending: a crash mid-fan-out must not loop and
            # re-ping the people already reached on the next tick.
            await database["cards"].update_one(
                {"card_id": card["card_id"]},
                {"$set": {"reminder_sent_for": occurrence}},
            )
            # Too old to be useful — a first-deploy backlog or a long outage.
            # Marked handled above so it never fires late; just don't send it.
            if now - trigger > REMINDER_CATCHUP:
                continue
            title_txt = (card.get("title") or "").strip()
            due_txt = due.strftime("%d/%m %H:%M")
            for uid in await _reminder_recipients(database, card):
                prefs = await database["notification_settings"].find_one(
                    {"user_id": uid}, {"_id": 0})
                if not alerts_enabled(prefs, "card_reminders"):
                    continue
                target = await database["users"].find_one({"user_id": uid}, {"_id": 0})
                lang = (target or {}).get("language") or "en"
                L = PUSH_I18N.get(lang, PUSH_I18N["en"])
                await send_push_to_user(
                    database, uid,
                    L["reminder_title"],
                    L["reminder_body"].format(title=title_txt, due=due_txt),
                    {"type": "card_reminder", "card_id": card.get("card_id")},
                )
                sent += 1
        except Exception as e:
            log.warning("due reminder failed for %s: %s", card.get("card_id"), e)
    return sent


async def _reminder_scheduler_loop():
    while True:
        try:
            await send_due_card_reminders(get_db())
        except Exception as e:  # a tick must never kill the loop
            log.warning("reminder scheduler tick failed: %s", e)
        await asyncio.sleep(REMINDER_SCAN_INTERVAL)


@app.on_event("startup")
async def start_reminder_scheduler():
    # Only under a real server with a database — a plain import (the test suite)
    # never fires startup, so this never runs there. Off switch: REMINDER_
    # SCHEDULER=off, in case it ever needs killing without a code change.
    if db is None or not REMINDER_SCHEDULER_ENABLED:
        return
    asyncio.create_task(_reminder_scheduler_loop())
    log.info("Reminder scheduler started (every %ss)", REMINDER_SCAN_INTERVAL)


async def require_user(authorization: str = Header(default=""),
                       x_client_platform: str = Header(default="")):
    database = get_db()

    if not authorization.startswith("Bearer "):
        raise HTTPException(status_code=401, detail="Missing bearer token")

    token = authorization.replace("Bearer ", "", 1).strip()
    session = await database["user_sessions"].find_one(
        {"token_hash": sha256(token), "expires_at": {"$gt": utcnow()}},
        {"_id": 0},
    )
    if not session:
        raise HTTPException(status_code=401, detail="Invalid or expired session")

    # A child's session is deliberately worthless here. Kid access rides on the
    # parent's account, so a child token that satisfied require_user would
    # inherit the whole household — the vault, the calendar, the co-parent's
    # private items. Refusing it in the one place every route already depends
    # on means the ~150 endpoints below are all protected by default, and a
    # new endpoint written next month is protected without anyone remembering
    # to think about it. Fail closed, once, centrally.
    if session.get("kind") == "child":
        raise HTTPException(status_code=403, detail="Not available in kid mode")

    # Using the app keeps you signed in — but only past the kid-mode refusal
    # above. Renewing first meant a hand-over session on someone else's phone
    # kept itself alive forever on the strength of its own rejected requests.
    await _touch_session(database, sha256(token), session)

    user = await database["users"].find_one({"user_id": session["user_id"]}, {"_id": 0})
    if not user:
        raise HTTPException(status_code=401, detail="User not found")

    # A teen has a real login but a restricted world. Like a kid session, their
    # token is worthless on every parent route — refusing it here, in the one
    # gate all ~150 endpoints share, means a teen can never read the family
    # calendar, vault, other members or settings, even on an endpoint written
    # next month. Teens live entirely on the /api/teen/* allowlist below.
    if user.get("is_teen"):
        raise HTTPException(status_code=403, detail="teen_mode")

    # Daily-active tracking, throttled to one write per user per day.
    today = utcnow().strftime("%Y-%m-%d")
    if user.get("last_active_day") != today:
        try:
            await database["users"].update_one(
                {"user_id": user["user_id"]},
                {"$set": {"last_active_day": today}},
            )
            await database["metrics_daily"].update_one(
                {"date": today, "name": "active_users"},
                {"$inc": {"count": 1}},
                upsert=True,
            )
            # ...and the same count split by platform, so "web vs app" is
            # answerable. Web users can't buy through the store, so if the daily
            # actives are mostly web that alone explains a flat subscriber line.
            plat = (x_client_platform or "").strip().lower()
            plat = plat if plat in ("web", "android", "ios") else "other"
            await database["metrics_daily"].update_one(
                {"date": today, "name": f"active_{plat}"},
                {"$inc": {"count": 1}},
                upsert=True,
            )
        except Exception:
            pass  # metrics must never break auth

    return user


async def require_full_member(user=Depends(require_user)):
    """Parent/co-parent only — the gate for the sensitive surfaces a helper must
    not reach: the document vault, billing, member management, invites and
    expenses. A helper (grandparent/carer) uses the normal app for the shared
    day but is refused here, deny-by-default. (Teens never reach this: require_
    user already 403s a teen token, so this only has to stop helpers.)"""
    if user.get("is_helper"):
        raise HTTPException(status_code=403, detail="helper_mode")
    return user


# -----------------------------------------------------------------------------
# Schemas
# -----------------------------------------------------------------------------
class SessionIn(BaseModel):
    session_id: str
    invite_token: Optional[str] = None
    language: Optional[str] = None  # device/browser language for a new account



class EmailRegisterIn(BaseModel):
    name: str
    email: str
    password: str
    invite_token: Optional[str] = None
    # The device/browser language the app opened in, so a new account is created
    # in it (a French sign-up stays French on their next device too).
    language: Optional[str] = None


class EmailLoginIn(BaseModel):
    email: str
    password: str
    invite_token: Optional[str] = None


class InviteIn(BaseModel):
    email: str
    # Free text from the inviter: "Grandma", "Nanny", "Brother"... Becomes
    # the member's displayed role when the invite is accepted.
    relationship: Optional[str] = None
    # A 13-17 teen gets their own login but a restricted, teen-only view.
    # Under-13 never reaches here — they stay a managed child profile (no
    # account, no email), which is what keeps us COPPA / Families-policy clean.
    is_teen: Optional[bool] = False
    age: Optional[int] = None
    # Which child this invite is for, when it was started from their page. It is
    # what lets the age claimed here be checked against the age already on
    # record, instead of taking the typed number on trust.
    member_id: Optional[str] = None
    # A helper is a trusted adult with a limited view — a grandparent, nanny or
    # babysitter. A real family member (uses the normal app), but denied the
    # sensitive surfaces (vault, billing, member management) server-side.
    is_helper: Optional[bool] = False


class LanguageIn(BaseModel):
    language: str


class PinIn(BaseModel):
    pin: str


class AiAssignIn(BaseModel):
    title: str
    description: str = ""
    type: str = "TASK"


# A photo the app captures at quality 0.55 is well under a megabyte; the vault
# caps uploads against the storage quota, but a JSON body is buffered whole in
# memory before any quota check can run, so an image field with no ceiling is a
# body sized to exhaust memory. ~10 MB of image is ~14 MB of base64 — generous
# for a phone photo, and a firm wall rejected at validation time.
MAX_IMAGE_B64_CHARS = 14 * 1024 * 1024


class CardIn(BaseModel):
    type: str = "TASK"
    title: str
    description: Optional[str] = None
    assignee: Optional[str] = None
    due_date: Optional[str] = None
    source: str = "MANUAL"
    image_base64: Optional[str] = Field(default=None, max_length=MAX_IMAGE_B64_CHARS)
    recurrence: str = "none"
    reminder_minutes: int = 60
    # Shared unless someone says otherwise. It was private by default, which
    # inverted the whole product: a task added with the + button was invisible
    # to the rest of the household and notified nobody, so a co-parent could
    # add the school run and the other parent would never see it or hear about
    # it. The two failures are not equal - sharing something you meant to keep
    # back is visible and reversible in one tap, while a task that silently
    # reaches nobody is never discovered at all.
    shared: bool = True


class CardPatchIn(BaseModel):
    type: Optional[str] = None
    title: Optional[str] = None
    description: Optional[str] = None
    assignee: Optional[str] = None
    due_date: Optional[str] = None
    status: Optional[str] = None
    recurrence: Optional[str] = None
    reminder_minutes: Optional[int] = None
    shared: Optional[bool] = None


class VaultIn(BaseModel):
    title: str
    category: str
    image_base64: str = Field(max_length=MAX_IMAGE_B64_CHARS)
    mime_type: Optional[str] = None
    file_name: Optional[str] = None
    # "private" (default) or "shared". A co-parent joining a household must
    # not inherit sight of documents nobody chose to share.
    visibility: Optional[str] = None


class VaultVisibilityIn(BaseModel):
    visibility: str


class RewardIn(BaseModel):
    title: str
    # PATCH already refuses anything under 1; create did not, and a negative
    # cost inverts the redemption maths into handing out stars.
    cost_stars: int = Field(ge=1, le=100000)
    icon: Optional[str] = None
    weekend: bool = False


class ChildIn(BaseModel):
    name: str
    starting_stars: int = 0
    pin: Optional[str] = None


class MemberPatchIn(BaseModel):
    """Only what a parent can safely correct in place.

    Stars are deliberately absent: they move through the audited endpoint that
    writes a ledger entry, and a silent $set here would break the balance the
    history is supposed to explain.
    """

    name: Optional[str] = None
    avatar: Optional[str] = None
    # A child's age, set after the fact: children are added in a hurry at setup
    # and the age was never asked for, so it had to be correctable later. Send
    # 0 to clear one that was set by mistake.
    age: Optional[int] = None


class StarAdjustmentIn(BaseModel):
    delta: int
    reason: Optional[str] = None
    # The day the job was done, when that is not today. Bounded to the current
    # week by the handler: editing a settled week would move a meter that has
    # already paid out a treat.
    awarded_for: Optional[str] = None


class WeeklyClaimIn(BaseModel):
    """What the week is being cashed in for. Free text, because the ideas are
    a family's own — "Bowling with Dad" is as valid as anything we ship."""
    title: str = Field(max_length=80)


class RewardPatchIn(BaseModel):
    title: Optional[str] = None
    cost_stars: Optional[int] = None
    icon: Optional[str] = None
    weekend: Optional[bool] = None


class RedeemIn(BaseModel):
    member_id: str


class SubscriptionChangeIn(BaseModel):
    plan: str
    billing_cycle: str


class CustodyConfigIn(BaseModel):
    # Alternating custody (garde alternée). enabled turns the Feed/Calendar
    # markers on; our_weeks is the parity of ISO week the children are in this
    # home ("even"/"odd", the paire/impaire of a French judgment); away_label is
    # an optional name for the other home ("Chez leur papa"), blank for generic.
    enabled: bool
    our_weeks: str
    away_label: str = Field(default="", max_length=CUSTODY_LABEL_MAX)


class VisionIn(BaseModel):
    image_base64: str = Field(max_length=MAX_IMAGE_B64_CHARS)

class CalendarImportIn(BaseModel):
    access_token: str
    days: int = 30


class NotificationTokenIn(BaseModel):
    token: str
    platform: Optional[str] = None
    # What build and runtime this device is on, for OTA-adoption reporting.
    # Optional so an older client that does not send them still registers.
    app_version: Optional[str] = None
    runtime_version: Optional[str] = None


class NotificationPrefsIn(BaseModel):
    card_reminders: Optional[bool] = None
    new_card_alerts: Optional[bool] = None


class HandoffNoteIn(BaseModel):
    member_id: Optional[str] = None
    text: str


class ShoppingItemIn(BaseModel):
    name: str
    category: str = "Other"


class ShoppingItemPatchIn(BaseModel):
    checked: Optional[bool] = None
    name: Optional[str] = None
    category: Optional[str] = None


class ExpenseIn(BaseModel):
    description: str = ""
    # A negative expense silently skews the split totals; nobody spends a
    # negative amount, and an upper bound keeps a fat-fingered figure out of
    # the aggregation.
    amount: float = Field(ge=0, le=1_000_000)
    category: str = "General"
    child_member_id: Optional[str] = None
    # Where it was spent. Optional, because a school trip has no shop.
    merchant: Optional[str] = None
    # The date on the receipt, YYYY-MM-DD. People photograph a pile of tickets on
    # a Sunday evening, so "when it was added" and "when it was spent" are not the
    # same day, and only one of them belongs in a monthly total.
    spent_on: Optional[str] = None


class TemplateIn(BaseModel):
    title: str
    description: Optional[str] = None
    recurrence: str = "daily"
    time_of_day: Optional[str] = None
    assignee: Optional[str] = None


class RoutineIn(BaseModel):
    name: str
    steps: list  # [{"label": str, "duration_seconds": int}]
    member_id: Optional[str] = None
    star_reward: int = 2


class RoutinePatchIn(BaseModel):
    name: Optional[str] = None
    steps: Optional[list] = None


class MealPlanIn(BaseModel):
    day: str  # "monday", "tuesday", etc.
    meal_type: str = "dinner"  # breakfast, lunch, dinner, snack
    title: str
    ingredients: list = []  # [str]
    notes: Optional[str] = None
    # Set when the meal came from the built-in suggestion library. The client
    # re-renders the title from that library in the current language, so a
    # plan created in French still reads correctly after switching to English.
    # Absent for meals the user typed themselves — those keep `title` verbatim.
    recipe_id: Optional[str] = None


class CarpoolIn(BaseModel):
    title: str
    day_of_week: str  # "monday" etc.
    time: str  # "08:00"
    driver_name: str
    pickup_kids: list = []  # [str] kid names
    notes: Optional[str] = None


class AllowanceIn(BaseModel):
    member_id: str
    amount: float
    frequency: str = "weekly"  # weekly, biweekly, monthly


class AllowanceTxnIn(BaseModel):
    member_id: str
    amount: float
    description: str
    txn_type: str = "deposit"  # deposit, withdrawal


class AnnouncementIn(BaseModel):
    # A board post is read by everyone, 50 at a time — an unbounded one is
    # paid for by every member on every load, not just its author.
    text: str = Field(max_length=4000)
    priority: str = "normal"  # normal, urgent


class ChoreIn(BaseModel):
    title: str
    frequency: str = "daily"  # daily, weekly
    assigned_members: list = []  # [member_id, ...]
    rotate: bool = True
    # Stars the assignee earns for finishing it. Per-chore, so taking the bins
    # out can be worth more than feeding the cat.
    star_reward: int = 3



# -----------------------------------------------------------------------------
# Health
# -----------------------------------------------------------------------------
@app.get("/")
async def root():
    # Deliberately minimal: this endpoint is public, so it must not disclose
    # which integrations/keys are configured (that inventory is a free recon
    # gift to an attacker). Config detail moved to the admin-only /api/health.
    return {"status": "online", "message": "Ahenora Backend is live"}


_AI_PROBE = {"last": None}


@app.get("/api/health/ai")
async def health_ai(probe: int = 0, user=Depends(require_user)):
    """Whether the AI features can work, verified from production itself.

    Every AI feature degrades gracefully, which is right for users and terrible
    for diagnosis: a retired model name failed every call for weeks while each
    feature quietly showed its fallback. This endpoint makes the plumbing
    observable. With ?probe=1 it performs one tiny real generation, so
    "working" means answered, not configured.

    Admin-gated: it names configured models, key state, and per-model errors —
    an integration map nobody but us needs, so nobody but us gets it.
    """
    if not is_admin_user(user):
        raise HTTPException(status_code=403, detail="Admin only")
    status = {
        "key_configured": bool(GOOGLE_API_KEY),
        "library_loaded": bool(genai),
        # The SDK importing and the client actually building are separate
        # failures with separate fixes, so report them separately.
        "client_ready": bool(_gemini_client()),
        "sdk": "google-genai",
        "model_env": GEMINI_MODEL or None,
        "model_resolved": _gemini_state["model"],
        "model_candidates": model_candidates(GEMINI_MODEL, _gemini_state["model"] or ""),
        "last_error": summarize_ai_error(_gemini_state["last_error"]),
        "model_errors": dict(_gemini_state["errors"]),
        # Everything below funnels through the same client, so one probe
        # vouches for the plumbing of all of them.
        "features": [
            "vision_extract", "voice_transcribe", "suggest_assignee",
            "morning_brief", "meal_recipe", "meal_suggestions",
        ],
    }

    # Always show what the key can actually use — the fastest way to diagnose
    # model_not_found is to see the real list.
    status["available_models"] = _discover_models()

    if probe:
        # Force a fresh discovery on an explicit probe.
        _gemini_state["discovered"] = None
        if not _gemini_client():
            status["probe"] = {"ok": False, "error": "GOOGLE_API_KEY missing or client unavailable"}
            return status
        # One probe a minute, globally. Enough to diagnose, too slow to farm.
        now = utcnow()
        last = _AI_PROBE["last"]
        if last and (now - last).total_seconds() < 60:
            status["probe"] = {"skipped": "rate limited, try again in a minute"}
            return status
        _AI_PROBE["last"] = now
        try:
            reply = await _gemini_generate("Reply with exactly one word: OK")
            status["probe"] = {
                "ok": "ok" in reply.lower(),
                "reply": reply[:40],
                "model": _gemini_state["model"],
            }
        except Exception as exc:  # noqa: BLE001 — categorised below, full text logged
            # Raw exception text can carry stack frames and internal detail;
            # an unauthenticated endpoint only gets the category. The full
            # text is in the logs for whoever operates the server.
            log.warning("ai health probe failed: %s", exc)
            status["probe"] = {"ok": False, "error": summarize_ai_error(str(exc))}
        # The probe may have just resolved (or changed) the model — report the
        # state after the probe, not the snapshot from before it.
        status["model_resolved"] = _gemini_state["model"]
        status["last_error"] = summarize_ai_error(_gemini_state["last_error"])
        status["model_errors"] = dict(_gemini_state["errors"])

    return status


@app.api_route("/api/health", methods=["GET", "HEAD"])
async def health():
    """Liveness + real database check for uptime monitoring.

    Actually pings MongoDB rather than just reporting that a URL is set, so a
    dead/stale connection surfaces here (503) instead of as user-facing 500s.

    HEAD is accepted as well as GET: uptime monitors (UptimeRobot and friends)
    default to HEAD, and a GET-only route answers 405, which they report as an
    outage even while the service is healthy.
    """
    if db is None:
        return JSONResponse(status_code=503, content={"status": "error", "database": "unconfigured"})
    try:
        start = utcnow()
        await asyncio.wait_for(db.command("ping"), timeout=5)
        elapsed_ms = int((utcnow() - start).total_seconds() * 1000)
        # invite_flow is a deploy marker: proves which invite generation this
        # running process carries when a user-side failure needs diagnosing.
        return {"status": "ok", "database": "ok", "db_latency_ms": elapsed_ms,
                "invite_flow": "v4"}
    except (asyncio.TimeoutError, Exception) as exc:  # noqa: B014 - report any failure
        log.warning("Health check database ping failed: %s", exc)
        return JSONResponse(status_code=503, content={"status": "error", "database": "unreachable"})


class ClientErrorIn(BaseModel):
    endpoint: str
    method: Optional[str] = None
    status: Optional[int] = None
    message: Optional[str] = None
    platform: Optional[str] = None


@app.get("/api/telemetry/invite-routes")
async def invite_route_stats(user=Depends(require_user)):
    """Which of the five acceptance doors real devices actually get through.

    The point of counting: the routes exist because one blocked iPhone
    refused specific request shapes, and they cannot be deleted on a hunch —
    old app builds still call them. These counts are what turns "probably
    nobody uses this one" into something a later session can act on.
    """
    if not is_admin_email(user.get("email", "")):
        raise HTTPException(status_code=403, detail="Admins only")
    database = get_db()
    rows = []
    async for row in database["invite_route_stats"].find({}, {"_id": 0}):
        rows.append({"route": row.get("route") or "unknown",
                     "count": int(row.get("count") or 0)})
    rows.sort(key=lambda r: -r["count"])
    return {"routes": rows, "total": sum(r["count"] for r in rows)}


@app.post("/api/telemetry/client-error")
async def report_client_error(payload: ClientErrorIn, user=Depends(require_user)):
    """Failed requests phone home, so silent breakage stops being silent.

    The invite-accept bug lived on a family iPhone for days and was reported
    by screenshot; this records the same facts automatically. Clients only
    send network-level failures and 5xx — semantic 4xx already surface in
    their own UI.
    """
    database = get_db()
    await database["client_errors"].insert_one({
        "error_id": new_id("cerr"),
        "user_id": user["user_id"],
        "family_id": user.get("family_id"),
        "name": user.get("name"),
        "endpoint": str(payload.endpoint)[:120],
        "method": (payload.method or "")[:8],
        "status": payload.status,
        "message": (payload.message or "")[:300],
        "platform": (payload.platform or "")[:20],
        "created_at": utcnow(),
    })
    # Bounded retention: two weeks is plenty for diagnosis.
    await database["client_errors"].delete_many(
        {"created_at": {"$lt": utcnow() - timedelta(days=14)}}
    )
    return {"ok": True}


@app.get("/api/telemetry/client-errors")
async def list_client_errors(user=Depends(require_user)):
    if not is_admin_user(user):
        raise HTTPException(status_code=403, detail="Admin only")
    database = get_db()
    rows = []
    cursor = database["client_errors"].find({}, {"_id": 0}).sort("created_at", -1).limit(50)
    async for item in cursor:
        item["created_at"] = iso(item.get("created_at"))
        rows.append(item)
    return rows


# Synthetic accounts used by the post-deploy production smoke test. Only
# emails under this reserved pattern can self-destruct; a real user cannot
# match it by accident, and matching it on purpose only deletes yourself.
_SMOKE_EMAIL_RE = re.compile(r"^smoke-[a-z0-9\-]+@household-coo\.smoke$")


class SmokeCleanupIn(BaseModel):
    # Families this account created and abandoned (its pre-join solo family);
    # deleted only if actually empty by the time everything above is gone.
    family_ids: Optional[list] = None



async def _resend_send(payload: dict) -> dict:
    """POST an email to Resend. Shared by the invite and the deletion receipt.
    The explicit User-Agent matters: Resend's Cloudflare edge 403s the default
    urllib agent before the request ever reaches the API."""
    def _send():
        req = urllib.request.Request(
            "https://api.resend.com/emails",
            data=json.dumps(payload).encode("utf-8"),
            headers={
                "Authorization": f"Bearer {RESEND_API_KEY}",
                "Content-Type": "application/json",
                "User-Agent": f"HouseholdCOO/1.0 (+{APP_NAME})",
            },
            method="POST",
        )
        with urllib.request.urlopen(req, timeout=15) as response:
            raw = response.read().decode("utf-8")
            return {"sent": True, "provider": "resend",
                    "response": (json.loads(raw) if raw else {})}
    return await asyncio.to_thread(_send)


async def send_account_deleted_email(to_email: str, name: str) -> dict:
    """A receipt for a deletion — sent best-effort, never blocking the delete."""
    if not RESEND_API_KEY or not INVITE_FROM_EMAIL or not to_email:
        return {"sent": False}
    safe_app = html.escape(APP_NAME)
    safe_name = html.escape(name or "there")
    subject = f"Your {APP_NAME} account has been deleted"
    text = (
        f"Hi {name or 'there'},\n\n"
        f"Your {APP_NAME} account and its data have been permanently deleted, as you requested.\n\n"
        "If this was not you, contact us immediately — but note the data cannot be recovered.\n"
    )
    html_body = f"""
<div style="font-family:-apple-system,'Segoe UI',Roboto,Arial,sans-serif;background:#f4f5f2;padding:24px;">
  <div style="max-width:520px;margin:0 auto;background:#fff;border:1px solid #e6e1da;border-radius:16px;padding:28px;">
    <p style="color:#202323;font-size:16px;line-height:1.55;margin:0 0 14px;">Hi {safe_name},</p>
    <p style="color:#202323;font-size:16px;line-height:1.55;margin:0 0 16px;">
      Your <strong>{safe_app}</strong> account and all of its data have been permanently deleted,
      as you requested. There is nothing left to recover.
    </p>
    <p style="color:#747b7c;font-size:13px;line-height:1.5;margin:18px 0 0;">
      If you did not ask for this, reply to this email straight away.
    </p>
  </div>
</div>""".strip()
    try:
        return await _resend_send({"from": INVITE_FROM_EMAIL, "to": [to_email],
                                   "subject": subject, "text": text, "html": html_body})
    except Exception as exc:  # noqa: BLE001 — a receipt must never fail a delete
        log.warning("account-deleted email skipped: %s", exc)
        return {"sent": False}


async def send_password_reset_email(to_email: str, name: str, code: str) -> dict:
    """The one-time code that lets someone who has forgotten their password
    prove they own the inbox. Sent best-effort; a delivery failure must not tell
    the caller whether the account exists, so the endpoint ignores the result."""
    if not RESEND_API_KEY or not INVITE_FROM_EMAIL or not to_email:
        return {"sent": False}
    safe_app = html.escape(APP_NAME)
    safe_name = html.escape(name or "there")
    safe_code = html.escape(code)
    minutes = PASSWORD_RESET_TTL_MINUTES
    subject = f"Your {APP_NAME} password reset code"
    text = (
        f"Hi {name or 'there'},\n\n"
        f"Your {APP_NAME} password reset code is: {code}\n\n"
        f"Enter it in the app to set a new password. It expires in {minutes} minutes.\n\n"
        "If you did not ask to reset your password, you can ignore this email — "
        "your password has not changed.\n"
    )
    html_body = f"""
<div style="font-family:-apple-system,'Segoe UI',Roboto,Arial,sans-serif;background:#f4f5f2;padding:24px;">
  <div style="max-width:520px;margin:0 auto;background:#fff;border:1px solid #e6e1da;border-radius:16px;padding:28px;">
    <p style="color:#202323;font-size:16px;line-height:1.55;margin:0 0 14px;">Hi {safe_name},</p>
    <p style="color:#202323;font-size:16px;line-height:1.55;margin:0 0 12px;">
      Here is your <strong>{safe_app}</strong> password reset code:
    </p>
    <p style="font-size:34px;font-weight:700;letter-spacing:8px;color:#202323;text-align:center;
              background:#f4f5f2;border-radius:12px;padding:16px 0;margin:0 0 16px;">{safe_code}</p>
    <p style="color:#202323;font-size:15px;line-height:1.55;margin:0 0 16px;">
      Enter it in the app to set a new password. It expires in {minutes} minutes.
    </p>
    <p style="color:#747b7c;font-size:13px;line-height:1.5;margin:18px 0 0;">
      If you did not ask to reset your password, you can ignore this email — your password has not changed.
    </p>
  </div>
</div>""".strip()
    try:
        return await _resend_send({"from": INVITE_FROM_EMAIL, "to": [to_email],
                                   "subject": subject, "text": text, "html": html_body})
    except Exception as exc:  # noqa: BLE001 — never reveal delivery outcome to caller
        log.warning("password-reset email skipped: %s", exc)
        return {"sent": False}


# Every collection whose rows belong to a household. Purged by family_id when
# the last account-holder leaves. Listed exhaustively on purpose: a forgotten
# collection is orphaned personal data, which is the one thing account deletion
# exists to prevent. Deleting by family_id from a collection that lacks the
# field is a harmless no-op, so erring long is safe; erring short is not.
_FAMILY_SCOPED_COLLECTIONS = (
    "activity", "allowance_txns", "allowances", "announcements", "calendar_contacts",
    "cards", "carpools", "chore_logs", "chores", "expenses", "family_invites",
    "family_members", "handoff_notes", "meal_plans_saved", "meals", "redemptions",
    "rewards", "routine_logs", "routines", "shopping_history", "shopping_list",
    "star_transactions", "templates", "vault",
)


async def _purge_family(database: Any, family_id: str) -> None:
    for collection in _FAMILY_SCOPED_COLLECTIONS:
        await database[collection].delete_many({"family_id": family_id})
    await database["families"].delete_many({"family_id": family_id})


async def _purge_user_account(database: Any, user_id: str, email: str) -> None:
    for collection in ("user_sessions", "notification_tokens", "notification_settings",
                       "support_tickets"):
        await database[collection].delete_many({"user_id": user_id})
    if email:
        await database["family_invites"].delete_many({"email": email, "status": "pending"})
    await database["users"].delete_many({"user_id": user_id})


@app.post("/api/auth/smoke-cleanup")
async def smoke_cleanup(payload: Optional[SmokeCleanupIn] = Body(None), user=Depends(require_user)):
    """Self-deletion for smoke-test accounts, so repeated production runs
    leave no residue (members, invites, sessions, tokens, user, empty family)."""
    database = get_db()
    email = (user.get("email") or "").strip().lower()
    if not _SMOKE_EMAIL_RE.match(email):
        raise HTTPException(status_code=403, detail="Not a smoke-test account")
    uid = user["user_id"]
    await database["family_members"].delete_many({"user_id": uid})
    await database["family_invites"].delete_many({"email": email})
    await database["user_sessions"].delete_many({"user_id": uid})
    await database["notification_tokens"].delete_many({"user_id": uid})
    await database["users"].delete_many({"user_id": uid})
    for fid in ((payload.family_ids if payload else None) or [])[:5]:
        remaining = await database["family_members"].find_one({"family_id": fid}, {"_id": 0})
        if remaining is None:
            await database["families"].delete_many({"family_id": fid})
    return {"ok": True}


class DeleteAccountIn(BaseModel):
    password: Optional[str] = None
    confirm: Optional[bool] = None


@app.post("/api/auth/delete-account")
async def delete_account(payload: DeleteAccountIn, user=Depends(require_user)):
    """Delete your own account and its data, for real and at once.

    Self-service and immediate — the store requires it and it is simply how it
    should work. Two shapes, decided by whether anyone else has a sign-in:

      - You are the only account in the household → the whole household goes:
        every card, chore, reward, vault document, the lot. Child profiles are
        not accounts; they go with the household they lived in.
      - Someone else has an account too (a co-parent) → only you leave. The
        household and its shared data continue under them; the founder role
        recomputes to the next-earliest parent on its own.
    """
    database = get_db()
    fresh = await database["users"].find_one({"user_id": user["user_id"]}, {"_id": 0}) or user

    # A password account must prove it is really them. An OAuth account has no
    # password to check; the app gates it behind a typed confirmation instead.
    if fresh.get("password_hash"):
        if not payload.password or not verify_password(payload.password, fresh["password_hash"]):
            raise HTTPException(status_code=403, detail="That password is not correct.")
    elif not payload.confirm:
        raise HTTPException(status_code=400, detail="Confirm the deletion to continue.")

    family_id = fresh.get("family_id")
    email = (fresh.get("email") or "").strip().lower()

    # "Is anyone else left in this household?" is asked of the users collection,
    # not family_members: every account-holder has a users doc carrying a
    # user_id and this family_id, whereas a founder's MEMBER row can predate
    # user_id linkage. Asking family_members for another user_id would miss such
    # a founder and wrongly purge the whole household when a co-parent leaves.
    other_account = None
    if family_id:
        other_account = await database["users"].find_one(
            {"family_id": family_id, "user_id": {"$ne": fresh["user_id"]}},
            {"_id": 0})

    deleted_household = False
    if family_id and other_account is None:
        # Last account-holder: the household leaves with them.
        await _purge_family(database, family_id)
        deleted_household = True
    elif family_id:
        # A co-parent remains — take only this member out; the family stays.
        await database["family_members"].delete_many(
            {"family_id": family_id, "user_id": fresh["user_id"]})

    await _purge_user_account(database, fresh["user_id"], email)

    # A receipt, after the fact and never in the way.
    await send_account_deleted_email(email, fresh.get("name") or "")

    return {"ok": True, "deleted_household": deleted_household}


@app.get("/api/health/config")
async def health_config(user=Depends(require_user)):
    """Admin-only configuration inventory (was previously public on `/`)."""
    if not is_admin_user(user):
        raise HTTPException(status_code=403, detail="Admin only")
    return {
        "api_configured": bool(GOOGLE_API_KEY),
        "db_configured": bool(MONGO_URL),
        "backend_version": "pricing_gating_v1",
        "email_configured": bool(RESEND_API_KEY and INVITE_FROM_EMAIL),
        "admin_access_enabled": bool(ADMIN_EMAILS),
        "voice_configured": bool(GOOGLE_API_KEY and genai),
        "google_web_configured": bool(GOOGLE_WEB_CLIENT_ID),
        "google_android_configured": bool(GOOGLE_ANDROID_CLIENT_ID),
        "google_client_ids_count": len(GOOGLE_CLIENT_IDS),
        "billing_configured": bool(os.environ.get("RC_WEBHOOK_SECRET")),
    }

# -----------------------------------------------------------------------------
# Auth
# -----------------------------------------------------------------------------
@app.post("/api/auth/session")
async def exchange_session(payload: SessionIn):
    database = get_db()

    if not GOOGLE_CLIENT_IDS:
        raise HTTPException(status_code=500, detail="Google OAuth client IDs are missing")

    token_info = None
    last_error = None
    for client_id in GOOGLE_CLIENT_IDS:
        try:
            token_info = await asyncio.wait_for(
                asyncio.to_thread(
                    google_id_token.verify_oauth2_token,
                    payload.session_id,
                    GoogleRequest(),
                    client_id,
                ),
                timeout=10,
            )
            break
        except Exception as e:
            last_error = e

    if not token_info:
        if isinstance(last_error, asyncio.TimeoutError):
            raise HTTPException(status_code=504, detail="Timed out reaching Google to verify sign-in")
        # Log the specific reason server-side; return a generic message so the
        # accepted audience / client-ID details aren't reflected to the client.
        log.warning("Google token verification failed: %s", last_error)
        raise HTTPException(status_code=401, detail="Could not verify Google sign-in. Please try again.")

    google_sub = token_info["sub"]
    # Normalize like register/login so the same inbox is one account, never two
    # rows that differ only by casing.
    email = (token_info.get("email") or "").strip().lower()
    # Only a Google-*verified* email may be used to link into an existing
    # account: verify_oauth2_token proves the token is genuine, not that Google
    # confirmed the address. Without this, a token bearing an unverified email
    # equal to a victim's account (possible on Workspace/custom domains) could
    # graft its google_sub onto that account and hijack it. Consumer @gmail is
    # always verified, so this only refuses the risky federated case.
    email_verified = token_info.get("email_verified") is True
    name = token_info.get("name", email.split("@")[0] if email else "Parent")
    picture = token_info.get("picture")

    invite = None
    target_family_id = None

    if payload.invite_token:
        invite = await database["family_invites"].find_one(
            {"token": payload.invite_token},
            {"_id": 0},
        )
        if not invite:
            raise HTTPException(status_code=404, detail="Invite not found")

        if invite.get("expires_at") and _expired(invite["expires_at"]):
            await database["family_invites"].update_one(
                {"invite_id": invite["invite_id"]},
                {"$set": {"status": "expired", "updated_at": utcnow()}},
            )
            raise HTTPException(status_code=410, detail="Invite has expired")

        if invite.get("status") == "accepted" and invite.get("accepted_by_email") != email:
            raise HTTPException(status_code=409, detail="Invite has already been accepted")

        target_family_id = invite["family_id"]

    user = await database["users"].find_one({"google_sub": google_sub}, {"_id": 0})

    if not user and email and email_verified:
        # Never mint a second row for an inbox that already has an account (an
        # email/password sign-up, most importantly). Link Google to it instead —
        # otherwise the new passwordless row shadows theirs and email login then
        # answers "no password account for this email", locking them out.
        # Gated on email_verified so an unverified address can never link into
        # (and hijack) someone else's existing account.
        existing = await database["users"].find_one(
            {"email": {"$regex": f"^{re.escape(email)}$", "$options": "i"}}, {"_id": 0})
        if existing:
            await database["users"].update_one(
                {"user_id": existing["user_id"]},
                {"$set": {"google_sub": google_sub, "email": email,
                          "picture": existing.get("picture") or picture,
                          "updated_at": utcnow()}})
            user = {**existing, "google_sub": google_sub, "email": email,
                    "picture": existing.get("picture") or picture}

    if not user:
        family_id = target_family_id or new_id("family")
        user = {
            "user_id": new_id("user"),
            "google_sub": google_sub,
            "email": email,
            "name": name,
            "picture": picture,
            "family_id": family_id,
            "language": payload.language if payload.language in ("en", "es", "fr", "de") else "en",
            "onboarding_completed": False,
            "created_at": utcnow(),
            "updated_at": utcnow(),
        }
        await database["users"].insert_one(user)

        if not target_family_id:
            await database["families"].insert_one(
                {
                    "family_id": family_id,
                    "plan": "village",
                    "billing_cycle": "monthly",
                    "grandfathered": False,
                    "updated_at": utcnow(),
                    "ai_scans_used": 0,
                    "ai_scans_period_start": utcnow(),
                    "vault_bytes_used": 0,
                }
            )

            await database["family_members"].insert_one({
                "member_id": new_id("member"),
                "family_id": family_id,
                "user_id": user["user_id"],
                "email": email,
                "name": name,
                "role": invite_member_role(invite),
                "avatar": picture,
                "stars": 0,
                "pin_hash": None,
                "created_at": utcnow(),
            })
        else:
            await add_user_to_family_if_needed(database, user, target_family_id, invite_member_role(invite))
    else:
        updates = {
            "email": email,
            "name": name,
            "picture": picture,
            "updated_at": utcnow(),
        }
        old_family_id = user.get("family_id")
        if target_family_id:
            updates["family_id"] = target_family_id

        await database["users"].update_one(
            {"user_id": user["user_id"]},
            {"$set": updates},
        )
        user = await database["users"].find_one({"user_id": user["user_id"]}, {"_id": 0})

        if target_family_id:
            await add_user_to_family_if_needed(database, user, target_family_id, invite_member_role(invite))
            await _bring_children_along(database, old_family_id, target_family_id)

    if invite:
        await database["family_invites"].update_one(
            {"invite_id": invite["invite_id"]},
            {
                "$set": {
                    "status": "accepted",
                    "accepted_at": utcnow(),
                    "accepted_by_user_id": user["user_id"],
                    "accepted_by_email": email,
                    "updated_at": utcnow(),
                }
            },
        )
        await notify_invite_accepted(database, invite, user.get("name") or email)

    raw_session = secrets.token_urlsafe(32)
    await database["user_sessions"].insert_one(
        {
            "session_id": new_id("sess"),
            "user_id": user["user_id"],
            "token_hash": sha256(raw_session),
            "expires_at": utcnow() + timedelta(days=SESSION_DAYS),
            "created_at": utcnow(),
        }
    )

    return {"user": public_user(user), "session_token": raw_session}


async def _issue_session(database, user_id: str) -> str:
    raw_session = secrets.token_urlsafe(32)
    await database["user_sessions"].insert_one({
        "session_id": new_id("sess"),
        "user_id": user_id,
        "token_hash": sha256(raw_session),
        "expires_at": utcnow() + timedelta(days=SESSION_DAYS),
        "created_at": utcnow(),
    })
    return raw_session


async def _seed_new_family(database, user: dict, family_id: str, email: str, name: str, picture=None):
    await database["families"].insert_one({
        "family_id": family_id,
        "plan": "village",
        "billing_cycle": "monthly",
        "grandfathered": False,
        "updated_at": utcnow(),
        "ai_scans_used": 0,
        "ai_scans_period_start": utcnow(),
        "vault_bytes_used": 0,
    })
    await database["family_members"].insert_one({
        "member_id": new_id("member"),
        "family_id": family_id,
        "user_id": user["user_id"],
        "email": email,
        "name": name,
        "role": "Parent",
        "avatar": picture,
        "stars": 0,
        "pin_hash": None,
        "created_at": utcnow(),
    })


async def _resolve_invite(database, invite_token: Optional[str], email: str):
    """Validate an invite token and return (invite_doc, target_family_id)."""
    if not invite_token:
        return None, None
    invite = await database["family_invites"].find_one({"token": invite_token}, {"_id": 0})
    if not invite:
        raise HTTPException(status_code=404, detail="Invite not found")
    if invite.get("expires_at") and _expired(invite["expires_at"]):
        await database["family_invites"].update_one(
            {"invite_id": invite["invite_id"]},
            {"$set": {"status": "expired", "updated_at": utcnow()}},
        )
        raise HTTPException(status_code=410, detail="Invite has expired")
    if invite.get("status") == "accepted" and invite.get("accepted_by_email") != email:
        raise HTTPException(status_code=409, detail="Invite has already been accepted")
    return invite, invite["family_id"]


async def _bring_children_along(database, old_family_id: Optional[str], new_family_id: str):
    """A parent joining a household brings their kids' profiles with them.

    Field case: the invitee had set up all three children in her own family
    before joining — accepting must not strand them (or their stars) in an
    abandoned household. Children whose name already exists in the target
    family are skipped, so co-parents who each created "Richard" don't end
    up with two of him.
    """
    if not old_family_id or old_family_id == new_family_id:
        return
    existing = set()
    async for member in database["family_members"].find({"family_id": new_family_id}, {"_id": 0}):
        existing.add((member.get("name") or "").strip().lower())
    async for member in database["family_members"].find({"family_id": old_family_id}, {"_id": 0}):
        if (member.get("role") or "").strip().lower() != "child":
            continue
        name = (member.get("name") or "").strip().lower()
        if not name or name in existing:
            continue
        existing.add(name)
        await database["family_members"].update_one(
            {"member_id": member["member_id"]},
            {"$set": {"family_id": new_family_id}},
        )
        # Their history follows, so balances and past rewards stay visible.
        for collection in ("star_transactions", "redemptions", "allowances", "allowance_txns"):
            await database[collection].update_many(
                {"member_id": member["member_id"]},
                {"$set": {"family_id": new_family_id}},
            )


async def _accept_invite_for_user(database, user: dict, invite: Optional[dict], target_family_id: Optional[str]):
    """Move `user` into the inviting family and mark the invite accepted.

    Shared by email login and the signed-in accept endpoint so the two paths
    can never drift. Returns (fresh_user, joined).
    """
    email = user.get("email", "")
    joined = False
    if target_family_id and target_family_id != user.get("family_id"):
        old_family_id = user.get("family_id")
        await database["users"].update_one(
            {"user_id": user["user_id"]},
            {"$set": {"family_id": target_family_id, "updated_at": utcnow()}},
        )
        user = await database["users"].find_one({"user_id": user["user_id"]}, {"_id": 0})
        await add_user_to_family_if_needed(database, user, target_family_id, invite_member_role(invite))
        await _bring_children_along(database, old_family_id, target_family_id)
        await log_activity(database, {"family_id": target_family_id,
                                      "user_id": user["user_id"],
                                      "name": user.get("name")},
                           "member_joined", user.get("name") or "")
        joined = True
    if invite and invite.get("status") != "accepted":
        await database["family_invites"].update_one(
            {"invite_id": invite["invite_id"]},
            {"$set": {
                "status": "accepted",
                "accepted_at": utcnow(),
                "accepted_by_user_id": user["user_id"],
                "accepted_by_email": email,
                "updated_at": utcnow(),
            }},
        )
        # Bounded: the join must never hang on push delivery. Five seconds is
        # plenty for exp.host; past that the inviter just misses one push.
        try:
            await asyncio.wait_for(
                notify_invite_accepted(database, invite, user.get("name") or email),
                timeout=5.0,
            )
        except Exception as exc:
            log.warning("invite acceptance notification skipped: %s", exc)
    return user, joined


def _expired(value) -> bool:
    """Has this stored timestamp already passed?

    Always via _coerce_dt, never a bare `stored < utcnow()`. BSON carries no
    timezone, so a datetime read back from Mongo can be naive while utcnow()
    is aware, and comparing the two raises TypeError rather than returning
    False. That is not hypothetical: it is what took invite acceptance down in
    production, as an unhandled 500 on the one endpoint a new co-parent needs.

    The client now sets tz_aware=True, which fixes the same thing one layer
    lower — but correctness here should not depend on a connection flag set
    600 lines away, and a stored value can also be an ISO string.
    """
    parsed = _coerce_dt(value)
    return bool(parsed and parsed < utcnow())


@app.post("/api/auth/register")
async def register_email(payload: EmailRegisterIn):
    database = get_db()

    email = payload.email.strip().lower()
    name = payload.name.strip() or (email.split("@")[0] if email else "Parent")
    password = payload.password or ""

    if not email or "@" not in email:
        raise HTTPException(status_code=400, detail="A valid email is required")
    if len(password) < 8:
        raise HTTPException(status_code=400, detail="Password must be at least 8 characters")

    existing = await database["users"].find_one({"email": email}, {"_id": 0})
    if existing:
        raise HTTPException(
            status_code=409,
            detail="An account with this email already exists. Try signing in instead.",
        )

    invite, target_family_id = await _resolve_invite(database, payload.invite_token, email)

    family_id = target_family_id or new_id("family")
    user = {
        "user_id": new_id("user"),
        "email": email,
        "name": name,
        "picture": None,
        "password_hash": hash_password(password),
        "family_id": family_id,
        "language": payload.language if payload.language in ("en", "es", "fr", "de") else "en",
        "onboarding_completed": False,
        "created_at": utcnow(),
        "updated_at": utcnow(),
    }
    await database["users"].insert_one(user)

    if not target_family_id:
        await _seed_new_family(database, user, family_id, email, name)
    else:
        await add_user_to_family_if_needed(database, user, target_family_id, invite_member_role(invite))

    if invite:
        await database["family_invites"].update_one(
            {"invite_id": invite["invite_id"]},
            {"$set": {
                "status": "accepted",
                "accepted_at": utcnow(),
                "accepted_by_user_id": user["user_id"],
                "accepted_by_email": email,
                "updated_at": utcnow(),
            }},
        )
        await notify_invite_accepted(database, invite, name or email)

    raw_session = await _issue_session(database, user["user_id"])
    return {"user": public_user(user), "session_token": raw_session}


@app.post("/api/auth/login")
async def login_email(payload: EmailLoginIn):
    database = get_db()

    email = payload.email.strip().lower()
    identity = f"login:{email}"
    if _auth_locked(identity):
        raise HTTPException(status_code=429, detail="Too many attempts. Please try again later.")

    # A stray duplicate can shadow the real password account under one email: a
    # Google row created before account-linking (no password), or a legacy row
    # stored with different casing. Gather every match case-insensitively and
    # pick the one that actually carries a password, so those users can get in.
    matches = []
    async for candidate in database["users"].find(
        {"email": {"$regex": f"^{re.escape(email)}$", "$options": "i"}}, {"_id": 0}
    ):
        matches.append(candidate)
    pw_matches = [u for u in matches if u.get("password_hash")]
    if not pw_matches:
        # Distinct hint is intentional UX (many users sign up with Google);
        # email existence is low-sensitivity here. Still counts toward lockout.
        _auth_record_fail(identity)
        raise HTTPException(
            status_code=401,
            detail="No password account found for this email. Try Google sign-in.",
        )
    # Verify against every password row for this email, not just the first: if a
    # legacy duplicate exists, the one whose password actually matches wins, so
    # the account is never a spurious 401 just for being second in the list.
    user = next((u for u in pw_matches if verify_password(payload.password or "", u["password_hash"])), None)
    if not user:
        _auth_record_fail(identity)
        raise HTTPException(status_code=401, detail="Incorrect email or password")

    _auth_clear(identity)

    # Login never consumed invite tokens, so an existing email account
    # clicking an invite link could never join the inviting family.
    # Same semantics as Google sign-in and registration now.
    invite, target_family_id = await _resolve_invite(database, payload.invite_token, email)
    user, _ = await _accept_invite_for_user(database, user, invite, target_family_id)

    raw_session = await _issue_session(database, user["user_id"])
    return {"user": public_user(user), "session_token": raw_session}


class ChangePasswordIn(BaseModel):
    current_password: Optional[str] = None
    new_password: Optional[str] = None


@app.post("/api/auth/sign-out-everywhere")
async def sign_out_everywhere(user=Depends(require_user), authorization: str = Header(default="")):
    """End every other session on every other device, keeping this one.

    Changing your password already does this, but that is the wrong tool when
    nothing is wrong with the password — a phone left at a friend's house, a
    tablet handed on. With sessions renewing themselves for as long as they are
    used, there has to be a way to say "not that device, not any more".
    """
    database = get_db()
    raw_auth = authorization if isinstance(authorization, str) else ""
    this_token = raw_auth.replace("Bearer ", "", 1).strip()
    result = await database["user_sessions"].delete_many(
        {"user_id": user["user_id"], "token_hash": {"$ne": sha256(this_token)}})
    return {"ok": True, "ended": getattr(result, "deleted_count", 0)}


@app.post("/api/auth/change-password")
async def change_password(payload: ChangePasswordIn, user=Depends(require_user),
                          authorization: str = Header(default="")):
    """Change the password on a password account.

    Only for accounts that actually have a password — a Google account has
    nothing to change and is told so rather than left guessing. The current
    password must be proven before a new one is set, and the new one meets the
    same floor as registration.
    """
    database = get_db()
    stored = user.get("password_hash")
    if not stored:
        raise HTTPException(
            status_code=400,
            detail="This account signs in with Google, so it has no password to change.")
    if not verify_password(payload.current_password or "", stored):
        raise HTTPException(status_code=403, detail="Current password is incorrect")
    new_password = payload.new_password or ""
    if len(new_password) < 8:
        raise HTTPException(status_code=400, detail="Password must be at least 8 characters")
    if verify_password(new_password, stored):
        raise HTTPException(status_code=400, detail="New password must be different from the current one")
    await database["users"].update_one(
        {"user_id": user["user_id"]},
        {"$set": {"password_hash": hash_password(new_password), "updated_at": utcnow()}})

    # Changing your password is how a person locks out a device they no longer
    # control — a sold laptop, an ex-partner's tablet. Only the forgotten-password
    # reset used to revoke, so every other session survived; and now that a live
    # session renews itself on use, one left signed in would never lapse at all.
    # Every session but the one making this request is ended here.
    # Called directly (tests) the header default is a Header sentinel, not a
    # string. An unknown caller token hashes to something no row carries, so the
    # revoke sweeps every session — erring toward locking out rather than
    # leaving a device signed in, which is the safe direction here.
    raw_auth = authorization if isinstance(authorization, str) else ""
    this_token = raw_auth.replace("Bearer ", "", 1).strip()
    await database["user_sessions"].delete_many(
        {"user_id": user["user_id"], "token_hash": {"$ne": sha256(this_token)}})
    return {"ok": True}


class RequestPasswordResetIn(BaseModel):
    email: Optional[str] = None


@app.post("/api/auth/request-password-reset")
async def request_password_reset(payload: RequestPasswordResetIn):
    """Start a forgotten-password reset by emailing a one-time code.

    Always answers the same {"ok": True}, whether or not an account exists, so
    the endpoint never becomes an oracle for which emails are registered. The
    work — minting a code, storing its hash, sending the mail — happens only for
    a real password account, silently.
    """
    database = get_db()
    email = (payload.email or "").strip().lower()
    identity = f"reset-request:{email}"
    # Rate-limit per email so nobody can spray reset mail at an inbox. Locked
    # requests still return ok — the caller learns nothing either way.
    if not email or "@" not in email or _auth_locked(identity):
        return {"ok": True}
    _auth_record_fail(identity)

    user = await database["users"].find_one({"email": email}, {"_id": 0})
    # A Google account has no password to reset; treat it exactly like a missing
    # account so the response gives nothing away.
    if user and user.get("password_hash"):
        code = f"{secrets.randbelow(1_000_000):06d}"
        await database["password_resets"].update_one(
            {"user_id": user["user_id"]},
            {"$set": {
                "user_id": user["user_id"],
                "email": email,
                "code_hash": sha256(code),
                "expires_at": utcnow() + timedelta(minutes=PASSWORD_RESET_TTL_MINUTES),
                "attempts": 0,
                "created_at": utcnow(),
            }},
            upsert=True,
        )
        await send_password_reset_email(email, user.get("name") or "there", code)
    return {"ok": True}


class ResetPasswordIn(BaseModel):
    email: Optional[str] = None
    code: Optional[str] = None
    new_password: Optional[str] = None


@app.post("/api/auth/reset-password")
async def reset_password(payload: ResetPasswordIn):
    """Finish a reset: prove the emailed code, set a new password, sign in.

    A correct code is a full account takeover, so on success every existing
    session is dropped — anyone who was signed in on the old password is signed
    out — and a fresh session is issued to whoever completed the reset.
    """
    database = get_db()
    email = (payload.email or "").strip().lower()
    code = (payload.code or "").strip()
    new_password = payload.new_password or ""
    identity = f"reset-verify:{email}"
    if _auth_locked(identity):
        raise HTTPException(status_code=429, detail="Too many attempts. Please try again later.")
    if len(new_password) < 8:
        raise HTTPException(status_code=400, detail="Password must be at least 8 characters")

    record = await database["password_resets"].find_one({"email": email}, {"_id": 0})
    expires = _coerce_dt(record.get("expires_at")) if record else None
    bad_code = "That code is incorrect or has expired. Request a new one."
    if not record or not code or expires is None or expires < utcnow():
        _auth_record_fail(identity)
        raise HTTPException(status_code=400, detail=bad_code)
    if int(record.get("attempts", 0)) >= PASSWORD_RESET_MAX_ATTEMPTS:
        raise HTTPException(status_code=400, detail=bad_code)
    if not secrets.compare_digest(record.get("code_hash", ""), sha256(code)):
        await database["password_resets"].update_one(
            {"user_id": record["user_id"]}, {"$inc": {"attempts": 1}})
        _auth_record_fail(identity)
        raise HTTPException(status_code=400, detail=bad_code)

    user = await database["users"].find_one({"user_id": record["user_id"]}, {"_id": 0})
    if not user:
        raise HTTPException(status_code=400, detail=bad_code)

    await database["users"].update_one(
        {"user_id": user["user_id"]},
        {"$set": {"password_hash": hash_password(new_password), "updated_at": utcnow()}})
    # The code is spent, and every session opened before the reset is now stale.
    await database["password_resets"].delete_many({"user_id": user["user_id"]})
    await database["user_sessions"].delete_many({"user_id": user["user_id"]})
    _auth_clear(identity)
    _auth_clear(f"login:{email}")

    fresh = await database["users"].find_one({"user_id": user["user_id"]}, {"_id": 0})
    raw_session = await _issue_session(database, user["user_id"])
    return {"user": public_user(fresh or user), "session_token": raw_session}


@app.get("/api/auth/me")
async def me(user=Depends(require_user), database=Depends(get_db)):
    # Cheap retention heartbeat: the app calls /auth/me on open, so stamping
    # last_active here is enough to compute daily-active and D1/D7 return
    # without a per-request events pipeline. Best-effort — never block sign-in.
    try:
        await database["users"].update_one(
            {"user_id": user["user_id"]},
            {"$set": {"last_active_at": utcnow()}},
        )
    except Exception:
        pass
    return public_user(user)


@app.post("/api/auth/logout")
async def logout(user=Depends(require_user), authorization: str = Header(default="")):
    database = get_db()
    token = authorization.replace("Bearer ", "", 1).strip()
    await database["user_sessions"].delete_many(
        {"user_id": user["user_id"], "token_hash": sha256(token)}
    )
    return {"ok": True}


@app.post("/api/auth/complete-onboarding")
async def complete_onboarding(user=Depends(require_user)):
    database = get_db()
    await database["users"].update_one(
        {"user_id": user["user_id"]},
        {"$set": {"onboarding_completed": True, "updated_at": utcnow()}},
    )
    user = await database["users"].find_one({"user_id": user["user_id"]}, {"_id": 0})
    return public_user(user)


@app.patch("/api/auth/language")
async def set_language(payload: LanguageIn, user=Depends(require_user)):
    database = get_db()
    await database["users"].update_one(
        {"user_id": user["user_id"]},
        {"$set": {"language": payload.language, "updated_at": utcnow()}},
    )
    user = await database["users"].find_one({"user_id": user["user_id"]}, {"_id": 0})
    return public_user(user)


# -----------------------------------------------------------------------------
# Family
# -----------------------------------------------------------------------------
@app.get("/api/family/members")
async def family_members(user=Depends(require_user)):
    database = get_db()
    raw = [item async for item in
           database["family_members"].find({"family_id": user["family_id"]}, {"_id": 0})]
    # The founder is the earliest parent-level member — the only one a co-parent
    # cannot remove. Computed once here so the client does not have to guess.
    parents = sorted(
        (m for m in raw if _is_parent_role(m.get("role"))),
        key=lambda m: (_coerce_dt(m.get("created_at")) or _OLDEST_DT))
    founder_id = parents[0]["member_id"] if parents else None
    # Which row is "me", resolved the same resilient way removal is — by user_id,
    # else by email — so a founder whose row predates user_id linkage is still
    # recognised as themselves (and so never offered a Remove button on it).
    me = await _member_for_user(database, user["family_id"], user)
    my_member_id = me.get("member_id")
    my_email = str(user.get("email") or "").strip().lower()
    rows = []
    for item in raw:
        # Roll a child's weekly meter over on read, so the Kids screen always
        # shows this week's earnings rather than a stale count from last week.
        if str(item.get("role", "")).lower() == "child":
            item = await roll_week_if_stale(database, item)
        row = public_member(item)
        row["is_me"] = (
            (item.get("user_id") is not None and item.get("user_id") == user["user_id"])
            or (my_member_id is not None and item.get("member_id") == my_member_id)
            or (bool(my_email) and str(item.get("email") or "").strip().lower() == my_email)
        )
        row["is_founder"] = item.get("member_id") == founder_id
        rows.append(row)
    return rows


@app.put("/api/family/members/{member_id}/weekend-goal")
async def set_weekend_goal(member_id: str, payload: dict = Body(...), user=Depends(require_user)):
    """Pin (or clear) the weekend treat a child is working toward this week.

    Only drives the progress ring; it commits nothing and costs nothing. A
    reward_id of null clears it. The reward must be a weekend treat in this
    family, so the ring can never point at a saved-up reward the weekly meter
    does not gate.
    """
    database = get_db()
    member = await database["family_members"].find_one(
        {"member_id": member_id, "family_id": user["family_id"]}, {"_id": 0})
    if not member:
        raise HTTPException(status_code=404, detail="Family member not found")

    reward_id = payload.get("reward_id")
    # A raw body means this could be {"$ne": null} rather than an id, which
    # would match an arbitrary reward and then be stored as the child's goal.
    if reward_id is not None and not isinstance(reward_id, str):
        raise HTTPException(status_code=400, detail="Invalid reward")
    if reward_id:
        reward = await database["rewards"].find_one(
            {"reward_id": reward_id, "family_id": user["family_id"]}, {"_id": 0})
        if not reward or not reward.get("weekend"):
            raise HTTPException(status_code=400, detail="That isn't a weekend treat.")

    await database["family_members"].update_one(
        {"member_id": member_id, "family_id": user["family_id"]},
        {"$set": {"weekend_goal_reward_id": reward_id or None}},
    )
    updated = await database["family_members"].find_one({"member_id": member_id}, {"_id": 0})
    return public_member(updated)


@app.post("/api/family/members")
async def create_family_member(payload: ChildIn, user=Depends(require_full_member)):
    database = get_db()
    name = payload.name.strip()
    if not name:
        raise HTTPException(status_code=400, detail="Child name is required")
    if len(name) > 60:
        raise HTTPException(status_code=400, detail="Name is too long")
    # Same reason as the rename: cards address a child by name, so two children
    # sharing one makes "who earned this star" unanswerable.
    clash = await database["family_members"].find_one({
        "family_id": user["family_id"],
        "name": {"$regex": f"^{re.escape(name)}$", "$options": "i"},
    })
    if clash:
        raise HTTPException(
            status_code=400,
            detail="Somebody in this household already has that name.",
        )

    starting_stars = max(0, int(payload.starting_stars or 0))
    pin = (payload.pin or "").strip()
    if pin and (not pin.isdigit() or len(pin) != 4):
        raise HTTPException(status_code=400, detail="PIN must be 4 digits")

    subscription = await build_subscription(user["family_id"])
    if not is_admin_user(user):
        # Role-aware metering: kids and teens share one cap (parents/caregivers
        # are never the meter), so a household with two teens can't add a third
        # young person for free. Free = 2, Premium = 5.
        young_people = await count_young_people(database, user["family_id"])
        max_young = subscription["limits"].get("max_children", 2)
        if young_people >= max_young:
            plan_limit_error(
                feature="max_children",
                current_plan=subscription["plan"],
                message="Upgrade to Premium to add more (kids and teens share your plan's limit).",
                limit=max_young,
                used=young_people,
            )

    member = {
        "member_id": new_id("member"),
        "family_id": user["family_id"],
        "name": name,
        "role": "Child",
        "avatar": None,
        "stars": starting_stars,
        "pin_hash": hash_pin(pin) if pin else None,
        "created_at": utcnow(),
    }
    await database["family_members"].insert_one(member)

    if starting_stars:
        transaction = {
            "transaction_id": new_id("star"),
            "family_id": user["family_id"],
            "member_id": member["member_id"],
            "delta": starting_stars,
            "reason": "Starting stars",
            "created_by_user_id": user["user_id"],
            "created_by_name": user.get("name"),
            "created_at": utcnow(),
        }
        await database["star_transactions"].insert_one(transaction)

    return public_member(member)


@app.patch("/api/family/members/{member_id}")
async def update_family_member(member_id: str, payload: MemberPatchIn, user=Depends(require_full_member)):
    """Correct a child's details — a mistyped name, mostly.

    Without this a typo at setup was permanent short of deleting the child,
    which would have taken their stars and their history with it.
    """
    database = get_db()
    member = await database["family_members"].find_one(
        {"member_id": member_id, "family_id": user["family_id"]},
        {"_id": 0},
    )
    if not member:
        raise HTTPException(status_code=404, detail="Family member not found")

    changes: dict = {}
    old_name = member.get("name") or ""

    if payload.name is not None:
        name = payload.name.strip()
        if not name:
            raise HTTPException(status_code=400, detail="Name is required")
        if len(name) > 60:
            raise HTTPException(status_code=400, detail="Name is too long")
        # Task cards address a child by name, so two children answering to the
        # same one makes "who earned this star" unanswerable. Only a new clash
        # is refused — a household that already has two Sams keeps them.
        if name.lower() != old_name.lower():
            clash = await database["family_members"].find_one({
                "family_id": user["family_id"],
                "member_id": {"$ne": member_id},
                "name": {"$regex": f"^{re.escape(name)}$", "$options": "i"},
            })
            if clash:
                raise HTTPException(
                    status_code=400,
                    detail="Somebody in this household already has that name.",
                )
        changes["name"] = name

    if payload.avatar is not None:
        avatar = payload.avatar.strip()
        if len(avatar) > 200:
            raise HTTPException(status_code=400, detail="Avatar is too long")
        changes["avatar"] = avatar or None

    # getattr, not attribute access: callers construct their own payload stubs
    # and a hard reference makes adding a field break every one of them.
    payload_age = getattr(payload, "age", None)
    if payload_age is not None:
        age = int(payload_age)
        # 0 clears it. Otherwise 1-17: this is a child's row, and an age at or
        # above 18 would claim a grown-up lives behind a profile that has no
        # login and cannot consent to anything.
        if age == 0:
            changes["age"] = None
        elif not (1 <= age <= 17):
            raise HTTPException(status_code=400, detail="Age must be between 1 and 17.")
        elif age < 13 and member.get("user_id"):
            # The contradiction runs both ways: recording an under-13 age on
            # someone who already holds an account would leave the household
            # asserting two incompatible things about the same child. Say so
            # rather than quietly keeping both.
            raise HTTPException(
                status_code=400,
                detail=f"{member.get('name') or 'They'} has their own account, which is for 13 "
                       "and over. Remove the account first if that age is right.")
        else:
            changes["age"] = age

    if not changes:
        return public_member(member)

    await database["family_members"].update_one(
        {"member_id": member_id, "family_id": user["family_id"]},
        {"$set": changes},
    )

    # Several records address this member by name rather than by id — task
    # cards decide who earned a star by looking the name up. Without this a
    # rename would quietly break every card already assigned to them: the
    # lookup misses, the card is still flagged as paid, and nobody is.
    new_name = changes.get("name")
    if new_name and new_name != old_name and old_name:
        await database["cards"].update_many(
            {"family_id": user["family_id"], "assignee": old_name},
            {"$set": {"assignee": new_name}},
        )
        await database["handoff_notes"].update_many(
            {"family_id": user["family_id"], "member_id": member_id},
            {"$set": {"member_name": new_name}},
        )
        await database["expenses"].update_many(
            {"family_id": user["family_id"], "child_member_id": member_id},
            {"$set": {"child_name": new_name}},
        )

    updated = await database["family_members"].find_one(
        {"member_id": member_id, "family_id": user["family_id"]}, {"_id": 0}
    )
    return public_member(updated)


@app.delete("/api/family/members/{member_id}")
async def delete_family_member(member_id: str, user=Depends(require_full_member)):
    database = get_db()
    family_id = user["family_id"]
    member = await database["family_members"].find_one(
        {"member_id": member_id, "family_id": family_id},
        {"_id": 0},
    )
    if not member:
        raise HTTPException(status_code=404, detail="Member not found")

    # Removing a co-parent — an account-holder — is a different act from
    # removing a child, and gated accordingly. A child profile (no account)
    # skips all of this.
    if member.get("user_id"):
        remover = await _member_for_user(database, family_id, user)
        # Only a parent runs the household. Authority is decided like this:
        #  - a matched member row must carry a parent-level role;
        #  - a requester that matches NO member row at all is the founder whose
        #    row predates user_id linkage (or was never written). Every INVITED
        #    member — co-parent or grandparent/nanny — gets a user_id-linked row
        #    on join, so "no row" can only be the owner. Treat them as the
        #    parent they are; the founder / last-parent / self guards below still
        #    protect the household.
        role = remover.get("role")
        authorised = _is_parent_role(role) if remover else True
        if not authorised:
            raise HTTPException(
                status_code=403,
                detail=f"Only a parent can remove a co-parent. (role read here: {role or 'none'})")
        if member["user_id"] == user["user_id"] or (
                remover.get("member_id") and remover["member_id"] == member["member_id"]):
            raise HTTPException(status_code=400, detail="To leave the household yourself, use account deletion.")
        # The founder cannot be removed by a co-parent they added — otherwise a
        # co-parent could evict the owner. The founder is the EARLIEST parent, so
        # a row with no usable created_at must sort as oldest (a legacy founder
        # row may lack one); sorting it as "now" would hand founder status to a
        # later co-parent and strip the real owner's protection.
        parent_rows = [m async for m in database["family_members"].find(
            {"family_id": family_id,
             "role": {"$regex": "^(parent|co-parent)$", "$options": "i"}}, {"_id": 0})]
        founder = min(parent_rows,
                      key=lambda m: (_coerce_dt(m.get("created_at")) or _OLDEST_DT),
                      default=None)
        # A requester with no member row is the rowless owner (every invited
        # member gets one), so they are the founder — the member they are
        # removing, which does have a row, cannot be. Skip the founder and
        # last-parent guards for them: removing a co-parent still leaves the
        # owner in place.
        requester_is_rowless_owner = not remover.get("member_id")
        if not requester_is_rowless_owner:
            if founder and member["member_id"] == founder.get("member_id"):
                raise HTTPException(status_code=403, detail="The household founder cannot be removed.")
            # Never leave the household with no parent at all.
            if await _count_parents(database, family_id) <= 1:
                raise HTTPException(status_code=400, detail="A household must keep at least one parent.")

        # They keep their account — they are ejected from THIS household, not
        # deleted. A fresh, empty household means they are never left pointing
        # at one they can no longer see.
        removed = await database["users"].find_one({"user_id": member["user_id"]}, {"_id": 0})
        if removed:
            new_family_id = new_id("family")
            await _seed_new_family(database, removed, new_family_id,
                                   removed.get("email", ""), removed.get("name") or "Parent",
                                   removed.get("picture"))
            await database["users"].update_one(
                {"user_id": removed["user_id"]},
                {"$set": {"family_id": new_family_id, "updated_at": utcnow()}})
        # Any invite they accepted to get here is spent; drop stale pending ones.
        await database["family_invites"].update_many(
            {"family_id": family_id, "accepted_by_user_id": member["user_id"]},
            {"$set": {"status": "removed", "updated_at": utcnow()}})
        # Sever every live thread back into this household. Moving their user
        # doc to a new family_id is not enough on its own: an open session still
        # carries a cached view of the old household until it refreshes, and
        # their notification tokens still carry the OLD family_id, so they would
        # keep receiving this household's pushes. Drop both — the next sign-in
        # re-issues a session and re-registers a token against their new,
        # empty family, so removal is a clean break in both directions.
        await database["user_sessions"].delete_many({"user_id": member["user_id"]})
        await database["notification_tokens"].delete_many({"user_id": member["user_id"]})

    await database["family_members"].delete_one({"member_id": member_id, "family_id": family_id})

    # Everything that was only ever this child's goes with them.
    for collection in ("star_transactions", "redemptions", "allowances", "allowance_txns"):
        await database[collection].delete_many({"member_id": member_id, "family_id": family_id})

    # Shared things stay, but must stop pointing at somebody who is gone. A
    # chore left holding a dead member id rendered as a raw "member_a3f9…" on
    # the wheel and, worse, paid nobody on completion while still reporting
    # success. Hand it to whoever is left instead.
    async for chore in database["chores"].find(
        {"family_id": family_id, "assigned_members": member_id}, {"_id": 0}
    ):
        remaining = [m for m in (chore.get("assigned_members") or []) if m != member_id]
        current = chore.get("current_assignee")
        await database["chores"].update_one(
            {"chore_id": chore["chore_id"], "family_id": family_id},
            {"$set": {
                "assigned_members": remaining,
                "current_assignee": (
                    remaining[0] if current == member_id and remaining
                    else None if current == member_id
                    else current
                ),
            }},
        )
    await database["chores"].update_many(
        {"family_id": family_id, "current_assignee": member_id},
        {"$set": {"current_assignee": None}},
    )

    # A routine is a checklist worth keeping; it just belongs to nobody now,
    # which the completion path already handles.
    await database["routines"].update_many(
        {"family_id": family_id, "member_id": member_id},
        {"$set": {"member_id": None}},
    )

    # Expenses and handoff notes keep their denormalised name, which is exactly
    # why it is stored — the historical record should survive the person.
    return {"ok": True}


async def award_stars_to_member(
    database,
    family_id: str,
    member_id: str,
    delta: int,
    reason: str,
    user: dict,
    awarded_for=None,
) -> Optional[dict]:
    """Award stars and write the matching ledger entry, atomically.

    Shared by everything that can earn a child stars — finished chores,
    completed routines — so every star movement lands in the ledger that
    Recent Activity and the weekly total are built from. Silent no-op for a
    missing member or a non-child (a chore can sit with a parent, and paying a
    parent in stars would be odd), so callers can award unconditionally
    without branching.
    """
    if not member_id or delta <= 0:
        return None
    member = await database["family_members"].find_one(
        {"member_id": member_id, "family_id": family_id}, {"_id": 0}
    )
    if not member or str(member.get("role", "")).lower() not in ("child", "teen"):
        return None

    # Roll the weekly meter first so a star earned in a fresh week counts toward
    # the new week, not the old one. Earning banks the star (stars) and also
    # ticks the this-week meter (week_earned) that gates weekend treats.
    member = await roll_week_if_stale(database, member)
    await database["family_members"].update_one(
        {"member_id": member_id, "family_id": family_id},
        {"$inc": {"stars": delta, "week_earned": delta}},
    )
    await log_activity(database, user, "stars_awarded", member.get("name", ""), delta)
    # `created_at` is when this was recorded; `awarded_for` is the day the job
    # was actually done. They differ whenever a parent catches up on Sunday,
    # and only the second one can honestly place a star on a weekday.
    when = _coerce_dt(awarded_for) or utcnow()
    transaction = {
        "transaction_id": new_id("star"),
        "family_id": family_id,
        "member_id": member_id,
        "delta": delta,
        "reason": reason,
        "created_by_user_id": user.get("user_id"),
        "created_by_name": user.get("name"),
        "created_at": utcnow(),
        "awarded_for": when,
    }
    await database["star_transactions"].insert_one(transaction)
    await _pay_full_week_bonus(database, family_id, member_id, user)
    return transaction


async def _pay_full_week_bonus(database, family_id: str, member_id: str, user: dict) -> None:
    """The last star of a complete week, paid once.

    Guarded on the member document rather than counted twice: the seventh
    active day can be reached by two awards landing together, and a bonus paid
    twice would put the meter above a target nobody earned.
    """
    week_start = current_week_start()
    days = set()
    async for txn in database["star_transactions"].find(
        {"family_id": family_id, "member_id": member_id},
        {"_id": 0, "awarded_for": 1, "created_at": 1, "delta": 1},
    ):
        if int(txn.get("delta") or 0) <= 0:
            continue
        when = _coerce_dt(txn.get("awarded_for")) or _coerce_dt(txn.get("created_at"))
        if when and when >= week_start:
            days.add(when.date())
    if len(days) < FULL_WEEK_DAYS:
        return

    claimed = await database["family_members"].update_one(
        {"member_id": member_id, "family_id": family_id,
         "week_bonus_for": {"$ne": week_start}},
        {"$set": {"week_bonus_for": week_start},
         "$inc": {"stars": FULL_WEEK_BONUS, "week_earned": FULL_WEEK_BONUS}},
    )
    if claimed.matched_count:
        await database["star_transactions"].insert_one({
            "transaction_id": new_id("star"), "family_id": family_id,
            "member_id": member_id, "delta": FULL_WEEK_BONUS,
            "reason": "Full week", "created_by_user_id": user.get("user_id"),
            "created_by_name": user.get("name"),
            "created_at": utcnow(), "awarded_for": utcnow()})


@app.post("/api/family/members/{member_id}/stars")
async def adjust_member_stars(member_id: str, payload: StarAdjustmentIn, user=Depends(require_full_member)):
    database = get_db()
    member = await database["family_members"].find_one(
        {"member_id": member_id, "family_id": user["family_id"]},
        {"_id": 0},
    )
    if not member:
        raise HTTPException(status_code=404, detail="Family member not found")

    delta = int(payload.delta or 0)
    if delta == 0:
        raise HTTPException(status_code=400, detail="Star adjustment cannot be zero")

    current_stars = int(member.get("stars", 0))
    if current_stars + delta < 0:
        raise HTTPException(status_code=400, detail="Stars cannot go below zero")

    reason = (payload.reason or "").strip()
    if delta < 0 and not reason:
        raise HTTPException(status_code=400, detail="Reason is required when removing stars")

    member = await roll_week_if_stale(database, member)

    # A star can be credited to an earlier day this week — a parent who catches
    # up on Sunday should not have Tuesday's job land on Sunday. Bounded to the
    # current week in both directions: the future is not a day anyone has
    # worked, and a settled week may already have paid out a treat.
    awarded_for = utcnow()
    if payload.awarded_for:
        parsed = _coerce_dt(parse_dt(payload.awarded_for))
        if not parsed:
            raise HTTPException(status_code=400, detail="Invalid date")
        if parsed < current_week_start() or parsed > utcnow():
            raise HTTPException(status_code=400, detail="Pick a day from this week")
        awarded_for = parsed

    transaction = {
        "transaction_id": new_id("star"),
        "family_id": user["family_id"],
        "member_id": member_id,
        "delta": delta,
        "reason": reason or ("Parent added stars" if delta > 0 else "Parent removed stars"),
        "created_by_user_id": user["user_id"],
        "created_by_name": user.get("name"),
        "created_at": utcnow(),
        "awarded_for": awarded_for,
    }

    # Increment rather than write a total computed from an earlier read: two
    # concurrent adjustments both read the same balance and the second $set
    # overwrote the first, silently losing one. For removals the filter also
    # carries the "cannot go below zero" rule, so the database enforces it even
    # when two removals race.
    #
    # A positive adjustment is a child earning — it banks AND ticks the weekly
    # meter, exactly like a finished chore. A removal is a correction against
    # the bank only; it leaves the meter alone, so undoing a mis-award does not
    # quietly hand back a weekend treat the child never earned.
    star_filter = {"member_id": member_id, "family_id": user["family_id"]}
    inc = {"stars": delta}
    if delta > 0:
        inc["week_earned"] = delta
    else:
        star_filter["stars"] = {"$gte": -delta}
    result = await database["family_members"].update_one(star_filter, {"$inc": inc})
    if result.matched_count == 0:
        raise HTTPException(status_code=400, detail="Stars cannot go below zero")

    await database["star_transactions"].insert_one(transaction)
    if delta > 0:
        # Filling in a missed day can be the thing that completes the week.
        await _pay_full_week_bonus(database, user["family_id"], member_id, user)
    updated = await database["family_members"].find_one({"member_id": member_id}, {"_id": 0})
    new_total = int(updated.get("stars", 0)) if updated else current_stars + delta

    if delta > 0:
        await send_star_milestone_alert(user["family_id"], member.get("name", "Your child"), current_stars, new_total)

    return {
        "ok": True,
        "member": public_member(updated),
        "transaction": public_star_transaction(transaction),
    }


@app.post("/api/family/members/{member_id}/weekly-claim")
async def claim_weekly_treat(member_id: str, payload: WeeklyClaimIn, user=Depends(require_user)):
    """Cash the week in for a treat.

    A loyalty card, not a locked door. The child can claim their treat at any
    point in the week — reaching 50 is the celebration, not the price of entry.
    Which treat is a conversation, not a price list, so nothing here spends
    `stars`: the bank is savings, untouched, and this claims the WEEK.

    The one rule is one treat per week, guarded on the member document so two
    parents tapping together cannot hand out two for one week — and so an
    always-available treat cannot become an unlimited one. The meter is not
    reset: it records what was earned, and hitting 50 later still celebrates.
    """
    database = get_db()
    member = await database["family_members"].find_one(
        {"member_id": member_id, "family_id": user["family_id"]}, {"_id": 0})
    if not member:
        raise HTTPException(status_code=404, detail="Family member not found")

    member = await roll_week_if_stale(database, member)
    week_start = current_week_start()
    title = sanitize_user_text(payload.title, 80)
    if len(title) < 2:
        raise HTTPException(status_code=400, detail="Give the treat a name")

    # No star floor: a treat is claimable at any point in the week. The only
    # guard is one-per-week — the atomic `$ne` filter makes a double tap
    # (or two parents at once) a no-op rather than a second treat.
    claimed = await database["family_members"].update_one(
        {"member_id": member_id, "family_id": user["family_id"],
         "week_claimed_for": {"$ne": week_start}},
        {"$set": {"week_claimed_for": week_start}},
    )
    if claimed.matched_count == 0:
        raise HTTPException(status_code=400, detail="This week's treat has already been claimed.")

    # Recorded as a redemption so it appears wherever "what are we still owed?"
    # is answered — at no star cost, because the week already paid for it.
    redemption = {
        "redemption_id": new_id("redemption"), "family_id": user["family_id"],
        "member_id": member_id, "reward_id": None,
        "reward_title": title, "reward_icon": None,
        "cost_stars": 0, "status": "pending", "created_at": utcnow(),
        "weekly": True, "week_start": week_start,
        "created_by_user_id": user["user_id"], "fulfilled_at": None}
    await database["redemptions"].insert_one(redemption)

    return {"ok": True, "redemption": public_redemption(redemption)}


@app.get("/api/family/members/{member_id}/star-history")
async def member_star_history(member_id: str, user=Depends(require_user)):
    database = get_db()
    member = await database["family_members"].find_one(
        {"member_id": member_id, "family_id": user["family_id"]},
        {"_id": 0},
    )
    if not member:
        raise HTTPException(status_code=404, detail="Family member not found")

    # 200, not 50. The Kids screen buckets these by weekday to draw the week's
    # momentum row, and a family using the quick-add chips daily crosses 50
    # entries by midweek — so Monday and Tuesday silently rendered as empty
    # boxes next to a meter that counted them. "You did nothing on Monday", on
    # a child's own screen, because of a query limit.
    cursor = database["star_transactions"].find(
        {"member_id": member_id, "family_id": user["family_id"]},
        {"_id": 0},
    ).sort("created_at", -1).limit(200)

    return [public_star_transaction(item) async for item in cursor]



@app.put("/api/family/members/{member_id}/pin")
async def set_member_pin(member_id: str, payload: PinIn, user=Depends(require_full_member)):
    database = get_db()
    member = await database["family_members"].find_one(
        {"member_id": member_id, "family_id": user["family_id"]},
        {"_id": 0},
    )
    if not member:
        raise HTTPException(status_code=404, detail="Member not found")

    pin = payload.pin.strip()
    if pin and (len(pin) != 4 or not pin.isdigit()):
        raise HTTPException(status_code=400, detail="PIN must be exactly 4 digits")

    pin_hash = hash_pin(pin) if pin else None
    await database["family_members"].update_one(
        {"member_id": member_id, "family_id": user["family_id"]},
        {"$set": {"pin_hash": pin_hash}},
    )
    return {"ok": True, "has_pin": bool(pin_hash)}




@app.delete("/api/family/members/{member_id}/pin")
async def remove_member_pin(member_id: str, user=Depends(require_full_member)):
    database = get_db()
    member = await database["family_members"].find_one(
        {"member_id": member_id, "family_id": user["family_id"]},
        {"_id": 0},
    )
    if not member:
        raise HTTPException(status_code=404, detail="Family member not found")

    await database["family_members"].update_one(
        {"member_id": member_id, "family_id": user["family_id"]},
        {"$set": {"pin_hash": None}},
    )
    return {"ok": True, "has_pin": False}

@app.post("/api/family/members/{member_id}/verify-pin")
async def verify_member_pin(member_id: str, payload: PinIn, user=Depends(require_user)):
    database = get_db()
    member = await database["family_members"].find_one(
        {"member_id": member_id, "family_id": user["family_id"]},
        {"_id": 0},
    )
    if not member:
        raise HTTPException(status_code=404, detail="Member not found")

    if not member.get("pin_hash"):
        return {"ok": True, "has_pin": False}

    identity = f"pin:{member_id}"
    if _auth_locked(identity):
        raise HTTPException(status_code=429, detail="Too many attempts. Please try again later.")

    if not await _verify_pin_and_upgrade(database, member, payload.pin.strip()):
        _auth_record_fail(identity)
        raise HTTPException(status_code=401, detail="Invalid PIN")

    _auth_clear(identity)
    return {"ok": True, "has_pin": True}


async def _enforce_member_slot_limit(database, user) -> None:
    """Raise a plan-limit error when the household (members + pending
    invites) has no free slot. Admins bypass the check."""
    sub = await build_subscription(user["family_id"])
    limit = sub["limits"]["max_members"]
    used = sub["members_count"]
    pending_invites_count = await database["family_invites"].count_documents(
        {
            "family_id": user["family_id"],
            "status": "pending",
            "expires_at": {"$gt": utcnow()},
        }
    )
    used_with_pending = used + pending_invites_count
    if not is_admin_user(user) and used_with_pending >= limit:
        plan_limit_error(
            feature="family_members",
            current_plan=sub["plan"],
            limit=limit,
            used=used_with_pending,
            message=f"Your current plan allows {limit} family member slots including pending invites. Upgrade to add more.",
        )


def _new_invite_doc(user, email=None, relationship=None, label=None, is_teen=False, age=None, is_helper=False) -> dict:
    now = utcnow()
    return {
        "invite_id": new_id("invite"),
        "family_id": user["family_id"],
        "email": email,
        "relationship": relationship,
        "is_teen": bool(is_teen),
        "is_helper": bool(is_helper),
        "age": age,
        "label": label,
        "token": secrets.token_urlsafe(24),
        "status": "pending",
        "created_by_user_id": user["user_id"],
        "created_by_name": user.get("name"),
        "created_by_email": user.get("email"),
        "created_at": now,
        "updated_at": now,
        "expires_at": now + timedelta(days=INVITE_DAYS),
        "accepted_at": None,
        "accepted_by_user_id": None,
        "accepted_by_email": None,
    }


class InviteLinkIn(BaseModel):
    # Who this link is meant for and what role they get on accepting —
    # a bare anonymous link gives the inviter no way to tell links apart
    # or control what the holder becomes.
    relationship: Optional[str] = None
    label: Optional[str] = None
    is_helper: Optional[bool] = False


def _clean_invite_text(value: Optional[str], cap: int = 32) -> Optional[str]:
    return re.sub(r"\s+", " ", (value or "").strip())[:cap] or None


@app.post("/api/family/invite/link")
async def family_invite_link(payload: Optional[InviteLinkIn] = Body(None), user=Depends(require_full_member)):
    """Create a shareable invite link with no email attached — used by the
    Phone (SMS) and Share-link options, which deliver the link from the
    inviter's own device. Carries an optional intended-recipient label and
    relationship; whoever accepts is recorded (accepted_by_email), so the
    inviter can always see which account used which link."""
    database = get_db()
    is_helper = bool(payload.is_helper if payload else False)
    await _enforce_member_slot_limit(database, user)
    # A helper is never a co-parent, so it skips the two-parent cap (an empty
    # relationship would otherwise be read as a co-parent link).
    if not is_helper:
        await _enforce_parent_limit(database, user["family_id"],
                                    payload.relationship if payload else None)
    invite = _new_invite_doc(
        user,
        email=None,
        relationship=_clean_invite_text(payload.relationship if payload else None),
        label=_clean_invite_text(payload.label if payload else None, cap=48),
        is_helper=is_helper,
    )
    await database["family_invites"].insert_one(invite)
    public = public_invite(invite)
    return {"ok": True, "invite": public, "invite_url": public["invite_url"]}


@app.post("/api/family/invite")
async def family_invite(payload: InviteIn, user=Depends(require_full_member)):
    database = get_db()
    await _enforce_member_slot_limit(database, user)
    # A teen or a helper is never a parent, so both skip the co-parent cap.
    # Without this, an invite with no relationship set is read as a co-parent and
    # rejected once the household already has two parents.
    if not payload.is_teen and not payload.is_helper:
        await _enforce_parent_limit(database, user["family_id"], payload.relationship)

    email = payload.email.strip().lower()
    if not email or "@" not in email:
        raise HTTPException(status_code=400, detail="Valid email is required")

    # The age gate is the COPPA/Families-policy line, enforced HERE and not only
    # in the client: a restricted young-person account is for 13-17. The 13
    # floor is the hard rule (no under-13 independent account); 17 is the ceiling
    # (at 18 they're an adult, not a dependent teen).
    if payload.is_teen and (payload.age is None or not (13 <= int(payload.age) <= 17)):
        raise HTTPException(status_code=400, detail="A restricted account is for ages 13 to 17.")

    # And it has to agree with what the household already recorded about this
    # child. Without this the 13 floor is only as good as the number typed into
    # the invite: a child on record as eight could be handed an account by
    # claiming fifteen a screen later. The recorded age wins.
    if payload.is_teen and payload.member_id:
        child = await database["family_members"].find_one(
            {"member_id": payload.member_id, "family_id": user["family_id"]}, {"_id": 0})
        if not child:
            raise HTTPException(status_code=404, detail="Family member not found")
        recorded = child.get("age")
        if recorded is not None:
            if int(recorded) < 13:
                raise HTTPException(
                    status_code=400,
                    detail=f"{child.get('name') or 'This child'} is recorded as {int(recorded)}. "
                           "An account of their own is for 13 and over — they can keep using "
                           "kid mode on your phone.")
            if int(recorded) != int(payload.age):
                raise HTTPException(
                    status_code=400,
                    detail=f"{child.get('name') or 'This child'} is recorded as {int(recorded)}, "
                           "not {}. Correct their age first.".format(int(payload.age)))

    # Teens share the kids' plan cap — enforce it here so the invite can't be the
    # free way around the limit the managed-child add already guards. The client
    # turns this 402 into the "household is full -> see plans" prompt.
    if payload.is_teen and not is_admin_user(user):
        sub = await build_subscription(user["family_id"])
        max_young = sub["limits"].get("max_children", 2)
        young = await count_young_people(database, user["family_id"])
        if young >= max_young:
            plan_limit_error(
                feature="max_children",
                current_plan=sub["plan"],
                message="Upgrade to add more (kids and teens share your plan's limit).",
                limit=max_young,
                used=young,
            )

    existing = await database["family_invites"].find_one(
        {
            "family_id": user["family_id"],
            "email": email,
            "status": "pending",
            "expires_at": {"$gt": utcnow()},
        },
        {"_id": 0},
    )

    relationship = _clean_invite_text(payload.relationship)

    if existing:
        invite = existing
        updates = {}
        if relationship and invite.get("relationship") != relationship:
            # Re-sending with a (new) relationship refreshes it — the last
            # thing the inviter typed is what they meant.
            updates["relationship"] = relationship
        # Keep the teen flag/age in sync with THIS send. Without it, sending a
        # teen invite to an address that already had a plain invite would reuse
        # the old doc and silently drop is_teen — the restriction failing open.
        if bool(invite.get("is_teen")) != bool(payload.is_teen) or invite.get("age") != payload.age:
            updates["is_teen"] = bool(payload.is_teen)
            updates["age"] = payload.age
        # Same fail-closed sync for the helper flag: reusing a plain invite for a
        # helper send must carry the restriction, not silently drop it.
        if bool(invite.get("is_helper")) != bool(payload.is_helper):
            updates["is_helper"] = bool(payload.is_helper)
        if updates:
            updates["updated_at"] = utcnow()
            await database["family_invites"].update_one(
                {"invite_id": invite["invite_id"]}, {"$set": updates})
            invite.update(updates)
    else:
        invite = _new_invite_doc(user, email=email, relationship=relationship,
                                 is_teen=bool(payload.is_teen), age=payload.age,
                                 is_helper=bool(payload.is_helper))
        await database["family_invites"].insert_one(invite)

    public = public_invite(invite)

    # If this address already belongs to an account, tell that person inside
    # the app too — invitation emails are where invitations go to die.
    invited_user = await database["users"].find_one({"email": email}, {"_id": 0})
    if invited_user and invited_user.get("family_id") != user["family_id"]:
        lang = invited_user.get("language") or "en"
        L = PUSH_I18N.get(lang, PUSH_I18N["en"])
        await send_push_to_user(
            database, invited_user["user_id"],
            L["invited_title"].format(name=user.get("name") or "A parent"),
            L["invited_body"],
            {"type": "family_invite", "token": invite["token"]},
        )

    email_result = await send_invite_email(
        email,
        public["invite_url"],
        user.get("name") or user.get("email") or "A family member",
        user.get("email") or "",
        relationship=invite.get("relationship") or "",
    )

    if email_result.get("sent"):
        message = f"Invitation email sent to {email}."
    else:
        message = "Invitation link created, but email delivery failed. Share the link manually."

    return {
        "ok": True,
        "sent": bool(email_result.get("sent")),
        "status": public["status"],
        "message": message,
        "invite": public,
        "invite_url": public["invite_url"],
        "email_provider": email_result.get("provider"),
        "email_error": email_result.get("error"),
    }


@app.get("/api/family/invites")
async def family_invites(user=Depends(require_user)):
    database = get_db()
    rows = []
    cursor = database["family_invites"].find(
        {"family_id": user["family_id"]},
        {"_id": 0},
    ).sort("created_at", -1)

    async for item in cursor:
        expires_at = ensure_aware_utc(item.get("expires_at"))
        if item.get("status") == "pending" and expires_at and _expired(expires_at):
            item["expires_at"] = expires_at
            item["status"] = "expired"
            await database["family_invites"].update_one(
                {"invite_id": item["invite_id"]},
                {"$set": {"status": "expired", "updated_at": utcnow()}},
            )
        rows.append(public_invite(item))

    return rows


@app.get("/api/family/invites/for-me")
async def invites_for_me(
    redeem: Optional[str] = None,
    x_redeem: Optional[str] = Header(None, alias="X-Redeem"),
    user=Depends(require_user),
):
    """Pending invites addressed to the signed-in user's own email.

    The join prompt asks the server directly instead of relying on the
    invite link surviving the trip through a junk folder, a cached page
    or a copy-pasted URL — signing in is enough to be told.

    ?redeem=<token> — or the X-Redeem header — accepts over THIS URL. The
    once-and-for-all invariant: the card only appears because this exact
    request succeeded moments earlier, so this request shape provably passes
    whatever blockers and middleboxes the device has — anyone who can SEE
    the offer can take it. The header variant exists because one field
    blocker also matched query strings; content-blocker rules match URLs
    and cannot see headers, making the header request byte-identical in URL
    and method to the proven one.
    """
    # isinstance guard: called outside FastAPI (tests), x_redeem is the
    # Header default object, not a string.
    redeem_token = redeem or (x_redeem if isinstance(x_redeem, str) else None)
    if redeem_token:
        return await _accept_invite_request(redeem_token, user)
    database = get_db()
    email = (user.get("email") or "").strip().lower()
    if not email:
        return []

    rows = []
    cursor = database["family_invites"].find(
        {"email": email, "status": "pending"},
        {"_id": 0},
    )
    async for item in cursor:
        expires_at = ensure_aware_utc(item.get("expires_at"))
        if expires_at and _expired(expires_at):
            continue
        if item.get("family_id") == user.get("family_id"):
            continue
        inviter = await database["users"].find_one(
            {"user_id": item.get("created_by_user_id")},
            {"_id": 0},
        )
        rows.append({
            "token": item["token"],
            "inviter_name": (inviter or {}).get("name")
                or item.get("created_by_name") or "A family member",
            "relationship": item.get("relationship"),
        })
    return rows


@app.get("/api/activity")
async def list_activity(limit: int = 12, user=Depends(require_user)):
    """What happened in this household lately, newest first.

    Everything you do is yours: the feed carries the household's shared lines
    plus your own, and nothing a co-parent did privately. `shared: {$ne:
    False}` keeps legacy rows (written before privacy existed, with no flag)
    visible to all, exactly as they were — only rows explicitly marked private
    are held back. `hidden_by` drops lines this person chose to clear from
    their own view without erasing them for anyone else.
    """
    database = get_db()
    me = user["user_id"]
    rows = []
    cursor = database["activity"].find(
        {
            "family_id": user["family_id"],
            "hidden_by": {"$ne": me},
            "$or": [{"shared": {"$ne": False}}, {"actor_user_id": me}],
        },
        {"_id": 0},
    ).sort("created_at", -1).limit(max(1, min(limit, 50)))
    async for row in cursor:
        rows.append(public_activity(row))
    return rows


@app.delete("/api/activity/{activity_id}")
async def delete_activity(activity_id: str, user=Depends(require_user)):
    """Clear a line from the feed.

    A private line is yours alone — deleting it removes the record. A shared
    line belongs to the household, so "delete" means "hide it from my view":
    the row stays for the co-parent, it just leaves your feed. Either way you
    can only touch lines your own feed actually shows.
    """
    database = get_db()
    me = user["user_id"]
    row = await database["activity"].find_one(
        {"activity_id": activity_id, "family_id": user["family_id"]}, {"_id": 0}
    )
    if not row:
        raise HTTPException(status_code=404, detail="Not found")
    if row.get("shared", True) is False:
        # Private: only the actor ever saw it, and only the actor may remove it.
        if row.get("actor_user_id") != me:
            raise HTTPException(status_code=403, detail="Not yours")
        await database["activity"].delete_one({"activity_id": activity_id})
        return {"ok": True, "deleted": True}
    # Shared: hide for me, keep for everyone else.
    await database["activity"].update_one(
        {"activity_id": activity_id},
        {"$addToSet": {"hidden_by": me}},
    )
    return {"ok": True, "hidden": True}


SEARCH_LIMIT = 40


def _search_hit(kind: str, item_id: str, title: str, subtitle: str = "",
                when: str = "", status: str = "") -> dict:
    return {"kind": kind, "id": item_id, "title": title,
            "subtitle": subtitle, "when": when, "status": status}


# -----------------------------------------------------------------------------
# Kid mode
#
# A child does not get an account. They get a profile on a device a parent has
# already signed into, entered with their own PIN, and a much smaller app:
# their stars, their chores, their rewards. Not the household.
#
# The session they hold is a real session row marked kind="child", which
# require_user refuses outright — so the only doors a child token opens are
# the handful below, and each one re-checks the member it belongs to.
# -----------------------------------------------------------------------------


class ChildSessionIn(BaseModel):
    member_id: str
    pin: str


class ParentPinIn(BaseModel):
    pin: str


class KidForgotPinIn(BaseModel):
    email: str
    password: str


async def require_teen(authorization: str = Header(default="")):
    """Resolve a teen's own login, and prove it IS a teen account.

    The mirror image of require_user's teen refusal: this is the only gate the
    /api/teen/* allowlist trusts, and it accepts a token ONLY if the account is
    flagged is_teen. So a normal parent token can't reach teen routes and a
    teen token can't reach anything else — the two never overlap.
    """
    database = get_db()
    if not authorization.startswith("Bearer "):
        raise HTTPException(status_code=401, detail="Missing bearer token")
    token = authorization.replace("Bearer ", "", 1).strip()
    session = await database["user_sessions"].find_one(
        {"token_hash": sha256(token), "expires_at": {"$gt": utcnow()}}, {"_id": 0})
    if not session or session.get("kind") == "child":
        raise HTTPException(status_code=401, detail="Invalid or expired session")
    # A teen signs in on their own device; being logged out weekly is the same
    # papercut it is for a parent. Kid mode is deliberately left out — that is a
    # hand-over on someone else's phone and should lapse.
    await _touch_session(database, sha256(token), session)
    user = await database["users"].find_one({"user_id": session["user_id"]}, {"_id": 0})
    if not user or not user.get("is_teen"):
        raise HTTPException(status_code=403, detail="Not a teen account")
    member = await _member_for_user(database, user["family_id"], user)
    return {"user": user, "member": member, "family_id": user["family_id"]}


async def require_child(authorization: str = Header(default="")):
    """Resolve a kid-mode session to the child it belongs to."""
    database = get_db()
    if not authorization.startswith("Bearer "):
        raise HTTPException(status_code=401, detail="Missing bearer token")
    token = authorization.replace("Bearer ", "", 1).strip()
    session = await database["user_sessions"].find_one(
        {"token_hash": sha256(token), "expires_at": {"$gt": utcnow()}}, {"_id": 0})
    if not session or session.get("kind") != "child":
        raise HTTPException(status_code=401, detail="Not a kid session")
    member = await database["family_members"].find_one(
        {"member_id": session.get("member_id")}, {"_id": 0})
    if not member:
        raise HTTPException(status_code=401, detail="Profile not found")
    return {"member": member, "family_id": member["family_id"],
            "parent_user_id": session.get("user_id")}


async def _family_parent_pin_members(database, family_id: str) -> list:
    """The grown-ups in this family who have set a PIN. Whole rows, not just the
    hashes: a PIN stored in the old bare-hash form is upgraded when it is used,
    and that needs to know which member it belonged to."""
    members = []
    async for m in database["family_members"].find({"family_id": family_id}, {"_id": 0}):
        if (m.get("role") or "").lower() != "child" and m.get("pin_hash"):
            members.append(m)
    return members


@app.get("/api/family/profiles")
async def list_profiles(user=Depends(require_user)):
    """Who could be using this device — the "Who's this?" picker.

    Kid mode is only offered once a grown-up has a PIN of their own; without
    one, handing the tablet to a child would be a one-way door.
    """
    database = get_db()
    rows, parent_pin = [], False
    async for m in database["family_members"].find({"family_id": user["family_id"]}, {"_id": 0}):
        is_child = (m.get("role") or "").lower() == "child"
        if not is_child and m.get("pin_hash"):
            parent_pin = True
        rows.append({"member_id": m["member_id"], "name": m.get("name") or "",
                     "role": m.get("role") or "Parent", "is_child": is_child,
                     "has_pin": bool(m.get("pin_hash")),
                     # Which row is the person holding the phone — so the
                     # hand-over sheet can offer to set THEIR PIN inline
                     # rather than sending them off to find it.
                     "is_me": m.get("user_id") == user["user_id"]})
    return {"profiles": rows, "kid_mode_ready": parent_pin}


@app.post("/api/kid/session")
async def start_kid_session(payload: ChildSessionIn, user=Depends(require_user)):
    """Swap this device into a child's view. Requires the child's own PIN."""
    database = get_db()
    member = await database["family_members"].find_one(
        {"member_id": payload.member_id, "family_id": user["family_id"]}, {"_id": 0})
    if not member or (member.get("role") or "").lower() != "child":
        raise HTTPException(status_code=404, detail="Child not found")
    if not member.get("pin_hash"):
        raise HTTPException(status_code=400, detail="This child has no PIN yet")
    if not await _family_parent_pin_members(database, user["family_id"]):
        raise HTTPException(status_code=400, detail="Set a parent PIN first")

    identity = f"kid:{member['member_id']}"
    if _auth_locked(identity):
        raise HTTPException(status_code=429, detail="Too many attempts. Please try again later.")
    if not await _verify_pin_and_upgrade(database, member, payload.pin.strip()):
        _auth_record_fail(identity)
        raise HTTPException(status_code=401, detail="Invalid PIN")
    _auth_clear(identity)

    raw = secrets.token_urlsafe(32)
    await database["user_sessions"].insert_one({
        "session_id": new_id("sess"),
        "user_id": user["user_id"],
        "member_id": member["member_id"],
        "kind": "child",
        "token_hash": sha256(raw),
        # Short by design: a device left in kid mode returns to a locked
        # picker the next day rather than staying open indefinitely.
        "expires_at": utcnow() + timedelta(days=1),
        "created_at": utcnow(),
    })
    return {"session_token": raw, "member": public_member(member)}


@app.post("/api/kid/exit")
async def exit_kid_session(payload: ParentPinIn, child=Depends(require_child),
                           authorization: str = Header(default="")):
    """Leave kid mode. A grown-up's PIN, so a child cannot let themselves out."""
    database = get_db()
    grownups = await _family_parent_pin_members(database, child["family_id"])
    identity = f"kidexit:{child['family_id']}"
    if _auth_locked(identity):
        raise HTTPException(status_code=429, detail="Too many attempts. Please try again later.")
    given = payload.pin.strip()
    # Every grown-up's PIN is checked, not just the first that matches, so the
    # work done does not depend on whose PIN was typed.
    matched = False
    for m in grownups:
        if await _verify_pin_and_upgrade(database, m, given):
            matched = True
    if not matched:
        _auth_record_fail(identity)
        raise HTTPException(status_code=401, detail="Invalid PIN")
    _auth_clear(identity)
    # Exit means exit: kill the kid session server-side rather than trusting
    # the client to discard it — otherwise the token stayed valid for kid
    # routes for its full 24 h after the parent believed kid mode was over.
    token = authorization.replace("Bearer ", "", 1).strip()
    await database["user_sessions"].delete_many({"token_hash": sha256(token)})
    return {"ok": True}


@app.post("/api/kid/exit-forgot-pin")
async def exit_kid_forgot_pin(payload: KidForgotPinIn, child=Depends(require_child),
                              authorization: str = Header(default="")):
    """Leave kid mode when the parent PIN has been forgotten.

    Without this, a parent who forgot the PIN was locked in the child's app
    with no way out — the one-way door the PIN was meant to prevent, turned on
    the grown-up. A child cannot use it: it needs a PARENT'S account email and
    password (a real credential the child does not have), not the PIN. On
    success the forgotten PIN is cleared so a fresh one is set next time.
    """
    database = get_db()
    identity = f"kidexitpw:{child['family_id']}"
    if _auth_locked(identity):
        raise HTTPException(status_code=429, detail="Too many attempts. Please try again later.")

    email = (payload.email or "").strip().lower()
    account = await database["users"].find_one({"email": email}, {"_id": 0})
    member = None
    if account and account.get("family_id") == child["family_id"]:
        member = await database["family_members"].find_one(
            {"family_id": child["family_id"], "user_id": account["user_id"]}, {"_id": 0})

    ok = bool(
        account and member
        and (member.get("role") or "").lower() != "child"
        and account.get("password_hash")
        and verify_password(payload.password or "", account["password_hash"])
    )
    if not ok:
        _auth_record_fail(identity)
        raise HTTPException(status_code=401, detail="Email or password not recognised.")

    _auth_clear(identity)
    # Clear the forgotten PIN so the hand-over sheet prompts for a new one.
    await database["family_members"].update_one(
        {"member_id": member["member_id"]}, {"$set": {"pin_hash": None}})
    # Same rule as the PIN exit: leaving kid mode kills the kid session
    # server-side rather than trusting the client to discard the token.
    kid_token = authorization.replace("Bearer ", "", 1).strip()
    await database["user_sessions"].delete_many({"token_hash": sha256(kid_token)})
    return {"ok": True}


@app.get("/api/kid/home")
async def kid_home(child=Depends(require_child)):
    """Everything a child's app shows, in one call.

    Scoped to this child on the server, not filtered on the device: their
    stars, the chores with their name on them, what they can spend on, and
    what they are still owed.
    """
    database = get_db()
    member = await roll_week_if_stale(database, child["member"])
    name = (member.get("name") or "").strip().lower()

    chores = []
    async for card in database["cards"].find(
        {"family_id": child["family_id"], "status": {"$ne": "DONE"}, "shared": True},
        {"_id": 0},
    ):
        if (card.get("assignee") or "").strip().lower() == name:
            chores.append({"card_id": card["card_id"], "title": card.get("title") or "",
                           "due_date": iso(card.get("due_date"))})
    chores.sort(key=lambda c: (c["due_date"] is None, c["due_date"] or ""))

    rewards = [public_reward(r) async for r in
               database["rewards"].find({"family_id": child["family_id"]}, {"_id": 0})]
    rewards.sort(key=lambda r: r["cost_stars"])

    owed = [public_redemption(r) async for r in database["redemptions"].find(
        {"family_id": child["family_id"], "member_id": member["member_id"],
         "status": "pending"}, {"_id": 0})]

    return {"name": member.get("name") or "", "stars": int(member.get("stars") or 0),
            "week_earned": max(0, int(member.get("week_earned") or 0)),
            "weekend_goal_reward_id": member.get("weekend_goal_reward_id"),
            "chores": chores, "rewards": rewards, "owed": owed}


# ── Teen mode: a restricted 13-17 account. The ONLY endpoints a teen token can
#    reach (require_user refuses it everywhere else). Everything here is scoped
#    server-side to the teen themselves — never the family's calendar, vault,
#    other members or settings. ──

def _teen_can_see(card: dict, teen_name: str, teen_user_id: str) -> bool:
    """A teen sees a card only if it is family-wide (shared) OR theirs — assigned
    to them, or created by them. Another member's private card never matches."""
    if card.get("shared") is True:
        return True
    # A legacy card that predates the privacy model (no owner, no shared flag)
    # is family-wide — same rule public_card uses — so teens see it too.
    if card.get("shared") is None and card.get("created_by_user_id") is None:
        return True
    if (card.get("assignee") or "").strip().lower() == teen_name:
        return True
    if teen_user_id and card.get("created_by_user_id") == teen_user_id:
        return True
    return False


@app.get("/api/teen/me")
async def teen_me(teen=Depends(require_teen)):
    """The teen's own basic identity — no family roster, no other members."""
    user = teen["user"]
    return {"user_id": user["user_id"], "name": user.get("name") or "",
            "email": user.get("email"), "family_id": user["family_id"],
            "language": user.get("language", "en"), "is_teen": True}


@app.get("/api/teen/home")
async def teen_home(teen=Depends(require_teen)):
    """A teen's whole world in one call: their tasks and their agenda.

    Tasks = jobs with their name on them (or that they created). Agenda =
    family-wide events plus their own. Scoped on the server, so a teen device
    is never even sent a parent-private item to filter out.
    """
    database = get_db()
    user = teen["user"]
    name = (user.get("name") or "").strip().lower()
    uid = user.get("user_id")
    # Their star balance — teens earn stars like the younger kids, but a parent
    # approves the star after the teen ticks the task (no auto-award here).
    member = teen.get("member") or {}
    if member:
        member = await roll_week_if_stale(database, member)

    tasks, agenda = [], []
    async for card in database["cards"].find({"family_id": user["family_id"]}, {"_id": 0}):
        if not _teen_can_see(card, name, uid):
            continue
        ctype = (card.get("type") or "").upper()
        row = {"card_id": card["card_id"], "title": card.get("title") or "",
               "due_date": iso(card.get("due_date")),
               "status": card.get("status"), "assignee": card.get("assignee")}
        if ctype == "EVENT":
            agenda.append(row)
        else:  # TASK / note-style → the teen's to-dos
            if card.get("status") != "DONE":
                tasks.append(row)
    tasks.sort(key=lambda c: (c["due_date"] is None, c["due_date"] or ""))
    agenda.sort(key=lambda c: (c["due_date"] is None, c["due_date"] or ""))

    return {"name": user.get("name") or "", "tasks": tasks, "agenda": agenda,
            "stars": int(member.get("stars") or 0),
            "week_earned": max(0, int(member.get("week_earned") or 0))}


@app.post("/api/teen/tasks/{card_id}/done")
async def teen_finish_task(card_id: str, teen=Depends(require_teen)):
    """Tick off one of the teen's own tasks — only a card they can see and only
    ever to DONE. Ownership is re-checked here, never trusted from the client."""
    database = get_db()
    user = teen["user"]
    name = (user.get("name") or "").strip().lower()
    card = await database["cards"].find_one(
        {"card_id": card_id, "family_id": user["family_id"]}, {"_id": 0})
    if not card or not _teen_can_see(card, name, user.get("user_id")):
        raise HTTPException(status_code=404, detail="Task not found")
    # A teen marks it done — but the star waits for a parent to approve (unlike
    # a young kid, who auto-earns). Record who finished it so the parent's
    # approval list can show it and credit the right teen.
    # Never reset a task a parent already approved — the $ne guard makes
    # re-ticking idempotent, so it can't be re-approved for extra stars.
    await database["cards"].update_one(
        {"card_id": card_id, "teen_star_status": {"$ne": "approved"}}, {"$set": {
            "status": "DONE", "completed_at": utcnow(),
            "completed_by_user_id": user["user_id"],
            "completed_by_name": user.get("name") or "",
            "teen_star_status": "pending", "updated_at": utcnow()}})
    # Tell the parents a star is waiting on them. Without this a teen finishes a
    # chore and it just sits in an approval list nobody knows to open.
    try:
        who = user.get("name") or "Your teen"
        title = (card.get("title") or "").strip()
        await send_coparent_alert(
            user["family_id"],
            f"{who} finished a task",
            (f"“{title}” — approve their star?" if title else "Approve their star?"),
            "teen_approval",
            created_by_user_id=user["user_id"],
        )
    except Exception as e:
        log.warning("teen approval alert failed: %s", e)
    return {"ok": True}


class TeenApprovalIn(BaseModel):
    approve: bool = True
    stars: int = 1


@app.get("/api/family/teen-approvals")
async def teen_approvals(user=Depends(require_user)):
    """Tasks a teen finished that are waiting for a parent to award the star."""
    database = get_db()
    rows = []
    async for card in database["cards"].find(
        {"family_id": user["family_id"], "teen_star_status": "pending"}, {"_id": 0}):
        rows.append({
            "card_id": card["card_id"],
            "title": card.get("title") or "",
            "teen_name": card.get("completed_by_name") or card.get("assignee") or "",
            "completed_at": iso(card.get("completed_at")),
        })
    rows.sort(key=lambda r: (r["completed_at"] is None, r["completed_at"] or ""), reverse=True)
    return {"approvals": rows}


@app.post("/api/family/teen-approvals/{card_id}")
async def resolve_teen_approval(card_id: str, payload: TeenApprovalIn, user=Depends(require_user)):
    """Approve (award the star to the teen) or dismiss a finished teen task."""
    database = get_db()
    new_status = "approved" if payload.approve else "declined"
    # Claim the pending card atomically: the status flip IS the guard, so two
    # near-simultaneous taps can't both award the star (only the first matches).
    claim = await database["cards"].update_one(
        {"card_id": card_id, "family_id": user["family_id"], "teen_star_status": "pending"},
        {"$set": {"teen_star_status": new_status, "updated_at": utcnow()}})
    if claim.matched_count == 0:
        raise HTTPException(status_code=404, detail="Nothing to approve")

    if payload.approve:
        # Credit the teen who finished it — resolve their member row from the
        # user_id recorded at completion.
        card = await database["cards"].find_one(
            {"card_id": card_id, "family_id": user["family_id"]}, {"_id": 0})
        teen_member = await database["family_members"].find_one(
            {"family_id": user["family_id"], "user_id": (card or {}).get("completed_by_user_id")},
            {"_id": 0})
        stars = max(1, min(int(payload.stars or 1), 20))
        if teen_member:
            await award_stars_to_member(
                database, user["family_id"], teen_member["member_id"], stars,
                (card or {}).get("title") or "Task approved",
                {"family_id": user["family_id"], "user_id": user["user_id"], "name": user.get("name")})
            # Close the loop back to the teen — their side is a real account, so
            # they should hear that the star they were waiting on came through.
            try:
                teen_uid = (card or {}).get("completed_by_user_id")
                if teen_uid:
                    teen_user = await database["users"].find_one({"user_id": teen_uid}, {"_id": 0})
                    L = PUSH_I18N.get((teen_user or {}).get("language") or "en", PUSH_I18N["en"])
                    await send_push_to_user(
                        database, teen_uid,
                        L["teen_star_title"],
                        L["teen_star_body"].format(title=(card or {}).get("title") or ""),
                        {"type": "teen_star", "card_id": card_id})
            except Exception as e:
                log.warning("teen star notify failed: %s", e)
    return {"ok": True, "status": new_status}


# -----------------------------------------------------------------------------
# Family chat — parents talk to each other and to each teen, privacy preserved.
#
# Threads: 'adults' (parents/co-parents only) and one per teen (parents + that
# one teen). A teen only ever reaches their OWN thread — the server forces the
# thread key to their user id and never trusts a client-supplied one — so the
# "teens see nothing of the family" promise holds even with chat on. The
# messages collection stands alone; reading a thread joins no family data.
# -----------------------------------------------------------------------------
MAX_CHAT_LEN = 2000
ADULTS_THREAD = "adults"


class ChatMessageIn(BaseModel):
    text: str


def public_chat_message(m: dict, viewer_id: str) -> dict:
    return {
        "message_id": m["message_id"],
        "thread": m["thread"],
        "sender_kind": m.get("sender_kind"),
        "sender_name": m.get("sender_name") or "",
        "text": m.get("text") or "",
        "created_at": iso(m.get("created_at")),
        "mine": m.get("sender_user_id") == viewer_id,
        "read": (m.get("sender_user_id") == viewer_id) or (viewer_id in (m.get("read_by") or [])),
    }


HOUSEHOLD_THREAD = "household"
DM_PREFIX = "dm:"
KID_PREFIX = "kid:"


# "~" and not "_": ids are minted as user_<hex>, so an underscore separator
# made "dm:user_a_user_b" impossible to split back into two ids.
DM_SEP = "~"


def dm_thread(a: str, b: str) -> str:
    """The id of a one-to-one. Sorted, so the same pair can never end up with
    two threads depending on who opened it first."""
    return DM_PREFIX + DM_SEP.join(sorted([a, b]))


async def _family_accounts(database, family_id: str) -> list:
    """Everyone in the household who can hold a conversation — a member row with
    a login. A young child has none; they are reached through a kid: thread."""
    out = []
    async for m in database["family_members"].find({"family_id": family_id}, {"_id": 0}):
        uid = m.get("user_id")
        if not uid:
            continue
        out.append({"user_id": uid, "member_id": m.get("member_id"),
                    "name": m.get("name") or "", "role": (m.get("role") or "").strip().lower()})
    return out


async def _family_teens(database, family_id: str) -> list:
    """Teens with their own login. Still used by the teen endpoints below."""
    return [a for a in await _family_accounts(database, family_id) if a["role"] == "teen"]


async def _parent_user_ids(database, family_id: str, exclude: Optional[str] = None) -> list:
    return [a["user_id"] for a in await _family_accounts(database, family_id)
            if _is_parent_role(a["role"]) and a["user_id"] != exclude]


async def _thread_participants(database, family_id: str, thread: str):
    """Who is allowed in a thread, or None if it is not a thread for this family.

    This is the whole access model. Every kind of conversation is described by
    the set of people in it, and reading or writing is allowed exactly when the
    caller is one of them — so a new kind of thread is safe the day it is added
    rather than the day someone remembers to guard it.
    """
    accounts = await _family_accounts(database, family_id)
    by_id = {a["user_id"]: a for a in accounts}
    parents = {a["user_id"] for a in accounts if _is_parent_role(a["role"])}

    if thread == HOUSEHOLD_THREAD:
        return set(by_id)                      # everyone with a login
    if thread == ADULTS_THREAD:
        return parents                          # the grown-ups' own room

    if thread.startswith(DM_PREFIX):
        pair = set(thread[len(DM_PREFIX):].split(DM_SEP))
        if len(pair) != 2 or not pair <= set(by_id):
            return None
        # Every private thread has a parent in it. Two teens, or a teen and a
        # helper, would otherwise get a channel inside the family app that no
        # parent can see — a safeguarding surface this app should not open by
        # accident. Siblings still have the household room.
        if not (pair & parents):
            return None
        return pair

    if thread.startswith(KID_PREFIX):
        member_id = thread[len(KID_PREFIX):]
        child = await database["family_members"].find_one(
            {"member_id": member_id, "family_id": family_id}, {"_id": 0})
        # Only a managed child — one with no login of their own. A teen has an
        # account and gets a real one-to-one instead.
        if not child or child.get("user_id"):
            return None
        return parents

    # Threads opened before this model existed are keyed by a teen's user id.
    if thread in by_id:
        return parents | {thread}
    return None


async def _require_thread_member(database, family_id: str, thread: str, user_id: str):
    """404, not 403: whether a conversation exists is itself private."""
    participants = await _thread_participants(database, family_id, thread)
    if participants is None or user_id not in participants:
        raise HTTPException(status_code=404, detail="Conversation not found")
    return participants


async def _chat_insert(database, family_id: str, thread: str, sender_user_id: str,
                       sender_kind: str, sender_name: str, text: str) -> dict:
    clean = sanitize_message_text(text or "", MAX_CHAT_LEN)
    if not clean:
        raise HTTPException(status_code=400, detail="Message can\'t be empty.")
    msg = {
        "message_id": new_id("msg"),
        "family_id": family_id,
        "thread": thread,
        "sender_user_id": sender_user_id,
        "sender_kind": sender_kind,
        "sender_name": sender_name,
        "text": clean,
        "read_by": [sender_user_id],  # the sender has, of course, "read" it
        "created_at": utcnow(),
    }
    await database["messages"].insert_one(msg)
    return msg


async def _chat_thread_messages(database, family_id: str, thread: str, viewer_id: str) -> list:
    rows = []
    async for m in database["messages"].find(
            {"family_id": family_id, "thread": thread}, {"_id": 0}):
        rows.append(m)
    rows.sort(key=lambda r: r.get("created_at") or utcnow())
    return [public_chat_message(m, viewer_id) for m in rows]


async def _mark_read(database, family_id: str, thread: str, viewer_id: str) -> None:
    # Everything in the thread the viewer didn't send is now read by them.
    await database["messages"].update_many(
        {"family_id": family_id, "thread": thread,
         "sender_user_id": {"$ne": viewer_id}, "read_by": {"$ne": viewer_id}},
        {"$addToSet": {"read_by": viewer_id}})


async def _chat_notify(database, family_id: str, thread: str, sender_user_id: str,
                       sender_name: str, text: str) -> None:
    """Best-effort push to everyone else in the thread — which the participant
    set already tells us, whatever kind of thread it is."""
    try:
        participants = await _thread_participants(database, family_id, thread) or set()
        preview = (text or "")[:120]
        for uid in participants - {sender_user_id}:
            await send_push_to_user(
                database, uid, sender_name or "New message", preview,
                {"type": "chat", "thread": thread})
    except Exception:
        pass  # a failed notification must never fail the send


async def require_chat_account(authorization: str = Header(default="")):
    """Any signed-in member of a household — parent, co-parent, teen or helper.

    Chat is the one surface where all of them belong, so role is the wrong gate
    here: WHICH conversation you may open is decided by whether you are in it
    (_thread_participants), not by what you are. Kid-mode sessions are refused —
    a child reads their notes through /api/kid/notes instead.
    """
    database = get_db()
    if not authorization.startswith("Bearer "):
        raise HTTPException(status_code=401, detail="Missing bearer token")
    token = authorization.replace("Bearer ", "", 1).strip()
    session = await database["user_sessions"].find_one(
        {"token_hash": sha256(token), "expires_at": {"$gt": utcnow()}}, {"_id": 0})
    if not session or session.get("kind") == "child":
        raise HTTPException(status_code=401, detail="Invalid or expired session")
    await _touch_session(database, sha256(token), session)
    user = await database["users"].find_one({"user_id": session["user_id"]}, {"_id": 0})
    if not user:
        raise HTTPException(status_code=401, detail="User not found")
    return user


def _sender_kind(user: dict) -> str:
    if user.get("is_teen"):
        return "teen"
    if user.get("is_helper"):
        return "helper"
    return "parent"


@app.get("/api/family/chat/threads")
async def family_chat_threads(user=Depends(require_chat_account)):
    """Every conversation this person is actually in.

    A parent sees the household, the grown-ups' room, a one-to-one with each
    other member, and a note thread for each young child. A teen or helper sees
    the household and their own one-to-one with each parent — the list is built
    from the same participation rule that guards reading, so it can never offer
    a door that then refuses to open.
    """
    database = get_db()
    family_id = user["family_id"]
    viewer = user["user_id"]
    accounts = await _family_accounts(database, family_id)
    is_parent = any(a["user_id"] == viewer and _is_parent_role(a["role"]) for a in accounts)

    threads = [{"thread": HOUSEHOLD_THREAD, "title": None, "is_adults": False, "is_household": True}]
    if is_parent:
        threads.append({"thread": ADULTS_THREAD, "title": None, "is_adults": True, "is_household": False})

    if is_parent:
        for a in accounts:
            if a["user_id"] == viewer:
                continue
            # A teen keeps the thread key they already had — their own user id.
            # Moving them to a dm: key would have left every message sent before
            # today sitting in a thread nothing lists any more, which is losing
            # a family's conversation to a refactor.
            thread_id = a["user_id"] if a["role"] == "teen" else dm_thread(viewer, a["user_id"])
            threads.append({"thread": thread_id, "title": a["name"], "is_adults": False,
                            "is_household": False, "role": a["role"],
                            "member_id": a["member_id"]})
    else:
        # A teen has one conversation with their parents — the same thread the
        # teen screen has always used. A helper gets a one-to-one with each.
        if any(a["user_id"] == viewer and a["role"] == "teen" for a in accounts):
            threads.append({"thread": viewer, "title": None, "is_adults": False,
                            "is_household": False, "is_parents": True})
        else:
            for a in accounts:
                if _is_parent_role(a["role"]):
                    threads.append({"thread": dm_thread(viewer, a["user_id"]), "title": a["name"],
                                    "is_adults": False, "is_household": False, "role": a["role"],
                                    "member_id": a["member_id"]})

    if is_parent:
        async for m in database["family_members"].find(
                {"family_id": family_id, "role": {"$regex": "^child$", "$options": "i"}}, {"_id": 0}):
            if m.get("user_id"):
                continue  # has an account: they get a real one-to-one above
            threads.append({"thread": KID_PREFIX + m["member_id"], "title": m.get("name") or "",
                            "is_adults": False, "is_household": False, "role": "child",
                            "member_id": m["member_id"]})

    out = []
    for th in threads:
        msgs = await _chat_thread_messages(database, family_id, th["thread"], viewer)
        unread = sum(1 for m in msgs if not m["read"] and not m["mine"])
        last = msgs[-1] if msgs else None
        out.append({**th, "unread": unread,
                    "last_text": last["text"] if last else "",
                    "last_at": last["created_at"] if last else None})
    return {"threads": out}


@app.get("/api/family/chat/{thread}")
async def family_chat_get(thread: str, user=Depends(require_chat_account)):
    database = get_db()
    await _require_thread_member(database, user["family_id"], thread, user["user_id"])
    msgs = await _chat_thread_messages(database, user["family_id"], thread, user["user_id"])
    await _mark_read(database, user["family_id"], thread, user["user_id"])
    return {"messages": msgs}


@app.post("/api/family/chat/{thread}")
async def family_chat_send(thread: str, payload: ChatMessageIn, user=Depends(require_chat_account)):
    database = get_db()
    await _require_thread_member(database, user["family_id"], thread, user["user_id"])
    name = user.get("name") or "Someone"
    msg = await _chat_insert(database, user["family_id"], thread, user["user_id"],
                             _sender_kind(user), name, payload.text)
    await _chat_notify(database, user["family_id"], thread, user["user_id"], name, msg["text"])
    return {"ok": True, "message": public_chat_message(msg, user["user_id"])}


@app.post("/api/family/chat/{thread}/read")
async def family_chat_read(thread: str, user=Depends(require_chat_account)):
    database = get_db()
    await _require_thread_member(database, user["family_id"], thread, user["user_id"])
    await _mark_read(database, user["family_id"], thread, user["user_id"])
    return {"ok": True}


@app.get("/api/teen/chat")
async def teen_chat_get(teen=Depends(require_teen)):
    """A teen's own thread, kept for the teen screen. The key is forced to their
    user id, so it can only ever be theirs."""
    database = get_db()
    tuid = teen["user"]["user_id"]
    msgs = await _chat_thread_messages(database, teen["family_id"], tuid, tuid)
    await _mark_read(database, teen["family_id"], tuid, tuid)
    return {"messages": msgs}


@app.post("/api/teen/chat")
async def teen_chat_send(payload: ChatMessageIn, teen=Depends(require_teen)):
    database = get_db()
    tuid = teen["user"]["user_id"]
    name = teen["user"].get("name") or "Teen"
    msg = await _chat_insert(database, teen["family_id"], tuid, tuid, "teen", name, payload.text)
    await _chat_notify(database, teen["family_id"], tuid, tuid, name, msg["text"])
    return {"ok": True, "message": public_chat_message(msg, tuid)}


@app.post("/api/teen/chat/read")
async def teen_chat_read(teen=Depends(require_teen)):
    database = get_db()
    tuid = teen["user"]["user_id"]
    await _mark_read(database, teen["family_id"], tuid, tuid)
    return {"ok": True}


@app.get("/api/kid/notes")
async def kid_notes(child=Depends(require_child)):
    """What a parent has written to this child, read in kid mode.

    A young child has no login and no inbox, so this is the only honest way to
    say something to them in the app.

    Read-only, deliberately. A teen has their own account and a real two-way
    conversation; an under-13 does not, and giving one a way to send messages
    from a shared family phone is a line this app should not cross for the sake
    of a "seen" confirmation.
    """
    database = get_db()
    member = child["member"]
    thread = KID_PREFIX + member["member_id"]
    # The child is not a user, so "read" is tracked against their member id.
    viewer = f"member:{member['member_id']}"
    msgs = await _chat_thread_messages(database, child["family_id"], thread, viewer)
    await _mark_read(database, child["family_id"], thread, viewer)
    return {"messages": msgs}


@app.post("/api/kid/chores/{card_id}/done")
async def kid_finish_chore(card_id: str, child=Depends(require_child)):
    """Tick off a chore. Only one with this child's name on it, and only ever
    to DONE — a child can finish their own jobs, not reopen or edit them."""
    database = get_db()
    member = child["member"]
    card = await database["cards"].find_one(
        {"card_id": card_id, "family_id": child["family_id"]}, {"_id": 0})
    if not card or not card.get("shared"):
        raise HTTPException(status_code=404, detail="Not found")
    if (card.get("assignee") or "").strip().lower() != (member.get("name") or "").strip().lower():
        raise HTTPException(status_code=403, detail="That is not yours to finish")

    # The guard rides in the filter, not in a variable read moments earlier:
    # a double tap had both requests see an unfinished chore and both pay for
    # it, leaving two ledger rows for one job and a weekly meter inflated
    # enough to unlock a treat that was not earned.
    claim = await database["cards"].update_one(
        {"card_id": card_id, "family_id": child["family_id"],
         "stars_awarded": {"$ne": True}, "status": {"$ne": "DONE"}},
        {"$set": {
            "status": "DONE", "completed_at": utcnow(),
            "completed_by_name": member.get("name") or "", "completed_by_user_id": None,
            "stars_awarded": True}})
    already = claim.matched_count == 0
    await log_activity(database, {"family_id": child["family_id"], "user_id": None,
                                  "name": member.get("name") or ""},
                       "task_done", card.get("title", ""))
    # Finishing your own chore earns the same 5 stars the parent-marked path
    # grants, through the shared helper so it ticks the weekly meter and writes
    # the ledger. Without this a child in kid mode saw "done" but earned
    # nothing, and the card was already DONE so no parent award could follow.
    # Guarded so a re-tap on an already-finished chore can't pay twice.
    if not already:
        old_stars = int(member.get("stars", 0))
        actor = {"family_id": child["family_id"], "user_id": None, "name": member.get("name") or ""}
        await award_stars_to_member(database, child["family_id"], member["member_id"], 5,
                                    card.get("title") or "Chore done", actor)
        await send_star_milestone_alert(child["family_id"], member.get("name", "Your child"),
                                        old_stars, old_stars + 5)
    return {"ok": True}


@app.post("/api/kid/rewards/{reward_id}/request")
async def kid_request_reward(reward_id: str, child=Depends(require_child)):
    """Spend stars on a reward.

    The same atomic guarded decrement the parent-side path uses: checking the
    balance and then subtracting in two steps lets two taps both pass, so the
    filter carries the check and the database rejects the second.
    """
    database = get_db()
    member = child["member"]
    reward = await database["rewards"].find_one(
        {"reward_id": reward_id, "family_id": child["family_id"]}, {"_id": 0})
    if not reward:
        raise HTTPException(status_code=404, detail="Reward not found")
    cost = max(1, int(reward.get("cost_stars") or 0))

    # A weekend treat has to be earned THIS week — the rule the whole weekly
    # meter exists to express. It was enforced on the parent's redeem endpoint
    # and not here, on the screen children actually spend from: old savings
    # bought weekend treats, and because the meter was never debited, one
    # week's earnings could buy them without limit.
    member = await roll_week_if_stale(database, member)
    star_filter = {"member_id": member["member_id"], "family_id": child["family_id"],
                   "stars": {"$gte": cost}}
    inc = {"stars": -cost}
    if reward.get("weekend"):
        star_filter["week_earned"] = {"$gte": cost}
        inc["week_earned"] = -cost

    result = await database["family_members"].update_one(star_filter, {"$inc": inc})
    if result.matched_count == 0:
        raise HTTPException(
            status_code=400,
            detail=("Not enough stars earned this week for a weekend treat."
                    if reward.get("weekend") else "Not enough stars"))

    await database["star_transactions"].insert_one({
        "transaction_id": new_id("star"), "family_id": child["family_id"],
        "member_id": member["member_id"], "delta": -cost,
        "reason": reward.get("title") or "Reward redeemed",
        "created_by_user_id": None, "created_by_name": member.get("name"),
        "created_at": utcnow()})
    redemption = {
        "redemption_id": new_id("redemption"), "family_id": child["family_id"],
        "member_id": member["member_id"], "reward_id": reward["reward_id"],
        "reward_title": reward.get("title") or "", "reward_icon": reward.get("icon"),
        "cost_stars": cost, "status": "pending", "created_at": utcnow(),
        "weekend": bool(reward.get("weekend")),
        "weekend": bool(reward.get("weekend")),
        "created_by_user_id": None, "fulfilled_at": None}
    await database["redemptions"].insert_one(redemption)
    # A child spent their stars — both parents should know, so the treat gets
    # handed over and nobody double-pays. created_by is None (the child has no
    # login), so this reaches every adult in the household.
    try:
        await send_coparent_alert(
            child["family_id"],
            f"{member.get('name') or 'Your child'} redeemed a reward",
            reward.get("title") or "", "reward_redeemed", created_by_user_id=None)
    except Exception as e:
        log.warning("kid reward alert failed: %s", e)

    fresh = await database["family_members"].find_one(
        {"member_id": member["member_id"]}, {"_id": 0})
    return {"ok": True, "stars": int(fresh.get("stars") or 0),
            "redemption": public_redemption(redemption)}


@app.get("/api/search")
async def search_everything(q: str = Query(default=""), user=Depends(require_user)):
    """One box for "where did I put the school form".

    A household's knowledge is scattered across five screens by the time it is
    two weeks old — a task here, a photographed letter in the vault, a note
    about the plumber, a meal plan with the recipe in it. Remembering WHICH
    screen is the app's problem, not the parent's.

    Matching is done in Python rather than with a database regex on purpose:
    the visibility rules are not expressible as one query (private cards, the
    vault's owner/visibility/legacy triple), and re-implementing them in a
    query language is exactly how a search box leaks a co-parent's private
    documents. Every result goes through the same predicate the owning screen
    uses. A family's data is small enough that this costs nothing.
    """
    needle = (q or "").strip().lower()
    if len(needle) < 2:
        return {"query": q, "results": [], "truncated": False}

    database = get_db()
    hits: List[dict] = []

    def matches(*fields) -> bool:
        return any(needle in str(f or "").lower() for f in fields)

    # Tasks, events and notes — same visibility rule as the feed.
    cursor = database["cards"].find({
        "family_id": user["family_id"],
        "$or": [
            {"shared": True},
            {"created_by_user_id": user["user_id"]},
            {"created_by_user_id": {"$exists": False}},
        ],
    }, {"_id": 0}).sort("created_at", -1)
    async for row in cursor:
        if matches(row.get("title"), row.get("description"), row.get("assignee")):
            hits.append(_search_hit(
                (row.get("type") or "TASK").lower(),
                row["card_id"],
                row.get("title") or "",
                row.get("description") or row.get("assignee") or "",
                iso(row.get("due_date")) or "",
                row.get("status") or "",
            ))

    # Documents — never widen what the vault itself would show.
    async for row in database["vault"].find({"family_id": user["family_id"]}, {"_id": 0}).sort("created_at", -1):
        if not _may_see_vault_doc(row, user):
            continue
        if matches(row.get("title"), row.get("category"), row.get("file_name")):
            hits.append(_search_hit(
                "document", row["doc_id"], row.get("title") or "",
                row.get("category") or "", iso(row.get("created_at")) or "",
                row.get("visibility") or "shared",
            ))

    async for row in database["shopping_list"].find({"family_id": user["family_id"]}, {"_id": 0}):
        if matches(row.get("name"), row.get("category")):
            hits.append(_search_hit(
                "shopping", row["item_id"], row.get("name") or "",
                row.get("category") or "", "",
                "done" if row.get("checked") else "open",
            ))

    async for row in database["meals"].find({"family_id": user["family_id"]}, {"_id": 0}):
        if matches(row.get("title"), row.get("notes"), " ".join(row.get("ingredients") or [])):
            hits.append(_search_hit(
                "meal", row["meal_id"], row.get("title") or "",
                row.get("notes") or "", row.get("day") or "",
                row.get("meal_type") or "",
            ))

    # A title match is what someone typing two words is almost always after;
    # a body match is the fallback. Within each, keep the order the screens
    # use (newest first).
    hits.sort(key=lambda h: 0 if needle in h["title"].lower() else 1)
    return {
        "query": q,
        "results": hits[:SEARCH_LIMIT],
        "truncated": len(hits) > SEARCH_LIMIT,
    }


@app.get("/api/family/updates")
async def family_updates(
    x_confirm: Optional[str] = Header(None, alias="X-Confirm"),
    user=Depends(require_user),
):
    """Blocklist-neutral twin of the invite discovery + acceptance.

    The device-errors telemetry finally named the enemy: one family iPhone
    blocks EVERY URL containing the words "invite" or "membership" — GET and
    POST alike — while /telemetry, /cards and /family/members sail through.
    Keyword filter lists, not verbs or networks. This URL contains none of
    the trigger words; without X-Confirm it lists pending invites for the
    signed-in email, with X-Confirm: <token> it performs the acceptance.
    """
    token = x_confirm if isinstance(x_confirm, str) else None
    if token:
        return await _accept_invite_request(token, user, via="discovery")
    return await invites_for_me(user=user)


@app.delete("/api/family/invites/{invite_id}")
async def delete_family_invite(invite_id: str, user=Depends(require_full_member)):
    """Revoke a pending (or any) invite belonging to the caller's family.
    Deleting the record invalidates its token, so a shared link stops working."""
    database = get_db()
    result = await database["family_invites"].delete_one(
        {"invite_id": invite_id, "family_id": user["family_id"]}
    )
    if result.deleted_count == 0:
        raise HTTPException(status_code=404, detail="Invite not found")
    return {"ok": True}


@app.post("/api/family/invites/{invite_id}/complete")
async def complete_family_invite(invite_id: str, user=Depends(require_user)):
    """The inviter finishes the join from THEIR device.

    Field saga: the invitee's iPhone killed every write-shaped request no
    matter the verb, path or header, across Wi-Fi and 5G — her device could
    read the invitation but never deliver the acceptance. The join is a
    server-side operation; nothing about it requires the invitee's network
    to work. The inviter taps "Add now" on the pending invite, the server
    runs the exact same acceptance path (family switch, children migration,
    notifications), and the invitee's device only has to read the result.
    Guarded: the invite must belong to the caller's family, be pending, be
    addressed to an email, and that email must already have an account.
    """
    database = get_db()
    invite = await database["family_invites"].find_one(
        {"invite_id": invite_id, "family_id": user["family_id"]}, {"_id": 0}
    )
    if not invite:
        raise HTTPException(status_code=404, detail="Invite not found")
    if invite.get("status") == "accepted":
        raise HTTPException(status_code=409, detail="Invite has already been accepted")
    email = (invite.get("email") or "").strip().lower()
    if not email:
        raise HTTPException(status_code=400, detail="This invite has no email address")
    invitee = await database["users"].find_one({"email": email}, {"_id": 0})
    if not invitee:
        raise HTTPException(
            status_code=404,
            detail="No account with this email yet — they need to sign up first.",
        )
    if invitee.get("family_id") == user["family_id"]:
        raise HTTPException(status_code=409, detail="They are already in your household")
    fresh, joined = await _accept_invite_for_user(
        database, invitee, invite, user["family_id"]
    )
    # Tell the invitee in-app, best effort — their device reads fine.
    try:
        lang = fresh.get("language") or "en"
        L = PUSH_I18N.get(lang, PUSH_I18N["en"])
        await asyncio.wait_for(send_push_to_user(
            database, fresh["user_id"],
            L["invited_title"].format(name=user.get("name") or "A parent"),
            L["invited_body"],
            {"type": "family_joined"},
        ), timeout=5.0)
    except Exception as exc:
        log.warning("invite completion push skipped: %s", exc)
    return {"ok": True, "joined": joined, "member_email": email}


@app.get("/api/family/invite/{token}")
async def family_invite_lookup(token: str):
    database = get_db()
    invite = await database["family_invites"].find_one({"token": token}, {"_id": 0})
    if not invite:
        raise HTTPException(status_code=404, detail="Invite not found")

    if invite.get("expires_at") and _expired(invite["expires_at"]):
        raise HTTPException(status_code=410, detail="Invite has expired")

    inviter = await database["users"].find_one(
        {"user_id": invite.get("created_by_user_id")},
        {"_id": 0},
    )

    return {
        "invite_id": invite["invite_id"],
        "status": invite.get("status", "pending"),
        "email": invite.get("email"),
        "inviter_name": (inviter or {}).get("name") or invite.get("created_by_name") or "A family member",
        "relationship": invite.get("relationship"),
        "expires_at": iso(invite.get("expires_at")),
    }


class InviteAcceptIn(BaseModel):
    token: str


# Five doors onto one room, and none of them is redundant.
#
# _accept_invite_request is the whole feature; everything below is a different
# URL and verb pointing at it, because a real family iPhone with a content
# blocker refused specific SHAPES of request rather than failing outright.
# Which door a given device can walk through is not knowable in advance.
#
# The order the CLIENT tries them is where the thinking now lives (see
# InviteJoinPrompt): the discovery URL first, because the join card is only on
# screen if that exact request just succeeded.
#
# Before deleting any of these: they are reachable by app builds already on
# people's phones, which update on their own schedule and not ours. Before
# ADDING a sixth: check the client's ordered list first — the answer is
# almost certainly ordering, not another door.
@app.post("/api/family/invite/accept")
async def family_invite_accept(payload: InviteAcceptIn, user=Depends(require_user)):
    """Accept an invite while already signed in.

    Invite links open the app for logged-in users too, and those users never
    pass through the sign-in screen where tokens are otherwise consumed.
    """
    return await _accept_invite_request(payload.token, user, via="invite-accept-post")


@app.get("/api/family/invite-accept")
async def family_invite_accept_get(token: str, user=Depends(require_user)):
    """GET twin of the accept endpoint, used as an automatic client fallback.

    Seen in the field: an iPhone on home Wi-Fi where every GET to this API
    succeeded while this one POST died in Safari's network layer ("Load
    failed"). Whatever drops those POSTs, GETs go through — so the app
    retries acceptance over the request shape that demonstrably works.
    Distinct path on purpose: GET /family/invite/accept would be captured
    by the /family/invite/{token} lookup route.
    """
    return await _accept_invite_request(token, user, via="invite-accept-get")


@app.post("/api/family/membership")
async def family_membership_post(payload: InviteAcceptIn, user=Depends(require_user)):
    """Same as the accept endpoint under a blocklist-proof name.

    Field case: an iPhone content blocker killed both verbs of the accept
    endpoint ("Load failed" on Wi-Fi AND 5G) while every other API call
    passed — ad-block filter lists target URLs containing words like
    "accept". "membership" appears on no list. Old routes stay for old
    clients.
    """
    return await _accept_invite_request(payload.token, user, via="membership-post")


@app.get("/api/family/membership")
async def family_membership_get(token: str, user=Depends(require_user)):
    return await _accept_invite_request(token, user, via="membership-get")


async def _accept_invite_request(token: str, user: dict, via: str = "unknown"):
    """`via` names the door this acceptance came through.

    Five routes exist because a blocked iPhone refused specific request
    SHAPES, and none can be deleted on a hunch: they are reachable from app
    builds already on people's phones. Counting which ones real devices
    actually succeed on turns that hunch into evidence — so a later session
    can retire the dead ones knowing whose phone it is about to break, or
    knowing it breaks nobody.
    """
    database = get_db()
    invite, target_family_id = await _resolve_invite(
        database, token, user.get("email", "")
    )
    if not invite:
        raise HTTPException(status_code=404, detail="Invite not found")
    try:
        fresh, joined = await _accept_invite_for_user(database, user, invite, target_family_id)
        # Counted on success, including the idempotent repeat where the user
        # was already a member: the question is which request SHAPES reach
        # the server on real devices, and a 200 answers that either way.
        # Deliberately not counted on failure — that would measure which
        # doors people knock on, not which ones open.
        try:
            await database["invite_route_stats"].update_one(
                {"route": via}, {"$inc": {"count": 1}}, upsert=True)
        except Exception as exc:  # noqa: BLE001 — a counter must never fail a join
            log.warning("invite route not counted (%s): %s", via, exc)
        # Serialising the result is part of accepting it. This sat outside the
        # try, so a failure here produced Starlette's bare "Internal Server
        # Error" instead of the readable message this handler exists to give —
        # and that difference is diagnostic: production returning the generic
        # text proves the failure was NOT in the join itself.
        result = {"ok": True, "joined": joined, "user": public_user(fresh)}
    except HTTPException:
        raise
    except Exception as exc:
        # A plain-text 500 reaches the join card as an unreadable mystery;
        # a JSON detail is shown to the user verbatim and names the problem.
        # The exception TYPE goes in the detail — not the message, which can
        # carry record contents — so a failure is diagnosable from the smoke
        # test's output alone, without shell access to the running container.
        log.exception("invite accept failed")
        raise HTTPException(
            status_code=500,
            detail=f"Join failed. Please try again. ({type(exc).__name__})",
        )
    return result


# -----------------------------------------------------------------------------
# Notifications
# -----------------------------------------------------------------------------
@app.get("/api/notifications/settings")
async def get_notification_settings(user=Depends(require_user)):
    settings = await get_notification_settings_doc(user["user_id"])
    return public_notification_settings(settings)


@app.put("/api/notifications/settings")
async def update_notification_settings(payload: NotificationPrefsIn, user=Depends(require_user)):
    database = get_db()
    current = await get_notification_settings_doc(user["user_id"])

    changes = {"updated_at": utcnow()}
    if payload.card_reminders is not None:
        changes["card_reminders"] = bool(payload.card_reminders)
    if payload.new_card_alerts is not None:
        changes["new_card_alerts"] = bool(payload.new_card_alerts)

    await database["notification_settings"].update_one(
        {"user_id": user["user_id"]},
        {"$set": changes},
        upsert=True,
    )

    current.update(changes)
    return public_notification_settings(current)


@app.post("/api/notifications/register")
async def register_notification_token(payload: NotificationTokenIn, user=Depends(require_user)):
    database = get_db()

    token = payload.token.strip()
    if not token:
        raise HTTPException(status_code=400, detail="Notification token is required")

    doc = {
        "token": token,
        "user_id": user["user_id"],
        "family_id": user["family_id"],
        "email": user.get("email"),
        "platform": payload.platform,
        # Which build/runtime this device is running — refreshed every time the
        # token is, so adoption reporting tracks the real fleet, not history.
        "app_version": (payload.app_version or "").strip()[:20] or None,
        "runtime_version": (payload.runtime_version or "").strip()[:20] or None,
        "active": True,
        "updated_at": utcnow(),
    }

    await database["notification_tokens"].update_one(
        {"token": token},
        {
            "$set": doc,
            "$setOnInsert": {"created_at": utcnow()},
        },
        upsert=True,
    )

    return {"ok": True}


@app.post("/api/notifications/test")
async def test_notification(user=Depends(require_user)):
    database = get_db()
    messages = []
    docs = [d async for d in database["notification_tokens"].find(
        {"user_id": user["user_id"], "active": True}, {"_id": 0})]

    for token_doc in _latest_token_per_user(docs):
        token = token_doc.get("token")
        if token and token.startswith("ExponentPushToken"):
            messages.append(
                {
                    "to": token,
                    "sound": "default",
                    "title": "Ahenora notifications are active",
                    "body": "You will receive card alerts and reminder notifications.",
                    "data": {"type": "notification_test"},
                }
            )

    result = await send_expo_push_messages(messages, database)
    return {"ok": True, "tokens": len(messages), "result": result}



# -----------------------------------------------------------------------------
# AI assign
# -----------------------------------------------------------------------------
@app.post("/api/ai/assign")
async def ai_assign(payload: AiAssignIn, user=Depends(require_user)):
    database = get_db()
    members = []
    async for m in database["family_members"].find({"family_id": user["family_id"]}, {"_id": 0}):
        members.append(m)

    if not members:
        return {"assignee": ""}

    names = [m["name"] for m in members]

    def local_pick():
        parent = next((m for m in members if m["role"].lower() == "parent"), members[0])
        return {"assignee": parent["name"]}

    # Assign is fired automatically as you type a card, not asked for the way a
    # document scan is. So when a family is out of AI scans it quietly falls
    # back to the local pick rather than raising the upgrade wall or burning a
    # document-scan slot on a helper. This also closes the only other unmetered
    # paid-AI route: an over-quota family can no longer drive Gemini calls here.
    if not GOOGLE_API_KEY or await ai_scans_remaining(user) <= 0:
        return local_pick()

    prompt = f"""
Choose the best assignee from this list only: {", ".join(names)}.
Task title: {payload.title}
Task description: {payload.description}
Return only one exact name from the list, or return an empty string.
""".strip()

    try:
        result = await _gemini_text(
            prompt,
            system="You are assigning family tasks. Return only one exact name or empty string.",
        )
    except Exception as exc:  # noqa: BLE001 — degrade like the no-key path, never 500
        log.warning("ai assign failed: %s", exc)
        return local_pick()
    result = result.strip().replace('"', "")
    if result not in names:
        result = ""
    return {"assignee": result}


# -----------------------------------------------------------------------------
# Cards
# -----------------------------------------------------------------------------
@app.get("/api/cards")
async def list_cards(status: Optional[str] = Query(default=None), user=Depends(require_user)):
    database = get_db()
    # Per-item privacy: each parent sees family-shared items, their own private
    # items, and legacy items (created before privacy existed, no owner stored).
    # A co-parent's private items stay hidden until they choose to share.
    query = {
        "family_id": user["family_id"],
        "$or": [
            {"shared": True},
            {"created_by_user_id": user["user_id"]},
            {"created_by_user_id": {"$exists": False}},
        ],
    }
    if status:
        query["status"] = status

    rows = []
    cursor = database["cards"].find(query, {"_id": 0}).sort("created_at", -1)
    async for item in cursor:
        rows.append(public_card(item))
    return rows


@app.get("/api/cards/shared")
async def list_shared_with_coparent(direction: str = "out", user=Depends(require_user)):
    """The two sides of a sharing agreement.

    `out` (the default, and what the old single-direction view showed): what
    the requester has shared — exactly what their co-parent can see from them.
    `in`: what the OTHER adults in the family have shared with the requester —
    so the transparency runs both ways and reads as a mutual arrangement rather
    than a one-way disclosure. Private items never appear in either direction.

    Legacy cards with no recorded owner were visible to the whole family; they
    belong to nobody in particular, so they are deliberately left out of both
    lists — "shared by you" and "shared by them" are both claims about a person,
    and there is no person to name.
    """
    database = get_db()
    if direction == "in":
        query = {
            "family_id": user["family_id"],
            "shared": True,
            "created_by_user_id": {"$nin": [user["user_id"], None]},
        }
    else:
        query = {
            "family_id": user["family_id"],
            "created_by_user_id": user["user_id"],
            "shared": True,
        }

    # Resolve owner user_id -> member name once, so the "in" view can say who
    # shared each item without the client holding a user_id lookup.
    names: dict = {}
    async for m in database["family_members"].find(
        {"family_id": user["family_id"], "user_id": {"$exists": True}},
        {"_id": 0, "user_id": 1, "name": 1},
    ):
        if m.get("user_id"):
            names[m["user_id"]] = m.get("name", "")

    rows = []
    cursor = database["cards"].find(query, {"_id": 0}).sort("due_date", 1)
    async for item in cursor:
        card = public_card(item)
        card["shared_by_name"] = names.get(item.get("created_by_user_id"), "")
        rows.append(card)
    return rows


# The oldest JavaScript bundle we are still willing to serve updates to, and
# the store version that supersedes it. Kept here rather than in the app so it
# can be changed without shipping a release — the whole point is to speak to
# clients that are already out of date.
MIN_SUPPORTED_RUNTIME = "2.0.0"
CURRENT_STORE_VERSION = "1.1.0"


@app.get("/api/app/version-info")
async def app_version_info():
    """What the app should compare itself against.

    An install whose runtime is below the minimum cannot receive over-the-air
    updates at all — a new native shell is the only way it moves — so it needs
    to be told to go to the store rather than left waiting for an update that
    can never arrive.

    Deliberately unauthenticated: a client that is too old to be updated is
    exactly the one that might also be too old to sign in cleanly, and this
    answer is not private.
    """
    # The running backend commit, so "is the fix deployed?" is answerable from
    # outside. Railway sets RAILWAY_GIT_COMMIT_SHA on every deploy; other hosts
    # may set GIT_COMMIT. Empty when neither is present (e.g. local dev).
    commit = (os.environ.get("RAILWAY_GIT_COMMIT_SHA")
              or os.environ.get("GIT_COMMIT") or "")[:12]
    return {
        "min_runtime": MIN_SUPPORTED_RUNTIME,
        "store_version": CURRENT_STORE_VERSION,
        "android_store_url": "https://play.google.com/store/apps/details?id=com.householdcoo.app",
        "backend_commit": commit,
    }


@app.get("/api/admin/dedupe-accounts")
async def admin_dedupe_preview(user=Depends(require_user)):
    """Dry-run report of the duplicate-account cleanup (admin only). Open this in
    a browser while signed in as an admin to see exactly what a merge would do —
    nothing is changed."""
    if not is_admin_user(user):
        raise HTTPException(status_code=403, detail="Admins only")
    lines: list = []
    summary = await dedupe_run(get_db(), apply=False,
                               log=lambda *a: lines.append(" ".join(str(x) for x in a)))
    return {"applied": False, "summary": summary, "report": lines}


@app.post("/api/admin/dedupe-accounts")
async def admin_dedupe_apply(payload: dict = Body(...), user=Depends(require_user)):
    """Apply the duplicate-account cleanup (admin only). Requires an explicit
    {"confirm": "APPLY"} body so it can never run by accident."""
    if not is_admin_user(user):
        raise HTTPException(status_code=403, detail="Admins only")
    if (payload or {}).get("confirm") != "APPLY":
        raise HTTPException(status_code=400, detail='Send {"confirm": "APPLY"} to apply the merge.')
    lines: list = []
    summary = await dedupe_run(get_db(), apply=True,
                               log=lambda *a: lines.append(" ".join(str(x) for x in a)))
    return {"applied": True, "summary": summary, "report": lines}


@app.post("/api/admin/run-due-reminders")
async def admin_run_due_reminders(user=Depends(require_user)):
    """Run one reminder scan now (admin only). The scheduler already does this
    every minute; this is a manual trigger for verifying it end to end and a
    fallback if the in-process loop is ever turned off."""
    if not is_admin_user(user):
        raise HTTPException(status_code=403, detail="Admins only")
    sent = await send_due_card_reminders(get_db())
    return {"sent": sent}


@app.get("/api/admin/version-adoption")
async def version_adoption(user=Depends(require_user)):
    """Who is on which build — the answer to "did the OTA reach everyone".

    An over-the-air update only lands on a device whose runtime matches what was
    published, so `by_runtime` is the real coverage number: the share on the
    current runtime is the share that can receive this and every future OTA;
    everyone else is frozen until they update from the store. Counted over
    active push tokens, one device each, deduped to distinct users so a person
    with two devices is not double-weighted.

    Admin-only: it reads across the whole install base, not one household.
    """
    if not is_admin_user(user):
        raise HTTPException(status_code=403, detail="Admins only")
    database = get_db()

    by_runtime: dict = {}
    by_version: dict = {}
    seen_users: set = set()
    seen_devices = 0
    reporting_devices = 0
    async for tok in database["notification_tokens"].find(
        {"active": True}, {"_id": 0, "user_id": 1, "app_version": 1, "runtime_version": 1}
    ):
        seen_devices += 1
        rt = tok.get("runtime_version")
        av = tok.get("app_version")
        if rt or av:
            reporting_devices += 1
        # Distinct users per bucket, so multi-device people count once. A token
        # with no user_id contributes no user to count — adding None would
        # inflate a bucket past the true user total and push the percentage
        # over 100. The bucket key is still created so the runtime/build shows
        # up (with a zero count) even if only anonymous devices report it.
        rt_key = rt or "unknown"
        av_key = av or "unknown"
        by_runtime.setdefault(rt_key, set())
        by_version.setdefault(av_key, set())
        uid = tok.get("user_id")
        if uid:
            by_runtime[rt_key].add(uid)
            by_version[av_key].add(uid)
            seen_users.add(uid)

    def _counts(buckets: dict) -> dict:
        return {k: len(v) for k, v in sorted(
            buckets.items(), key=lambda kv: (-len(kv[1]), kv[0]))}

    runtime_counts = _counts(by_runtime)
    on_current = runtime_counts.get(MIN_SUPPORTED_RUNTIME, 0)
    total_users = len(seen_users)
    return {
        "current_runtime": MIN_SUPPORTED_RUNTIME,
        "store_version": CURRENT_STORE_VERSION,
        "users_on_current_runtime": on_current,
        "total_users_with_a_device": total_users,
        "pct_on_current_runtime": round(100 * on_current / total_users, 1) if total_users else 0.0,
        "by_runtime": runtime_counts,
        "by_app_version": _counts(by_version),
        "devices_seen": seen_devices,
        "devices_reporting_version": reporting_devices,
    }


@app.get("/api/admin/plan-adoption")
async def plan_adoption(user=Depends(require_user)):
    """Who is actually paying — the answer to "lots of visits, no subscribers".

    Reads the whole install base and separates three populations that the app's
    own screens blur together:

      * paying — a household whose STORED plan is a paid tier. The only way to
        reach it is the RevenueCat webhook (a real purchase); new families seed
        as free, and self-serve upgrade is locked while billing is off. So this
        count is the true subscriber number.
      * free-premium — households getting Premium *limits* without paying, via
        the testing window (global, when RC_WEBHOOK_SECRET is unset → EVERY
        family) or a tester/admin household. These never see a paywall, so they
        can never convert; if this is most of the base, that is the finding.
      * free — genuinely on the free tier and gated normally.

    `billing_live` is the master switch: when False no gate fires anywhere, so a
    zero-subscriber number is expected, not a funnel problem. Counted per family,
    and separately per family that has at least one active device (a real, live
    household) so abandoned sign-ups don't dilute the conversion rate.

    Admin-only: it reads across every household, not one.
    """
    if not is_admin_user(user):
        raise HTTPException(status_code=403, detail="Admins only")
    database = get_db()

    billing_live = bool(os.environ.get("RC_WEBHOOK_SECRET"))

    # Which families own an active device — the population that actually opened
    # the app, so the conversion rate is measured against real households.
    families_with_device: set = set()
    async for tok in database["notification_tokens"].find(
        {"active": True}, {"_id": 0, "family_id": 1}
    ):
        fid = tok.get("family_id")
        if fid:
            families_with_device.add(fid)

    by_stored_plan: dict = {}
    paying = 0
    tester_households = 0
    total_families = 0
    active_total = 0
    active_paying = 0
    async for fam in database["families"].find(
        {}, {"_id": 0, "family_id": 1, "plan": 1}
    ):
        total_families += 1
        plan = fam.get("plan") or "village"
        by_stored_plan[plan] = by_stored_plan.get(plan, 0) + 1
        is_paying = plan != "village"
        if is_paying:
            paying += 1
        is_tester = await family_has_admin(database, fam.get("family_id", ""))
        if is_tester:
            tester_households += 1
        if fam.get("family_id") in families_with_device:
            active_total += 1
            if is_paying:
                active_paying += 1

    # Free-premium: households handed Premium limits for free. The testing window
    # is global, so when it is on every non-paying family qualifies; otherwise
    # only tester households do.
    free_premium = (total_families - paying) if not billing_live else tester_households
    active_free_premium = (active_total - active_paying) if not billing_live else min(
        tester_households, active_total)

    return {
        "billing_live": billing_live,
        "total_families": total_families,
        "by_stored_plan": dict(sorted(by_stored_plan.items(), key=lambda kv: -kv[1])),
        "paying_families": paying,
        "tester_households": tester_households,
        "free_premium_families": free_premium,
        "active_families_with_device": active_total,
        "active_paying_families": active_paying,
        "pct_active_paying": round(100 * active_paying / active_total, 1) if active_total else 0.0,
        "active_free_premium_families": active_free_premium,
    }


@app.get("/api/cards/sharing-summary")
async def sharing_summary(user=Depends(require_user)):
    """The three numbers the privacy panel states, counted in one place.

    The client used to derive these from the agenda it already had, which is a
    different population: /api/cards returns the whole family's SHARED items
    (so a co-parent's items inflated "yours") and only dated, open ones (so
    undated private items went uncounted). The banner and the panel therefore
    disagreed one tap apart — on the one screen whose entire job is to be
    believed. Counted here, against the same queries the lists use, they agree
    by construction.

    Legacy ownerless cards are excluded from all three for the same reason the
    lists exclude them: each number is a claim about a person.
    """
    database = get_db()
    fam = user["family_id"]
    me = user["user_id"]
    shared_out = await database["cards"].count_documents(
        {"family_id": fam, "created_by_user_id": me, "shared": True})
    shared_in = await database["cards"].count_documents(
        {"family_id": fam, "shared": True, "created_by_user_id": {"$nin": [me, None]}})
    private = await database["cards"].count_documents(
        {"family_id": fam, "created_by_user_id": me, "shared": {"$ne": True}})
    return {"shared_out": shared_out, "shared_in": shared_in, "private": private}


@app.get("/api/cards/mine")
async def list_assigned_to_me(user=Depends(require_user)):
    """What is actually on my plate, soonest first.

    Assignment used to be a word in a field that nobody read. This is the
    other half of a hand-off: a place to look that answers "what did they
    give me?" without scrolling the whole household's feed.

    Matched on name because that is what the assignee field holds — and it is
    kept in step by the member-rename path, which rewrites every card.
    """
    database = get_db()
    my_name = (user.get("name") or "").strip().lower()
    if not my_name:
        return []
    rows = []
    cursor = database["cards"].find({
        "family_id": user["family_id"],
        "status": {"$ne": "DONE"},
        "$or": [
            {"shared": True},
            {"created_by_user_id": user["user_id"]},
            {"created_by_user_id": {"$exists": False}},
        ],
    }, {"_id": 0})
    async for item in cursor:
        if (item.get("assignee") or "").strip().lower() == my_name:
            rows.append(public_card(item))
    # Anything with a date first, in date order; undated work after it.
    rows.sort(key=lambda c: (c["due_date"] is None, c["due_date"] or ""))
    return rows


@app.post("/api/cards")
async def create_card(payload: CardIn, user=Depends(require_user)):
    database = get_db()
    doc = {
        "card_id": new_id("card"),
        "family_id": user["family_id"],
        "type": payload.type,
        "title": payload.title,
        "description": payload.description,
        "assignee": payload.assignee,
        "due_date": parse_dt(payload.due_date),
        "status": "OPEN",
        "source": payload.source,
        "image_base64": payload.image_base64,
        "recurrence": payload.recurrence,
        "reminder_minutes": payload.reminder_minutes,
        "created_at": utcnow(),
        "completed_at": None,
        # Private to the creator by default — the co-parent only sees it if the
        # creator explicitly shares it (via /api/cards/{id}/share).
        "created_by_user_id": user["user_id"],
        "shared": bool(payload.shared),
    }
    await database["cards"].insert_one(doc)
    # Log it either way — creating a private item is still YOUR history, kept
    # for your own feed and marked private so a co-parent's feed never sees it.
    await log_activity(database, user, "task_created", doc.get("title", ""),
                       shared=bool(doc.get("shared")), ref=doc["card_id"])

    # Only ping the co-parent when the item is actually shared — private items
    # are silent by design.
    # The assignee gets the hand-off push below; excluding them here means one
    # person never receives both a "new card" alert and an "assigned to you"
    # push for the same card.
    assignee_uid = None
    if doc.get("assignee") and doc.get("shared"):
        assignee_uid = await resolve_member_user_id(database, user["family_id"], doc["assignee"])

    if doc.get("shared"):
        try:
            await send_new_card_alert(
                user["family_id"], doc, created_by_user_id=user["user_id"],
                exclude_user_ids={assignee_uid} if assignee_uid else None)
        except Exception as e:
            log.warning("new card alert failed: %s", e)

    if doc.get("assignee") and doc.get("shared"):
        await log_activity(database, user, "task_assigned", doc["title"], target=doc["assignee"])
        await notify_assignment(database, user, doc, doc["assignee"])

    return public_card(doc)


@app.patch("/api/cards/{card_id}")
async def update_card(card_id: str, payload: CardPatchIn, user=Depends(require_user)):
    database = get_db()
    card = await database["cards"].find_one(
        {"card_id": card_id, "family_id": user["family_id"]},
        {"_id": 0},
    )
    if not card:
        raise HTTPException(status_code=404, detail="Card not found")

    changes = {}
    award_child = False

    if payload.status is not None:
        if payload.status not in {"OPEN", "DONE"}:
            raise HTTPException(status_code=400, detail="Invalid card status")
        changes["status"] = payload.status
        done = payload.status == "DONE"
        changes["completed_at"] = utcnow() if done else None
        # Who ticked it off — the question a co-parent actually asks.
        changes["completed_by_user_id"] = user["user_id"] if done else None
        changes["completed_by_name"] = (user.get("name") or "") if done else None
        award_child = (
            card["status"] != "DONE"
            and payload.status == "DONE"
            and card["type"] == "TASK"
            and bool(card.get("assignee"))
            and not card.get("stars_awarded")
        )
        if award_child:
            changes["stars_awarded"] = True

    if payload.type is not None:
        if payload.type not in {"SIGN_SLIP", "RSVP", "TASK", "BIRTHDAY", "SCHOOL", "APPOINTMENT", "VACATION"}:
            raise HTTPException(status_code=400, detail="Invalid card type")
        changes["type"] = payload.type

    if payload.title is not None:
        title = payload.title.strip()
        if not title:
            raise HTTPException(status_code=400, detail="Title is required")
        changes["title"] = title

    if payload.description is not None:
        changes["description"] = payload.description.strip() or None

    if payload.assignee is not None:
        changes["assignee"] = payload.assignee.strip() or None

    if payload.due_date is not None:
        changes["due_date"] = parse_dt(payload.due_date) if payload.due_date else None

    if payload.recurrence is not None:
        if payload.recurrence not in {"none", "daily", "weekly", "monthly"}:
            raise HTTPException(status_code=400, detail="Invalid recurrence")
        changes["recurrence"] = payload.recurrence

    if payload.reminder_minutes is not None:
        if payload.reminder_minutes < 0:
            raise HTTPException(status_code=400, detail="Reminder must be zero or positive")
        changes["reminder_minutes"] = payload.reminder_minutes

    if payload.shared is not None:
        # Only the person who added a private item may change its sharing here
        # (making it private again). Sharing-with-notification goes through the
        # dedicated /share endpoint. Legacy items (no owner) are family-wide.
        owner = card.get("created_by_user_id")
        if owner and owner != user["user_id"]:
            raise HTTPException(status_code=403, detail="Only the person who added this can change its sharing")
        changes["shared"] = bool(payload.shared)

    if not changes:
        return public_card(card)

    if award_child:
        # Claim the award in the same write that marks it done, so only one
        # request can ever be the one that pays. A snapshot taken moments
        # earlier let two taps both believe the chore was unpaid.
        claim = await database["cards"].update_one(
            {"card_id": card_id, "family_id": user["family_id"],
             "stars_awarded": {"$ne": True}, "status": {"$ne": "DONE"}},
            {"$set": changes})
        award_child = claim.matched_count > 0
        if not award_child:
            await database["cards"].update_one({"card_id": card_id}, {"$set": changes})
    else:
        await database["cards"].update_one({"card_id": card_id}, {"$set": changes})
    updated = await database["cards"].find_one({"card_id": card_id}, {"_id": 0})

    if changes.get("status") == "DONE" and card["status"] != "DONE":
        # A private card's title must never surface in a co-parent's feed — a
        # private "Consult divorce lawyer" marked done would disclose its
        # title. So the completion is logged either way, but with the card's
        # own `shared` flag: a private one stays in YOUR feed (it is your
        # history) and `list_activity` keeps it out of everyone else's. `ref`
        # links the line to the card, so deleting the card clears the line.
        merged = {**card, **changes}
        await log_activity(database, user, "task_done", card.get("title", ""),
                           shared=bool(merged.get("shared")), ref=card_id)

    # A hand-off: the name changed, so somebody has just been given a job.
    # Only on a real change — re-saving a card without touching the assignee
    # must not ping them again.
    new_assignee = changes.get("assignee")
    if new_assignee and new_assignee != card.get("assignee"):
        merged = {**card, **changes}
        # Both the feed entry and the push carry the card title, so both are
        # gated: a private card handed to the co-parent must not announce
        # itself. (They could not open it anyway — it is private — so the
        # notification would only leak the title with no way to act on it.)
        if merged.get("shared"):
            await log_activity(database, user, "task_assigned",
                               merged.get("title", ""), target=new_assignee)
            await notify_assignment(database, user, merged, new_assignee)

    # When a recurring card is completed, spawn its next occurrence.
    if (
        changes.get("status") == "DONE"
        and card["status"] != "DONE"
        and card.get("recurrence") in {"daily", "weekly", "monthly"}
    ):
        base_due = ensure_aware_utc(card.get("due_date")) or utcnow()
        next_doc = {
            "card_id": new_id("card"),
            "family_id": user["family_id"],
            "type": card["type"],
            "title": card["title"],
            "description": card.get("description"),
            "assignee": card.get("assignee"),
            "due_date": advance_due_date(base_due, card["recurrence"]),
            "status": "OPEN",
            "source": card.get("source", "MANUAL"),
            "image_base64": card.get("image_base64"),
            "recurrence": card["recurrence"],
            "reminder_minutes": card.get("reminder_minutes", 0),
            "created_at": utcnow(),
            "completed_at": None,
            # The next occurrence keeps the same privacy as its parent.
            "created_by_user_id": card.get("created_by_user_id"),
            "shared": card.get("shared", False),
        }
        await database["cards"].insert_one(next_doc)

    if award_child:
        member = await database["family_members"].find_one(
            {
                "family_id": user["family_id"],
                "name": card["assignee"],
                "role": {"$regex": "^child$", "$options": "i"},
            },
            {"_id": 0},
        )
        if member:
            old_stars = int(member.get("stars", 0))
            # Route through the shared helper so a finished chore ticks the
            # weekly meter (week_earned, which gates weekend treats) and writes
            # the star ledger — a raw $inc banked the stars but left both
            # untouched, silently locking the weekend payoff and letting the
            # ledger sum diverge from the balance.
            await award_stars_to_member(
                database, user["family_id"], member["member_id"], 5,
                card.get("title") or "Chore done", user,
            )
            await send_star_milestone_alert(user["family_id"], member.get("name", "Your child"), old_stars, old_stars + 5)

    return public_card(updated)


@app.post("/api/cards/{card_id}/unshare")
async def unshare_card(card_id: str, user=Depends(require_user)):
    """Pull a shared item back to private.

    A twin of /share rather than a PATCH on purpose. The offline queue replays
    PATCHes, so `updateCard(id, {shared: false})` resolved optimistically with
    no signal — the row left the sharing panel, nothing failed, and the item
    stayed visible to the co-parent for as long as the phone was offline. A
    revoke is the one write that must never claim success it has not got, so it
    takes a path the queue will not swallow.

    No notification: telling someone the moment you stop sharing with them
    turns a private decision into an announcement.
    """
    database = get_db()
    card = await database["cards"].find_one(
        {"card_id": card_id, "family_id": user["family_id"]},
        {"_id": 0},
    )
    if not card:
        raise HTTPException(status_code=404, detail="Card not found")

    owner = card.get("created_by_user_id")
    if owner and owner != user["user_id"]:
        raise HTTPException(status_code=403, detail="Only the person who added this can change its sharing")

    if card.get("shared"):
        await database["cards"].update_one({"card_id": card_id}, {"$set": {"shared": False}})
        card["shared"] = False

    return public_card(card)


@app.post("/api/cards/{card_id}/share")
async def share_card(card_id: str, user=Depends(require_user)):
    """Share a private calendar item with the co-parent and notify them.
    Only the person who added the item can share it."""
    database = get_db()
    card = await database["cards"].find_one(
        {"card_id": card_id, "family_id": user["family_id"]},
        {"_id": 0},
    )
    if not card:
        raise HTTPException(status_code=404, detail="Card not found")

    owner = card.get("created_by_user_id")
    if owner and owner != user["user_id"]:
        raise HTTPException(status_code=403, detail="Only the person who added this can share it")

    if not card.get("shared"):
        await database["cards"].update_one({"card_id": card_id}, {"$set": {"shared": True}})
        card["shared"] = True
        try:
            sharer = user.get("name") or "A parent"
            await send_coparent_alert(
                user["family_id"],
                f"{sharer} shared a calendar item",
                card.get("title") or "New shared item",
                "shared_card",
                created_by_user_id=user["user_id"],
            )
        except Exception as e:
            log.warning("share card alert failed: %s", e)

    return public_card(card)


@app.delete("/api/cards/{card_id}")
async def delete_card(card_id: str, user=Depends(require_user)):
    database = get_db()
    await database["cards"].delete_one({"card_id": card_id, "family_id": user["family_id"]})
    # Deleting a card from the completed history takes its feed lines with it —
    # "done" and "created" for a thing that no longer exists would otherwise
    # linger. Scoped to the family and the card's own id, so nothing else moves.
    await database["activity"].delete_many(
        {"family_id": user["family_id"], "ref": card_id})
    return {"ok": True}


@app.get("/api/cards/conflicts")
async def card_conflicts(
    due_date: str,
    exclude_id: Optional[str] = None,
    user=Depends(require_user),
):
    database = get_db()
    target = parse_dt(due_date)
    if not target:
        return []

    start = target - timedelta(hours=2)
    end = target + timedelta(hours=2)

    query = {
        "family_id": user["family_id"],
        "due_date": {"$gte": start, "$lte": end},
        # Don't surface a co-parent's private items as conflicts.
        "$or": [
            {"shared": True},
            {"created_by_user_id": user["user_id"]},
            {"created_by_user_id": {"$exists": False}},
        ],
    }
    if exclude_id:
        query["card_id"] = {"$ne": exclude_id}

    rows = []
    async for item in database["cards"].find(query, {"_id": 0}):
        rows.append(public_card(item))
    return rows


# -----------------------------------------------------------------------------
# Vault
# -----------------------------------------------------------------------------
def _may_see_vault_doc(doc: dict, user: dict) -> bool:
    """Shared docs, your own docs, and unclaimed legacy docs (no owner)."""
    if (doc.get("visibility") or "shared") == "shared":
        return True
    return doc.get("owner_user_id") in (None, user["user_id"])


@app.get("/api/vault")
async def list_vault(user=Depends(require_full_member)):
    database = get_db()
    rows = []
    async for item in database["vault"].find({"family_id": user["family_id"]}, {"_id": 0}).sort("created_at", -1):
        if _may_see_vault_doc(item, user):
            rows.append(public_vault_doc(item))
    return rows


@app.patch("/api/vault/{doc_id}/visibility")
async def set_vault_visibility(doc_id: str, payload: VaultVisibilityIn, user=Depends(require_full_member)):
    """Flip a document between private and shared.

    Allowed for the owner, and for any member on an unclaimed legacy doc —
    claiming it makes the caller its owner, which is how a family converts
    documents uploaded before this feature into genuinely private ones.
    """
    database = get_db()
    visibility = "private" if payload.visibility == "private" else "shared"
    doc = await database["vault"].find_one(
        {"doc_id": doc_id, "family_id": user["family_id"]}, {"_id": 0}
    )
    if not doc:
        raise HTTPException(status_code=404, detail="Document not found")
    owner = doc.get("owner_user_id")
    if owner not in (None, user["user_id"]):
        raise HTTPException(status_code=403, detail="Only the owner can change this document")
    await database["vault"].update_one(
        {"doc_id": doc_id, "family_id": user["family_id"]},
        {"$set": {
            "visibility": visibility,
            "owner_user_id": owner or user["user_id"],
            "owner_name": doc.get("owner_name") or user.get("name"),
            "updated_at": utcnow(),
        }},
    )
    fresh = await database["vault"].find_one(
        {"doc_id": doc_id, "family_id": user["family_id"]}, {"_id": 0}
    )
    return public_vault_doc(fresh)


@app.post("/api/vault")
async def create_vault_doc(payload: VaultIn, user=Depends(require_full_member)):
    database = get_db()
    sub = await build_subscription(user["family_id"])
    family = await get_family_doc(user["family_id"])
    size = len(payload.image_base64.encode("utf-8"))

    if not is_admin_user(user) and family.get("vault_bytes_used", 0) + size > sub["limits"]["vault_bytes"]:
        plan_limit_error(
            feature="vault_storage",
            current_plan=sub["plan"],
            limit=sub["limits"]["vault_bytes"],
            used=family.get("vault_bytes_used", 0),
            message="Vault storage limit reached for your current plan.",
        )

    doc = {
        "doc_id": new_id("doc"),
        "family_id": user["family_id"],
        "title": payload.title,
        "category": payload.category,
        "image_base64": payload.image_base64,
        "mime_type": payload.mime_type or "image/jpeg",
        "file_name": payload.file_name,
        # Private unless the uploader says otherwise: a passport scan is not
        # household news just because the household shares an app.
        "visibility": "shared" if payload.visibility == "shared" else "private",
        "owner_user_id": user["user_id"],
        "owner_name": user.get("name"),
        "created_at": utcnow(),
    }
    await database["vault"].insert_one(doc)
    await database["families"].update_one(
        {"family_id": user["family_id"]},
        {"$inc": {"vault_bytes_used": size}, "$set": {"updated_at": utcnow()}},
    )
    return public_vault_doc(doc)


def _decode_doc_bytes(image_base64: str) -> bytes:
    data = image_base64 or ""
    if "," in data:
        data = data.split(",", 1)[1]
    return base64.b64decode(data)


def _docx_to_html(raw: bytes) -> str:
    from docx import Document  # local import keeps startup light
    d = Document(io.BytesIO(raw))
    parts = []
    for p in d.paragraphs:
        text = html.escape(p.text)
        style = (p.style.name if p.style else "") or ""
        if not p.text.strip():
            parts.append("<br/>")
        elif style == "Title" or style.startswith("Heading 1"):
            parts.append(f"<h1>{text}</h1>")
        elif style.startswith("Heading"):
            parts.append(f"<h2>{text}</h2>")
        else:
            parts.append(f"<p>{text}</p>")
    for table in d.tables:
        rows = []
        for row in table.rows:
            cells = "".join(f"<td>{html.escape(c.text)}</td>" for c in row.cells)
            rows.append(f"<tr>{cells}</tr>")
        if rows:
            parts.append(f"<table>{''.join(rows)}</table>")
    return "".join(parts) or "<p><em>(empty document)</em></p>"


def _xlsx_to_html(raw: bytes) -> str:
    import pandas as pd
    sheets = pd.read_excel(io.BytesIO(raw), sheet_name=None, engine="openpyxl")
    parts = []
    for name, df in sheets.items():
        parts.append(f"<h2>{html.escape(str(name))}</h2>")
        parts.append(df.to_html(index=False, na_rep="", border=0))
    return "".join(parts) or "<p><em>(empty spreadsheet)</em></p>"


DOCX_MIME = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
XLSX_MIME = "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"


@app.get("/api/vault/{doc_id}/render")
async def render_vault_doc(doc_id: str, user=Depends(require_full_member)):
    """Return an in-app-viewable form of a document. Images/PDFs are rendered
    client-side (the app already has the bytes); Word/Excel are converted to
    readable HTML here so they can be viewed without an external app."""
    database = get_db()
    doc = await database["vault"].find_one(
        {"doc_id": doc_id, "family_id": user["family_id"]}, {"_id": 0}
    )
    if not doc or not _may_see_vault_doc(doc, user):
        raise HTTPException(status_code=404, detail="Document not found")
    mime = doc.get("mime_type") or "image/jpeg"
    if mime.startswith("image/"):
        return {"kind": "image"}
    if mime == "application/pdf":
        return {"kind": "pdf"}
    try:
        raw = _decode_doc_bytes(doc.get("image_base64", ""))
        if mime == DOCX_MIME:
            return {"kind": "html", "html": _docx_to_html(raw)}
        if mime == XLSX_MIME:
            return {"kind": "html", "html": _xlsx_to_html(raw)}
    except Exception as e:
        log.warning("vault render failed for %s: %s", doc_id, e)
        return {"kind": "unsupported"}
    return {"kind": "unsupported"}


@app.delete("/api/vault/{doc_id}")
async def delete_vault_doc(doc_id: str, user=Depends(require_full_member)):
    database = get_db()
    doc = await database["vault"].find_one(
        {"doc_id": doc_id, "family_id": user["family_id"]},
        {"_id": 0},
    )
    if doc and doc.get("owner_user_id") not in (None, user["user_id"]):
        raise HTTPException(status_code=403, detail="Only the owner can delete this document")
    if doc:
        size = len(doc["image_base64"].encode("utf-8"))
        await database["vault"].delete_one({"doc_id": doc_id})
        await database["families"].update_one(
            {"family_id": user["family_id"]},
            {"$inc": {"vault_bytes_used": -size}, "$set": {"updated_at": utcnow()}},
        )
    return {"ok": True}


# -----------------------------------------------------------------------------
# Rewards
# -----------------------------------------------------------------------------
@app.get("/api/rewards")
async def list_rewards(user=Depends(require_user)):
    database = get_db()
    rows = []
    async for item in database["rewards"].find({"family_id": user["family_id"]}, {"_id": 0}).sort("created_at", -1):
        rows.append(public_reward(item))
    return rows


@app.post("/api/rewards")
async def create_reward(payload: RewardIn, user=Depends(require_user)):
    database = get_db()
    reward = {
        "reward_id": new_id("reward"),
        "family_id": user["family_id"],
        "title": payload.title,
        "cost_stars": payload.cost_stars,
        "icon": payload.icon,
        "weekend": bool(payload.weekend),
        "created_at": utcnow(),
    }
    await database["rewards"].insert_one(reward)
    return public_reward(reward)


@app.patch("/api/rewards/{reward_id}")
async def update_reward(reward_id: str, payload: RewardPatchIn, user=Depends(require_user)):
    database = get_db()
    reward = await database["rewards"].find_one(
        {"reward_id": reward_id, "family_id": user["family_id"]},
        {"_id": 0},
    )
    if not reward:
        raise HTTPException(status_code=404, detail="Reward not found")

    changes = {}

    if payload.title is not None:
        title = payload.title.strip()
        if not title:
            raise HTTPException(status_code=400, detail="Reward title is required")
        changes["title"] = title

    if payload.cost_stars is not None:
        if payload.cost_stars < 1:
            raise HTTPException(status_code=400, detail="Reward cost must be at least 1 star")
        changes["cost_stars"] = payload.cost_stars

    if payload.icon is not None:
        changes["icon"] = payload.icon.strip() or None

    if payload.weekend is not None:
        changes["weekend"] = bool(payload.weekend)

    if not changes:
        return public_reward(reward)

    await database["rewards"].update_one(
        {"reward_id": reward_id},
        {"$set": changes},
    )
    updated = await database["rewards"].find_one({"reward_id": reward_id}, {"_id": 0})
    return public_reward(updated)



@app.delete("/api/rewards/{reward_id}")
async def delete_reward(reward_id: str, user=Depends(require_user)):
    database = get_db()
    await database["rewards"].delete_one({"reward_id": reward_id, "family_id": user["family_id"]})
    return {"ok": True}


@app.post("/api/rewards/{reward_id}/redeem")
async def redeem_reward(reward_id: str, payload: RedeemIn, user=Depends(require_user)):
    database = get_db()
    reward = await database["rewards"].find_one(
        {"reward_id": reward_id, "family_id": user["family_id"]},
        {"_id": 0},
    )
    member = await database["family_members"].find_one(
        {"member_id": payload.member_id, "family_id": user["family_id"]},
        {"_id": 0},
    )
    if not reward or not member:
        raise HTTPException(status_code=404, detail="Reward or member not found")

    cost = int(reward["cost_stars"])
    member = await roll_week_if_stale(database, member)

    # Atomic guarded decrement. Checking the balance and then decrementing in
    # two steps let two taps (or two devices) both pass the check and both
    # spend, driving the balance negative. The filter carries the check, so the
    # database rejects the second one.
    #
    # A weekend treat carries a second gate: the child must have earned enough
    # THIS WEEK (week_earned), not merely have a big enough bank. Redeeming one
    # spends the bank and also draws down the weekly meter, so the same week's
    # earnings can't fund two weekend treats. A saved-up (non-weekend) reward
    # has no weekly gate and leaves the meter alone.
    star_filter = {
        "member_id": member["member_id"],
        "family_id": user["family_id"],
        "stars": {"$gte": cost},
    }
    inc = {"stars": -cost}
    if reward.get("weekend"):
        star_filter["week_earned"] = {"$gte": cost}
        inc["week_earned"] = -cost
    result = await database["family_members"].update_one(star_filter, {"$inc": inc})
    if result.matched_count == 0:
        if reward.get("weekend"):
            raise HTTPException(
                status_code=400,
                detail="Not enough stars earned this week for a weekend treat.")
        raise HTTPException(status_code=400, detail="Not enough stars")

    # A redemption is a star movement and belongs in the ledger. Without this
    # the balance dropped with nothing in Recent Activity to explain it, and
    # the weekly total — computed from history — counted awards but never
    # spends, so the ledger stopped reconciling with the balance.
    transaction = {
        "transaction_id": new_id("star"),
        "family_id": user["family_id"],
        "member_id": member["member_id"],
        "delta": -cost,
        "reason": reward.get("title") or "Reward redeemed",
        "created_by_user_id": user["user_id"],
        "created_by_name": user.get("name"),
        "created_at": utcnow(),
    }
    await database["star_transactions"].insert_one(transaction)

    # Spending the stars is only half of it — somebody still has to hand over
    # the ice cream. Without a record the promise lived in the parent's head,
    # and a child who had paid had no way to show they were still owed it.
    redemption = {
        "redemption_id": new_id("redemption"),
        "family_id": user["family_id"],
        "member_id": member["member_id"],
        "reward_id": reward["reward_id"],
        "reward_title": reward.get("title") or "",
        "reward_icon": reward.get("icon"),
        "cost_stars": cost,
        "status": "pending",
        "created_at": utcnow(),
        "created_by_user_id": user["user_id"],
        "fulfilled_at": None,
    }
    await database["redemptions"].insert_one(redemption)
    # The other parent hears that a reward was spent — the one who did it is
    # excluded (they just did it and are watching it happen).
    try:
        await send_coparent_alert(
            user["family_id"],
            f"{member.get('name') or 'Your child'} redeemed a reward",
            reward.get("title") or "", "reward_redeemed", created_by_user_id=user["user_id"])
    except Exception as e:
        log.warning("reward alert failed: %s", e)

    member = await database["family_members"].find_one({"member_id": member["member_id"]}, {"_id": 0})
    return {
        "ok": True,
        "member": public_member(member),
        "transaction": public_star_transaction(transaction),
        "redemption": public_redemption(redemption),
    }


# -----------------------------------------------------------------------------
# Redemptions — what has been paid for and not yet handed over
# -----------------------------------------------------------------------------
@app.get("/api/redemptions")
async def list_redemptions(status: Optional[str] = None, user=Depends(require_user)):
    database = get_db()
    query: dict = {"family_id": user["family_id"]}
    if status is not None:
        # Reject an unknown value rather than quietly ignoring the filter: a
        # typo returning every status looks like the filter worked.
        if status not in ("pending", "fulfilled", "cancelled"):
            raise HTTPException(status_code=400, detail="Unknown redemption status")
        query["status"] = status

    cursor = database["redemptions"].find(query, {"_id": 0}).sort("created_at", -1).limit(500)
    return [public_redemption(item) async for item in cursor]


@app.post("/api/redemptions/{redemption_id}/fulfil")
async def fulfil_redemption(redemption_id: str, user=Depends(require_user)):
    database = get_db()

    # The filter carries the "still pending" rule so two parents both tapping
    # Given produce one state change, not two, and the second gets a clear
    # answer instead of silently overwriting the first one's timestamp.
    result = await database["redemptions"].update_one(
        {
            "redemption_id": redemption_id,
            "family_id": user["family_id"],
            "status": "pending",
        },
        {
            "$set": {
                "status": "fulfilled",
                "fulfilled_at": utcnow(),
                "fulfilled_by_user_id": user["user_id"],
            }
        },
    )
    if result.matched_count == 0:
        raise HTTPException(status_code=404, detail="Redemption not found or already settled")

    updated = await database["redemptions"].find_one({"redemption_id": redemption_id}, {"_id": 0})
    return public_redemption(updated)


@app.post("/api/redemptions/{redemption_id}/cancel")
async def cancel_redemption(redemption_id: str, user=Depends(require_user)):
    """Give the stars back when a reward cannot be delivered.

    A parent may promise a trip to the cinema and then find it sold out. The
    honest resolution is to return what the child paid, not to quietly mark it
    given, so this refunds the stars and writes the refund to the ledger where
    the child can see it.
    """
    database = get_db()

    result = await database["redemptions"].update_one(
        {
            "redemption_id": redemption_id,
            "family_id": user["family_id"],
            "status": "pending",
        },
        {
            "$set": {
                "status": "cancelled",
                "cancelled_at": utcnow(),
                "cancelled_by_user_id": user["user_id"],
            }
        },
    )
    # Claim the redemption before touching the balance. If two parents cancel
    # at once only one update matches, so the stars are refunded exactly once.
    if result.matched_count == 0:
        raise HTTPException(status_code=404, detail="Redemption not found or already settled")

    redemption = await database["redemptions"].find_one({"redemption_id": redemption_id}, {"_id": 0})
    cost = int(redemption.get("cost_stars", 0) or 0)

    transaction = None
    if cost > 0:
        # A weekend treat costs the bank AND the week's meter, so a refund has
        # to restore both — otherwise a child whose cinema trip fell through
        # got their stars back but stayed locked out of weekend treats for the
        # rest of the week, having earned them. Only within the same week: a
        # cancellation after Monday must not credit last week's effort into a
        # fresh meter.
        refund = {"stars": cost}
        redeemed_at = _coerce_dt(redemption.get("created_at"))
        if redemption.get("weekend") and redeemed_at and redeemed_at >= current_week_start():
            refund["week_earned"] = cost
        credited = await database["family_members"].update_one(
            {"member_id": redemption["member_id"], "family_id": user["family_id"]},
            {"$inc": refund},
        )
        # Only write the ledger row if the balance actually moved. If the child
        # was removed between claiming the redemption and crediting it, an
        # unconditional insert would leave a +60 in the history that no balance
        # ever received, and the ledger would stop reconciling.
        if credited.matched_count:
            transaction = {
                "transaction_id": new_id("star"),
                "family_id": user["family_id"],
                "member_id": redemption["member_id"],
                "delta": cost,
                # The bare title, matching the spend entry. An English word like
                # "Refund:" would sit untranslated in a German family's ledger,
                # and the + sign against the earlier − already tells the story.
                "reason": redemption.get("reward_title") or "",
                "created_by_user_id": user["user_id"],
                "created_by_name": user.get("name"),
                "created_at": utcnow(),
            }
            await database["star_transactions"].insert_one(transaction)

    member = await database["family_members"].find_one(
        {"member_id": redemption["member_id"], "family_id": user["family_id"]}, {"_id": 0}
    )
    return {
        "ok": True,
        "redemption": public_redemption(redemption),
        "member": public_member(member) if member else None,
        "transaction": public_star_transaction(transaction) if transaction else None,
    }



# -----------------------------------------------------------------------------
# Google Calendar sync
# -----------------------------------------------------------------------------
def _parse_google_event_start(event: dict) -> Optional[datetime]:
    start = event.get("start") or {}
    raw = start.get("dateTime")

    if raw:
        return parse_dt(raw)

    raw_date = start.get("date")
    if raw_date:
        try:
            return datetime.fromisoformat(raw_date).replace(
                hour=9,
                minute=0,
                second=0,
                microsecond=0,
                tzinfo=timezone.utc,
            )
        except Exception:
            return None

    return None


def _event_attendee_contacts(event: dict) -> list[dict]:
    contacts = []
    seen = set()

    for key in ("organizer", "creator"):
        person = event.get(key) or {}
        email = str(person.get("email") or "").strip().lower()
        if email and email not in seen:
            seen.add(email)
            contacts.append({
                "email": email,
                "name": person.get("displayName") or email.split("@")[0],
            })

    for attendee in event.get("attendees") or []:
        email = str(attendee.get("email") or "").strip().lower()
        if email and email not in seen:
            seen.add(email)
            contacts.append({
                "email": email,
                "name": attendee.get("displayName") or email.split("@")[0],
            })

    return contacts


async def _fetch_google_calendar_events(access_token: str, days: int) -> list[dict]:
    now = utcnow()
    time_min = now.isoformat().replace("+00:00", "Z")
    time_max = (now + timedelta(days=days)).isoformat().replace("+00:00", "Z")

    params = urllib.parse.urlencode(
        {
            "timeMin": time_min,
            "timeMax": time_max,
            "singleEvents": "true",
            "orderBy": "startTime",
            "maxResults": "100",
            # Cancellations must arrive too, or a deleted meeting lives on
            # in the app forever. Google omits them unless asked.
            "showDeleted": "true",
        }
    )

    url = f"https://www.googleapis.com/calendar/v3/calendars/primary/events?{params}"

    def _request():
        req = urllib.request.Request(
            url,
            headers={"Authorization": f"Bearer {access_token}"},
            method="GET",
        )

        try:
            with urllib.request.urlopen(req, timeout=20) as response:
                raw = response.read().decode("utf-8")
                return json.loads(raw)
        except urllib.error.HTTPError as e:
            body = e.read().decode("utf-8", errors="replace")
            # Never let Google's own 401/403 propagate as OUR 401 — the app
            # treats a 401 as an expired session and signs the user out. Remap
            # an upstream auth failure to a clear, non-session error.
            if e.code in (401, 403):
                raise HTTPException(
                    status_code=400,
                    detail="Google Calendar access was denied or expired. Please reconnect your Google account and try again.",
                )
            raise HTTPException(status_code=502, detail=f"Google Calendar error: {body[:200]}")
        except HTTPException:
            raise
        except Exception as e:
            raise HTTPException(status_code=502, detail=f"Google Calendar request failed: {e}")

    data = await asyncio.to_thread(_request)
    return data.get("items") or []


@app.post("/api/calendar/import")
async def import_google_calendar(payload: CalendarImportIn, user=Depends(require_user)):
    database = get_db()

    token = payload.access_token.strip()
    if not token:
        raise HTTPException(status_code=400, detail="Google Calendar access token is required")

    days = max(1, min(payload.days or 30, 90))
    events = await _fetch_google_calendar_events(token, days)

    imported = 0
    updated = 0
    removed = 0
    skipped = 0
    contacts_found: dict[str, dict] = {}

    for event in events:
        event_id = event.get("id")

        if event.get("status") == "cancelled":
            # A meeting that no longer exists must not live on as a card.
            # Only open, unshared imports are removed — a completed card is
            # history, and a shared one is the family's now.
            if event_id:
                result = await database["cards"].delete_one({
                    "family_id": user["family_id"],
                    "google_event_id": event_id,
                    "status": "OPEN",
                    "shared": False,
                })
                if result.deleted_count:
                    removed += 1
                    continue
            skipped += 1
            continue

        if not event_id:
            skipped += 1
            continue

        start_dt = _parse_google_event_start(event)
        if not start_dt:
            skipped += 1
            continue

        title = (event.get("summary") or "Calendar event").strip()
        location = (event.get("location") or "").strip()
        html_link = event.get("htmlLink")
        contacts = _event_attendee_contacts(event)

        existing = await database["cards"].find_one(
            {
                "family_id": user["family_id"],
                "google_event_id": event_id,
            },
            {"_id": 0},
        )
        if existing:
            # The field bug: an already-imported event was skipped outright,
            # so a rescheduled meeting kept its old time in the app forever —
            # "sync" only ever discovered new events. Mirror what changed;
            # the card's status stays whatever the family made it.
            changed = {}
            if title != existing.get("title"):
                changed["title"] = title
            if ensure_aware_utc(existing.get("due_date")) != start_dt:
                changed["due_date"] = start_dt
            if changed:
                await database["cards"].update_one(
                    {"family_id": user["family_id"], "google_event_id": event_id},
                    {"$set": changed},
                )
                updated += 1
            else:
                skipped += 1
            continue

        for contact in contacts:
            email = contact["email"]
            contacts_found[email] = contact
            await database["calendar_contacts"].update_one(
                {"family_id": user["family_id"], "email": email},
                {
                    "$set": {
                        "name": contact.get("name") or email.split("@")[0],
                        "last_seen_at": utcnow(),
                        "last_event_title": title,
                    },
                    "$setOnInsert": {
                        "family_id": user["family_id"],
                        "email": email,
                        "first_seen_at": utcnow(),
                    },
                    "$inc": {"event_count": 1},
                },
                upsert=True,
            )

        contact_line = ""
        if contacts:
            contact_line = "People: " + ", ".join([c["email"] for c in contacts[:8]])

        description_parts = [
            (event.get("description") or "").strip(),
            f"Location: {location}" if location else "",
            contact_line,
            f"Google Calendar: {html_link}" if html_link else "",
        ]

        card = {
            "card_id": new_id("card"),
            "family_id": user["family_id"],
            "type": "TASK",
            "title": title,
            "description": "\n".join([p for p in description_parts if p]),
            "assignee": user.get("name"),
            "due_date": start_dt,
            "status": "OPEN",
            "source": "CALENDAR",
            "image_base64": None,
            "recurrence": "none",
            "reminder_minutes": 60,
            "google_event_id": event_id,
            "google_ical_uid": event.get("iCalUID"),
            "external_source": "google_calendar",
            "created_at": utcnow(),
            "completed_at": None,
            # A parent's Google Calendar is personal — imported events are
            # private to them until they choose to share.
            "created_by_user_id": user["user_id"],
            "shared": False,
        }

        await database["cards"].insert_one(card)
        imported += 1

    contacts = []
    if contacts_found:
        cursor = database["calendar_contacts"].find(
            {"family_id": user["family_id"], "email": {"$in": list(contacts_found.keys())}},
            {"_id": 0},
        ).sort("last_seen_at", -1)

        async for item in cursor:
            contacts.append(public_calendar_contact(item))

    return {
        "ok": True,
        "imported": imported,
        "updated": updated,
        "removed": removed,
        "skipped": skipped,
        "events_seen": len(events),
        "contacts_found": len(contacts_found),
        "contacts": contacts,
        "days": days,
    }


def _parse_ms_event_start(event: dict):
    start = event.get("start") or {}
    dt_str = str(start.get("dateTime") or "").strip()
    if not dt_str:
        return None
    # Graph returns up to 7 fractional digits; Python's fromisoformat takes 6.
    if "." in dt_str:
        head, frac = dt_str.split(".", 1)
        frac = "".join(ch for ch in frac if ch.isdigit())[:6]
        dt_str = f"{head}.{frac}" if frac else head
    try:
        dt = datetime.fromisoformat(dt_str)
    except ValueError:
        return None
    # We request Prefer: outlook.timezone="UTC", so naive values are UTC.
    if dt.tzinfo is None:
        dt = dt.replace(tzinfo=timezone.utc)
    return dt.astimezone(timezone.utc)


def _ms_event_contacts(event: dict) -> list[dict]:
    seen: set[str] = set()
    contacts: list[dict] = []
    org = (event.get("organizer") or {}).get("emailAddress") or {}
    people = [org] + [(a.get("emailAddress") or {}) for a in (event.get("attendees") or [])]
    for person in people:
        email = str(person.get("address") or "").strip().lower()
        if email and email not in seen:
            seen.add(email)
            contacts.append({"email": email, "name": person.get("name") or email.split("@")[0]})
    return contacts


async def _fetch_microsoft_calendar_events(access_token: str, days: int) -> list[dict]:
    now = utcnow()
    start = now.isoformat().replace("+00:00", "Z")
    end = (now + timedelta(days=days)).isoformat().replace("+00:00", "Z")
    params = urllib.parse.urlencode({
        "startDateTime": start,
        "endDateTime": end,
        "$orderby": "start/dateTime",
        "$top": "100",
        "$select": "id,subject,start,end,location,bodyPreview,webLink,attendees,isCancelled,organizer,iCalUId",
    })
    url = f"https://graph.microsoft.com/v1.0/me/calendarView?{params}"

    def _request():
        req = urllib.request.Request(
            url,
            headers={"Authorization": f"Bearer {access_token}", "Prefer": 'outlook.timezone="UTC"'},
            method="GET",
        )
        try:
            with urllib.request.urlopen(req, timeout=20) as response:
                return json.loads(response.read().decode("utf-8"))
        except urllib.error.HTTPError as e:
            body = e.read().decode("utf-8", errors="replace")
            # Remap upstream auth failures so the app doesn't read them as an
            # expired session (which would sign the user out).
            if e.code in (401, 403):
                raise HTTPException(
                    status_code=400,
                    detail="Outlook access was denied or expired. Please reconnect your Microsoft account and try again.",
                )
            raise HTTPException(status_code=502, detail=f"Outlook Calendar error: {body[:200]}")
        except HTTPException:
            raise
        except Exception as e:
            raise HTTPException(status_code=502, detail=f"Outlook Calendar request failed: {e}")

    data = await asyncio.to_thread(_request)
    return data.get("value") or []


@app.post("/api/calendar/import-microsoft")
async def import_microsoft_calendar(payload: CalendarImportIn, user=Depends(require_user)):
    """Import upcoming Outlook/Microsoft events via Graph — mirror of the Google
    import. Takes a delegated access token from the app (public-client PKCE),
    so no client secret lives on our side."""
    database = get_db()

    token = payload.access_token.strip()
    if not token:
        raise HTTPException(status_code=400, detail="Microsoft access token is required")

    days = max(1, min(payload.days or 30, 90))
    events = await _fetch_microsoft_calendar_events(token, days)

    imported = 0
    updated = 0
    removed = 0
    skipped = 0
    contacts_found: dict[str, dict] = {}

    for event in events:
        event_id = event.get("id")

        if event.get("isCancelled"):
            # Same rule as the Google import: a cancelled meeting's card goes,
            # but only while it is still open and unshared.
            if event_id:
                result = await database["cards"].delete_one({
                    "family_id": user["family_id"],
                    "ms_event_id": event_id,
                    "status": "OPEN",
                    "shared": False,
                })
                if result.deleted_count:
                    removed += 1
                    continue
            skipped += 1
            continue

        if not event_id:
            skipped += 1
            continue

        start_dt = _parse_ms_event_start(event)
        if not start_dt:
            skipped += 1
            continue

        existing = await database["cards"].find_one(
            {"family_id": user["family_id"], "ms_event_id": event_id},
            {"_id": 0},
        )
        if existing:
            # Mirror what changed rather than skipping — a rescheduled
            # meeting must move here too. Status stays the family's call.
            new_title = (event.get("subject") or "Calendar event").strip()
            changed = {}
            if new_title != existing.get("title"):
                changed["title"] = new_title
            if ensure_aware_utc(existing.get("due_date")) != start_dt:
                changed["due_date"] = start_dt
            if changed:
                await database["cards"].update_one(
                    {"family_id": user["family_id"], "ms_event_id": event_id},
                    {"$set": changed},
                )
                updated += 1
            else:
                skipped += 1
            continue

        title = (event.get("subject") or "Calendar event").strip()
        location = ((event.get("location") or {}).get("displayName") or "").strip()
        web_link = event.get("webLink")
        contacts = _ms_event_contacts(event)

        for contact in contacts:
            email = contact["email"]
            contacts_found[email] = contact
            await database["calendar_contacts"].update_one(
                {"family_id": user["family_id"], "email": email},
                {
                    "$set": {
                        "name": contact.get("name") or email.split("@")[0],
                        "last_seen_at": utcnow(),
                        "last_event_title": title,
                    },
                    "$setOnInsert": {
                        "family_id": user["family_id"],
                        "email": email,
                        "first_seen_at": utcnow(),
                    },
                    "$inc": {"event_count": 1},
                },
                upsert=True,
            )

        contact_line = ("People: " + ", ".join([c["email"] for c in contacts[:8]])) if contacts else ""
        description_parts = [
            (event.get("bodyPreview") or "").strip(),
            f"Location: {location}" if location else "",
            contact_line,
            f"Outlook: {web_link}" if web_link else "",
        ]

        card = {
            "card_id": new_id("card"),
            "family_id": user["family_id"],
            "type": "TASK",
            "title": title,
            "description": "\n".join([p for p in description_parts if p]),
            "assignee": user.get("name"),
            "due_date": start_dt,
            "status": "OPEN",
            "source": "CALENDAR",
            "image_base64": None,
            "recurrence": "none",
            "reminder_minutes": 60,
            "ms_event_id": event_id,
            "google_ical_uid": event.get("iCalUId"),
            "external_source": "microsoft_calendar",
            "created_at": utcnow(),
            "completed_at": None,
            # Personal by default — imported events stay private until shared.
            "created_by_user_id": user["user_id"],
            "shared": False,
        }

        await database["cards"].insert_one(card)
        imported += 1

    contacts = []
    if contacts_found:
        cursor = database["calendar_contacts"].find(
            {"family_id": user["family_id"], "email": {"$in": list(contacts_found.keys())}},
            {"_id": 0},
        ).sort("last_seen_at", -1)
        async for item in cursor:
            contacts.append(public_calendar_contact(item))

    return {
        "ok": True,
        "imported": imported,
        "updated": updated,
        "removed": removed,
        "skipped": skipped,
        "events_seen": len(events),
        "contacts_found": len(contacts_found),
        "contacts": contacts,
        "days": days,
    }


@app.get("/api/calendar/contacts")
async def calendar_contacts(user=Depends(require_user)):
    database = get_db()
    rows = []
    cursor = database["calendar_contacts"].find(
        {"family_id": user["family_id"]},
        {"_id": 0},
    ).sort("last_seen_at", -1).limit(50)

    async for item in cursor:
        rows.append(public_calendar_contact(item))

    return rows



# -----------------------------------------------------------------------------
# Subscription
# -----------------------------------------------------------------------------
@app.get("/api/subscription")
async def get_subscription(user=Depends(require_user)):
    sub = await build_subscription(user["family_id"])
    if is_admin_user(user):
        return apply_admin_subscription(sub)
    return sub


@app.post("/api/subscription/change")
async def change_subscription(payload: SubscriptionChangeIn, user=Depends(require_full_member)):
    database = get_db()
    # Once real billing is configured (RevenueCat webhook secret present),
    # plans are set ONLY by verified purchase events — the self-serve switcher
    # from the testing window locks itself for non-admins automatically.
    if os.environ.get("RC_WEBHOOK_SECRET") and not is_admin_user(user):
        raise HTTPException(status_code=403, detail="Plans change via the store purchase flow")
    if payload.plan not in PLAN_CATALOG:
        raise HTTPException(status_code=400, detail="Invalid plan")
    if payload.billing_cycle not in ("monthly", "yearly"):
        raise HTTPException(status_code=400, detail="Invalid billing cycle")

    await database["families"].update_one(
        {"family_id": user["family_id"]},
        {
            "$set": {
                "plan": payload.plan,
                "billing_cycle": payload.billing_cycle,
                "grandfathered": False,
                "updated_at": utcnow(),
            }
        },
        upsert=True,
    )
    return await build_subscription(user["family_id"])


@app.put("/api/family/custody")
async def set_custody(payload: CustodyConfigIn, user=Depends(require_full_member)):
    """Set the family's alternating-custody config. Parents only (require_full_
    member already refuses helpers and teens) — the schedule is a parent's to
    set, and it shows on every family member's Feed and Calendar."""
    if payload.our_weeks not in CUSTODY_WEEKS:
        raise HTTPException(status_code=400, detail="our_weeks must be 'even' or 'odd'")
    database = get_db()
    await database["families"].update_one(
        {"family_id": user["family_id"]},
        {
            "$set": {
                "custody_enabled": bool(payload.enabled),
                "custody_our_weeks": payload.our_weeks,
                "custody_away_label": payload.away_label.strip()[:CUSTODY_LABEL_MAX],
                "updated_at": utcnow(),
            }
        },
        upsert=True,
    )
    return await build_subscription(user["family_id"])


# -----------------------------------------------------------------------------
# Billing (RevenueCat)
# -----------------------------------------------------------------------------
# Events that mean the family currently holds (or regains) Premium, and events
# that mean access has actually ended. CANCELLATION only turns off auto-renew —
# access continues until EXPIRATION, so it does NOT downgrade.
RC_PREMIUM_EVENTS = {"INITIAL_PURCHASE", "RENEWAL", "UNCANCELLATION", "PRODUCT_CHANGE", "NON_RENEWING_PURCHASE"}
RC_DOWNGRADE_EVENTS = {"EXPIRATION"}


@app.post("/api/billing/revenuecat-webhook")
async def revenuecat_webhook(payload: dict, authorization: Optional[str] = Header(default=None)):
    """RevenueCat server-to-server events — the single source of truth for who
    has Premium once billing is live. app_user_id is our user_id (the app calls
    Purchases.logIn(user_id)), which maps to the family that gets the plan."""
    secret = os.environ.get("RC_WEBHOOK_SECRET", "")
    if not secret:
        raise HTTPException(status_code=503, detail="Billing not configured")
    supplied = (authorization or "").strip()
    if supplied.startswith("Bearer "):
        supplied = supplied[7:]
    if not secrets.compare_digest(supplied, secret):
        raise HTTPException(status_code=401, detail="Bad webhook signature")

    database = get_db()
    event = (payload or {}).get("event") or {}
    event_type = event.get("type", "")
    app_user_id = event.get("app_user_id") or ""

    user = await database["users"].find_one({"user_id": app_user_id}, {"_id": 0})
    if not user:
        # Unknown user (e.g. sandbox/anonymous id) — acknowledge so RC stops
        # retrying; nothing to update on our side.
        log.warning("RC webhook for unknown app_user_id=%s type=%s", app_user_id, event_type)
        return {"ok": True, "matched": False}

    product_id = (event.get("product_id") or "").lower()
    cycle = "yearly" if ("year" in product_id or "annual" in product_id) else "monthly"
    changes = {
        "rc_last_event": event_type,
        "rc_product_id": event.get("product_id"),
        "rc_event_at": utcnow(),
        "updated_at": utcnow(),
    }
    exp_ms = event.get("expiration_at_ms")
    if exp_ms:
        try:
            changes["rc_expires_at"] = datetime.fromtimestamp(int(exp_ms) / 1000, tz=timezone.utc)
        except (ValueError, TypeError, OSError):
            pass

    if event_type in RC_PREMIUM_EVENTS:
        changes["plan"] = "executive"
        changes["billing_cycle"] = cycle
    elif event_type in RC_DOWNGRADE_EVENTS:
        changes["plan"] = "village"
    # Other events (CANCELLATION, BILLING_ISSUE, TRANSFER, TEST) just record state.

    await database["families"].update_one(
        {"family_id": user["family_id"]}, {"$set": changes}
    )
    log.info("RC webhook applied: family=%s type=%s plan=%s", user["family_id"], event_type, changes.get("plan", "unchanged"))
    return {"ok": True, "matched": True}


async def _fetch_rc_subscriber(user_id: str, secret: str) -> dict:
    """Ask RevenueCat's REST API for a subscriber. Network only — the
    interpretation lives in rc_entitlement_state so it can be tested."""
    url = f"https://api.revenuecat.com/v1/subscribers/{urllib.parse.quote(user_id, safe='')}"

    def _request():
        req = urllib.request.Request(
            url, headers={"Authorization": f"Bearer {secret}"}, method="GET"
        )
        try:
            with urllib.request.urlopen(req, timeout=15) as resp:
                return json.loads(resp.read().decode("utf-8"))
        except urllib.error.HTTPError as e:
            raise HTTPException(status_code=502, detail=f"RevenueCat error {e.code}")
        except HTTPException:
            raise
        except Exception as e:
            raise HTTPException(status_code=502, detail=f"RevenueCat request failed: {e}")

    return await asyncio.to_thread(_request)


def rc_entitlement_state(subscriber: dict, now: datetime) -> tuple[bool, Optional[str]]:
    """Does this RevenueCat subscriber hold an active entitlement, and for
    which product? An entitlement with no expiry is lifetime; a malformed
    expiry is ignored rather than trusted."""
    entitlements = (subscriber or {}).get("entitlements") or {}
    for ent in entitlements.values():
        if not isinstance(ent, dict):
            continue
        product = ent.get("product_identifier")
        exp = ent.get("expires_date")
        if exp is None:
            return True, product
        try:
            exp_dt = ensure_aware_utc(parse_dt(str(exp)))
        except Exception:
            continue
        if exp_dt and exp_dt > now:
            return True, product
    return False, None


@app.post("/api/billing/reconcile")
async def reconcile_billing(user: dict = Depends(require_user)):
    """Ask RevenueCat directly what this user's subscription really is.

    Webhooks stay the normal source of truth, but a missed webhook used to be
    permanent: a paying family stayed 'free' until someone noticed. The app
    calls this quietly when the plans screen opens, so the server corrects
    itself from RevenueCat's answer — in both directions, with one guard:
    a downgrade is only applied when the plan verifiably came from billing
    (the family carries webhook state), so a manually granted plan is never
    silently revoked by a reconcile.
    """
    secret = os.environ.get("REVENUECAT_SECRET_KEY", "")
    if not secret:
        raise HTTPException(status_code=503, detail="Billing reconciliation not configured")

    database = get_db()
    data = await _fetch_rc_subscriber(user["user_id"], secret)
    active, product = rc_entitlement_state((data or {}).get("subscriber") or {}, utcnow())

    family = await get_family_doc(user["family_id"])
    changes = {"rc_reconciled_at": utcnow(), "updated_at": utcnow()}
    if active:
        product_id = (product or "").lower()
        changes["plan"] = "executive"
        changes["billing_cycle"] = "yearly" if ("year" in product_id or "annual" in product_id) else "monthly"
        changes["rc_product_id"] = product
    elif family.get("plan") == "executive" and family.get("rc_last_event"):
        changes["plan"] = "village"

    await database["families"].update_one(
        {"family_id": user["family_id"]}, {"$set": changes}
    )
    if changes.get("plan") and changes["plan"] != family.get("plan"):
        log.info(
            "RC reconcile corrected family=%s: %s -> %s",
            user["family_id"], family.get("plan"), changes["plan"],
        )

    sub = await build_subscription(user["family_id"])
    if is_admin_user(user):
        return apply_admin_subscription(sub)
    return sub


# -----------------------------------------------------------------------------
# Entitlements
# -----------------------------------------------------------------------------
@app.get("/api/subscription/entitlements")
async def get_entitlements(user=Depends(require_user)):
    database = get_db()
    sub = await build_subscription(user["family_id"])
    if is_admin_user(user):
        sub = apply_admin_subscription(sub)

    members_count = await database["family_members"].count_documents({"family_id": user["family_id"]})
    pending_invites = await database["family_invites"].count_documents(
        {
            "family_id": user["family_id"],
            "status": "pending",
            "expires_at": {"$gt": utcnow()},
        }
    )

    member_slots_used = members_count + pending_invites
    max_members = sub["limits"]["max_members"]

    return {
        "plan": sub["plan"],
        "admin_unlocked": bool(sub.get("admin_unlocked")),
        "members_count": members_count,
        "pending_invites": pending_invites,
        "member_slots_used": member_slots_used,
        "max_members": max_members,
        "can_invite": bool(sub.get("admin_unlocked")) or member_slots_used < max_members,
        "ai_scans_used": sub.get("ai_scans_used", 0),
        "ai_scans_limit": sub["limits"]["ai_scans_per_month"],
        "vault_bytes_used": sub.get("vault_bytes_used", 0),
        "vault_bytes_limit": sub["limits"]["vault_bytes"],
        "weekly_brief": sub["limits"].get("weekly_brief", False),
        "multi_property": sub["limits"].get("multi_property", False),
        "features": {
            "meal_planner": sub["limits"].get("meal_planner", False),
            "allowance": sub["limits"].get("allowance", False),
            "carpool": sub["limits"].get("carpool", False),
            "weekly_report": sub["limits"].get("weekly_report", False),
        },
    }



# -----------------------------------------------------------------------------
# Weekly brief
# -----------------------------------------------------------------------------
@app.post("/api/brief/weekly")
async def weekly_brief(user=Depends(require_user)):
    database = get_db()
    sub = await build_subscription(user["family_id"])
    if not is_admin_user(user) and not sub["limits"]["weekly_brief"]:
        plan_limit_error(
            feature="weekly_brief",
            current_plan=sub["plan"],
            message="Weekly Brief is available on Executive and Family Office plans.",
        )

    cards = []
    async for item in database["cards"].find({"family_id": user["family_id"], "status": "OPEN"}, {"_id": 0}):
        cards.append(item)

    if not cards:
        brief = "You have a clear runway this week. Use the space to reset routines, confirm calendars, and get ahead on one important family task."
        return {"brief": brief, "generated_at": iso(utcnow())}

    lines = []
    for c in cards[:12]:
        due = iso(c.get("due_date")) or "no due date"
        assignee = c.get("assignee") or "unassigned"
        lines.append(f"- {c['title']} | due: {due} | assignee: {assignee}")

    if not GOOGLE_API_KEY:
        brief = (
            "This week's household priorities are: "
            + "; ".join([c["title"] for c in cards[:5]])
            + ". Focus first on items with dates, assign open tasks clearly, and close one quick win today."
        )
        return {"brief": brief, "generated_at": iso(utcnow())}

    prompt = f"""
Write a warm, premium household chief-of-staff weekly brief in under 180 words.
Summarize priorities, likely bottlenecks, and one concrete action step.
Open items:
{chr(10).join(lines)}
""".strip()

    try:
        brief = await _gemini_text(prompt)
    except Exception as exc:  # noqa: BLE001 — degrade like the no-key path, never 500
        log.warning("weekly brief generation failed: %s", exc)
        brief = (
            "This week's household priorities are: "
            + "; ".join([c["title"] for c in cards[:5]])
            + ". Focus first on items with dates, assign open tasks clearly, and close one quick win today."
        )
    return {"brief": brief, "generated_at": iso(utcnow())}


# -----------------------------------------------------------------------------
# Vision
# -----------------------------------------------------------------------------
@app.post("/api/vision/extract")
async def vision_extract(payload: VisionIn, user=Depends(require_user)):
    database = get_db()
    sub = await build_subscription(user["family_id"])
    family = await get_family_doc(user["family_id"])

    if not is_admin_user(user) and family.get("ai_scans_used", 0) >= sub["limits"]["ai_scans_per_month"]:
        plan_limit_error(
            feature="ai_scans",
            current_plan=sub["plan"],
            limit=sub["limits"]["ai_scans_per_month"],
            used=family.get("ai_scans_used", 0),
            message="AI scan limit reached for this billing period.",
        )

    members = []
    async for m in database["family_members"].find({"family_id": user["family_id"]}, {"_id": 0}):
        members.append(m["name"])

    # What the family gets if the model says nothing usable: a card they can
    # edit, and no category. The old fallback guessed "School", which is how
    # a gas bill ended up filed with the permission slips — a guess presented
    # as an answer is worse than an honest blank.
    fallback = {
        "kind": "document",
        "type": "TASK",
        "title": "Scanned document",
        "description": "",
        "assignee": "",
        "due_date": None,
        "vault_category": "",
        "amount": None,
        "save_to_vault": True,
        "understood": False,
    }

    if not GOOGLE_API_KEY:
        return fallback

    extracted = None
    try:
        text = await _gemini_vision(
            "Read this household document.",
            payload.image_base64,
            system=build_document_scan_prompt(members),
        )
        parsed = extract_json(text)
        if parsed is None:
            raise UnsafeRecipe("unparseable")
        extracted = validate_document_scan(parsed, members)
    except UnsafeRecipe as exc:
        log.info("document scan rejected by safety gate: %s", exc.reason)
    except Exception as exc:
        log.warning("document scan failed: %s", exc)

    # The scan is charged whether or not it produced something, because it
    # cost what it cost. What must never happen is charging twice for one
    # photograph, which is why the recipe pass below runs inside this request
    # rather than sending the client off to /api/recipes/capture.
    if not is_admin_user(user):
        # Guarded so two concurrent scans cannot both push the counter past the
        # limit: the increment only lands while the family is still under it.
        # A request that raced past the last slot did its work already — it goes
        # through uncharged rather than over-counting.
        await database["families"].update_one(
            {"family_id": user["family_id"],
             "ai_scans_used": {"$lt": sub["limits"]["ai_scans_per_month"]}},
            {"$inc": {"ai_scans_used": 1}, "$set": {"updated_at": utcnow()}},
        )

    if extracted is None:
        return fallback

    result = {**extracted, "understood": True}

    if result["kind"] == "recipe":
        # A recipe is the one kind of document where filing it is not the
        # point — the ingredients are. Read it properly with the prompt that
        # already exists for photographed recipes; if that second pass fails,
        # this quietly stays a document, which is still a true description
        # of a photograph of a cookbook page.
        try:
            text = await _gemini_vision(
                "Read the recipe in this photo.",
                payload.image_base64,
                system=RECIPE_PHOTO_SYSTEM_PROMPT,
            )
            parsed = extract_json(text)
            if parsed is None:
                raise UnsafeRecipe("unparseable")
            result["recipe"] = validate_captured_recipe(parsed)
        except UnsafeRecipe as exc:
            log.info("recipe pass rejected by safety gate: %s", exc.reason)
            result["kind"] = "document"
        except Exception as exc:
            log.warning("recipe pass failed: %s", exc)
            result["kind"] = "document"

    return result



def _clean_json_text(text: str) -> str:
    value = (text or "").strip()

    if value.startswith("```json"):
        value = value.removeprefix("```json").strip()
    if value.startswith("```"):
        value = value.removeprefix("```").strip()
    if value.endswith("```"):
        value = value.removesuffix("```").strip()

    first = value.find("{")
    last = value.rfind("}")
    if first >= 0 and last > first:
        return value[first : last + 1]

    return value


def _safe_voice_draft(parsed: dict, fallback_transcript: str = "") -> dict:
    transcript = str(parsed.get("transcript") or fallback_transcript or "").strip()
    title = str(parsed.get("title") or "").strip()
    description = str(parsed.get("description") or "").strip()

    if not title:
        title = transcript[:70].strip() or "Voice task"
    if not description:
        description = transcript

    card_type = str(parsed.get("type") or "TASK").strip().upper()
    if card_type not in ("SIGN_SLIP", "RSVP", "TASK"):
        card_type = "TASK"

    due_date = parsed.get("due_date")
    if due_date in ("", "null", "None"):
        due_date = None

    return {
        "transcript": transcript,
        "type": card_type,
        "title": title,
        "description": description,
        "assignee": str(parsed.get("assignee") or "").strip(),
        "due_date": due_date,
    }


async def _voice_to_draft(audio_bytes: bytes, mime_type: str, members: list[str]) -> dict:
    if not GOOGLE_API_KEY or not genai:
        raise HTTPException(
            status_code=501,
            detail="Voice transcription requires GOOGLE_API_KEY in Railway.",
        )

    allowed_members = ", ".join(members) if members else ""

    prompt = f"""
You are Ahenora, a premium family chief-of-staff assistant.

Listen to the audio and return JSON only with these keys:
transcript, type, title, description, assignee, due_date

Rules:
- transcript: accurate transcription of the user's speech.
- type: one of SIGN_SLIP, RSVP, TASK.
- title: concise action title, max 70 characters.
- description: practical detail from the audio.
- assignee: one exact name from this list if clearly mentioned or obvious: {allowed_members}
- assignee may be empty string if unclear.
- due_date: ISO 8601 string if the audio clearly mentions a date/time, otherwise null.
- Return valid JSON only. No markdown.
""".strip()

    voice_system = "You convert spoken household instructions into structured task/card JSON."

    try:
        text = await _gemini_generate(
            [prompt, genai_types.Part.from_bytes(
                data=audio_bytes, mime_type=mime_type or "audio/aac")],
            system=voice_system,
        )
    except Exception as first_error:
        # Inline audio has a request-size ceiling; past it the file has to be
        # uploaded first. Falling back rather than failing keeps long voice
        # notes working.
        client = _gemini_client()
        if not client:
            raise HTTPException(
                status_code=500,
                detail=f"Voice transcription failed: {first_error}",
            )

        suffix = ".m4a"
        if "ogg" in (mime_type or ""):
            suffix = ".ogg"
        elif "webm" in (mime_type or ""):
            suffix = ".webm"
        elif "mpeg" in (mime_type or "") or "mp3" in (mime_type or ""):
            suffix = ".mp3"
        elif "wav" in (mime_type or ""):
            suffix = ".wav"

        tmp_path = None
        try:
            with tempfile.NamedTemporaryFile(delete=False, suffix=suffix) as tmp:
                tmp.write(audio_bytes)
                tmp_path = tmp.name

            uploaded = await client.aio.files.upload(
                file=tmp_path,
                config={"mime_type": mime_type} if mime_type else None,
            )
            # Routed through _gemini_generate so the upload path gets the same
            # model fallback as everything else — previously it was pinned to
            # one model and went dark on its own when that model was retired.
            text = await _gemini_generate([prompt, uploaded], system=voice_system)
        except Exception as second_error:
            raise HTTPException(
                status_code=500,
                detail=f"Voice transcription failed: {second_error}",
            )
        finally:
            if tmp_path:
                try:
                    os.remove(tmp_path)
                except Exception:
                    pass

    try:
        parsed = json.loads(_clean_json_text(text))
        if not isinstance(parsed, dict):
            raise ValueError("Model response was not an object")
        return _safe_voice_draft(parsed, fallback_transcript=text)
    except Exception:
        return _safe_voice_draft({"transcript": text, "description": text}, fallback_transcript=text)



# -----------------------------------------------------------------------------
# Voice placeholder
# -----------------------------------------------------------------------------
@app.post("/api/voice/transcribe")
async def voice_transcribe(audio: UploadFile = File(...), user=Depends(require_user)):
    database = get_db()

    audio_bytes = await audio.read()
    if not audio_bytes or len(audio_bytes) < 500:
        raise HTTPException(status_code=400, detail="Recording is too short")

    if len(audio_bytes) > MAX_VOICE_AUDIO_BYTES:
        raise HTTPException(status_code=413, detail="Recording is too large")

    # A voice note is a billed multimodal Gemini call, so it counts against the
    # monthly AI allowance like every other scan. This was the one paid AI
    # route that neither gated nor metered — a family out of scans could
    # transcribe without limit. Charged before the call, atomically.
    await charge_ai_scan(user)

    members = []
    async for member in database["family_members"].find(
        {"family_id": user["family_id"]},
        {"_id": 0, "name": 1},
    ):
        if member.get("name"):
            members.append(member["name"])

    mime_type = audio.content_type or "audio/aac"
    return await _voice_to_draft(audio_bytes, mime_type, members)


# -----------------------------------------------------------------------------
# Handoff Notes
# -----------------------------------------------------------------------------

@app.get("/api/handoff-notes")
async def list_handoff_notes(user=Depends(require_user)):
    database = get_db()
    rows = []
    async for note in database["handoff_notes"].find(
        {"family_id": user["family_id"]},
        {"_id": 0},
    ).sort("created_at", -1).limit(50):
        rows.append(public_handoff_note(note))
    return rows


@app.post("/api/handoff-notes")
async def create_handoff_note(payload: HandoffNoteIn, user=Depends(require_user)):
    database = get_db()
    member_name = None
    if payload.member_id:
        member = await database["family_members"].find_one(
            {"family_id": user["family_id"], "member_id": payload.member_id},
            {"_id": 0, "name": 1},
        )
        if member:
            member_name = member["name"]
    doc = {
        "note_id": new_id("note"),
        "family_id": user["family_id"],
        "member_id": payload.member_id,
        "member_name": member_name,
        "text": payload.text.strip(),
        "author_name": user.get("name", ""),
        "author_user_id": user["user_id"],
        "created_at": utcnow(),
    }
    await database["handoff_notes"].insert_one(doc)
    try:
        who = user.get("name") or "A co-parent"
        await send_coparent_alert(
            user["family_id"], f"{who} left a note", doc["text"], "handoff_note",
            created_by_user_id=user["user_id"],
        )
    except Exception as e:
        log.warning("handoff note alert failed: %s", e)
    return public_handoff_note(doc)


@app.delete("/api/handoff-notes/{note_id}")
async def delete_handoff_note(note_id: str, user=Depends(require_user)):
    database = get_db()
    result = await database["handoff_notes"].delete_one(
        {"note_id": note_id, "family_id": user["family_id"]}
    )
    if result.deleted_count == 0:
        raise HTTPException(404, "Note not found")
    return {"ok": True}


# -----------------------------------------------------------------------------
# Shopping List
# -----------------------------------------------------------------------------

SHOPPING_CATEGORIES = [
    "Produce", "Dairy", "Meat", "Bakery", "Frozen",
    "Pantry", "Drinks", "Snacks", "Baby", "Household",
    "Health", "School", "Other",
]


@app.get("/api/shopping")
async def list_shopping(user=Depends(require_user)):
    database = get_db()
    rows = []
    async for item in database["shopping_list"].find(
        {"family_id": user["family_id"]},
        {"_id": 0},
    ).sort("created_at", -1):
        rows.append(public_shopping_item(item))
    return rows


@app.post("/api/shopping")
async def add_shopping_item(payload: ShoppingItemIn, user=Depends(require_user)):
    database = get_db()
    doc = {
        "item_id": new_id("shop"),
        "family_id": user["family_id"],
        "name": payload.name.strip(),
        "category": payload.category if payload.category in SHOPPING_CATEGORIES else "Other",
        "checked": False,
        "added_by": user.get("name", ""),
        "created_at": utcnow(),
    }
    await database["shopping_list"].insert_one(doc)
    return public_shopping_item(doc)


@app.patch("/api/shopping/{item_id}")
async def update_shopping_item(item_id: str, payload: ShoppingItemPatchIn, user=Depends(require_user)):
    database = get_db()
    updates = {}
    if payload.checked is not None:
        updates["checked"] = payload.checked
    if payload.name is not None:
        updates["name"] = payload.name.strip()
    if payload.category is not None:
        updates["category"] = payload.category
    if not updates:
        raise HTTPException(400, "Nothing to update")
    result = await database["shopping_list"].update_one(
        {"item_id": item_id, "family_id": user["family_id"]},
        {"$set": updates},
    )
    if result.matched_count == 0:
        raise HTTPException(404, "Item not found")
    doc = await database["shopping_list"].find_one({"item_id": item_id}, {"_id": 0})
    return public_shopping_item(doc)


# Registered before /{item_id} so "all" isn't captured as an item id.
@app.delete("/api/shopping/all")
async def clear_all_shopping(user=Depends(require_user)):
    database = get_db()
    # Archive the list before wiping so it can be restored from history.
    items = await database["shopping_list"].find(
        {"family_id": user["family_id"]}, {"_id": 0, "name": 1}
    ).to_list(500)
    names = [i.get("name", "") for i in items if i.get("name")]
    # "Clear all" means all: older archives go too, so the restore banner
    # never lingers with lists from weeks ago. The list being cleared right
    # now is still archived, because an accidental clear must be undoable.
    await database["shopping_history"].delete_many({"family_id": user["family_id"]})
    if names:
        await database["shopping_history"].insert_one({
            "history_id": new_id("shist"),
            "family_id": user["family_id"],
            "items": names,
            "created_at": utcnow(),
        })
    result = await database["shopping_list"].delete_many({"family_id": user["family_id"]})
    if names:
        await log_activity(database, user, "list_cleared", "", len(names))
    return {"deleted": result.deleted_count}


@app.delete("/api/shopping/{item_id}")
async def delete_shopping_item(item_id: str, user=Depends(require_user)):
    database = get_db()
    result = await database["shopping_list"].delete_one(
        {"item_id": item_id, "family_id": user["family_id"]}
    )
    if result.deleted_count == 0:
        raise HTTPException(404, "Item not found")
    return {"ok": True}


def public_shopping_history(h: dict) -> dict:
    return {"history_id": h["history_id"], "items": h.get("items", []), "created_at": iso(h["created_at"])}


@app.delete("/api/shopping")
async def clear_checked_shopping(user=Depends(require_user)):
    database = get_db()
    # Archive the finished trip (checked items) before clearing, so it can be reused.
    checked = await database["shopping_list"].find(
        {"family_id": user["family_id"], "checked": True}, {"_id": 0, "name": 1}
    ).to_list(500)
    names = [c.get("name", "") for c in checked if c.get("name")]
    if names:
        await database["shopping_history"].insert_one({
            "history_id": new_id("shist"),
            "family_id": user["family_id"],
            "items": names,
            "created_at": utcnow(),
        })
    result = await database["shopping_list"].delete_many(
        {"family_id": user["family_id"], "checked": True}
    )
    return {"deleted": result.deleted_count}


@app.get("/api/shopping/history")
async def list_shopping_history(user=Depends(require_user)):
    database = get_db()
    rows = await database["shopping_history"].find(
        {"family_id": user["family_id"]}, {"_id": 0}
    ).sort("created_at", -1).to_list(30)
    return [public_shopping_history(h) for h in rows]


@app.post("/api/shopping/history/{history_id}/reuse")
async def reuse_shopping_history(history_id: str, user=Depends(require_user)):
    database = get_db()
    h = await database["shopping_history"].find_one(
        {"history_id": history_id, "family_id": user["family_id"]}, {"_id": 0}
    )
    if not h:
        raise HTTPException(404, "Not found")
    added = 0
    for name in h.get("items", []):
        if not name:
            continue
        await database["shopping_list"].insert_one({
            "item_id": new_id("shop"),
            "family_id": user["family_id"],
            "name": name,
            "category": "Other",
            "checked": False,
            "added_by": user.get("name", ""),
            "created_at": utcnow(),
        })
        added += 1
    return {"ok": True, "added": added}


@app.delete("/api/shopping/history")
async def clear_shopping_history(user=Depends(require_user)):
    """Wipe every archived list for the family — the restore banner's
    "no thanks, ever" for households that never reuse an old list."""
    database = get_db()
    result = await database["shopping_history"].delete_many({"family_id": user["family_id"]})
    return {"ok": True, "deleted": result.deleted_count}


@app.delete("/api/shopping/history/{history_id}")
async def delete_shopping_history(history_id: str, user=Depends(require_user)):
    database = get_db()
    await database["shopping_history"].delete_one(
        {"history_id": history_id, "family_id": user["family_id"]}
    )
    return {"ok": True}


@app.post("/api/shopping/scan")
async def scan_shopping_list(
    payload: dict = Body(...),
    user: dict = Depends(require_user),
    database=Depends(get_db),
):
    """Read a photo of a paper shopping list into items.

    Returns candidates only — nothing is added here. The app shows them in a
    ticked preview (unsure reads come unticked) and commits the selection
    through the ordinary bulk add, so a misread never lands silently.
    """
    image_b64 = str(payload.get("image_base64") or "")
    if "," in image_b64:
        image_b64 = image_b64.split(",")[-1]
    if len(image_b64) < 100:
        raise HTTPException(400, "No photo attached.")
    if len(image_b64) > MAX_IMAGE_B64_CHARS:
        raise HTTPException(413, "That photo is too large.")

    if not GOOGLE_API_KEY:
        raise HTTPException(503, "Photo scanning is unavailable right now.")

    sub = await build_subscription(user["family_id"])
    family = await get_family_doc(user["family_id"])
    if not is_admin_user(user) and family.get("ai_scans_used", 0) >= sub["limits"]["ai_scans_per_month"]:
        plan_limit_error(
            feature="ai_scans",
            current_plan=sub["plan"],
            limit=sub["limits"]["ai_scans_per_month"],
            used=family.get("ai_scans_used", 0),
            message="AI scan limit reached for this billing period.",
        )

    try:
        text = await _gemini_vision(
            "Read the shopping list in this photo.",
            image_b64,
            system=SHOPPING_SCAN_SYSTEM_PROMPT,
            fast=True,
        )
        parsed = extract_json(text)
        if parsed is None:
            raise UnsafeRecipe("unparseable")
        items = validate_shopping_scan(parsed)
    except UnsafeRecipe as exc:
        log.info("shopping scan rejected by safety gate: %s", exc.reason)
        raise HTTPException(422, "We could not read a shopping list in that photo.")
    except Exception as exc:
        log.warning("shopping scan failed: %s", exc)
        raise HTTPException(502, "We could not read a shopping list in that photo.")

    if not is_admin_user(user):
        # Guarded so two concurrent scans cannot both push the counter past the
        # limit: the increment only lands while the family is still under it.
        # A request that raced past the last slot did its work already — it goes
        # through uncharged rather than over-counting.
        await database["families"].update_one(
            {"family_id": user["family_id"],
             "ai_scans_used": {"$lt": sub["limits"]["ai_scans_per_month"]}},
            {"$inc": {"ai_scans_used": 1}, "$set": {"updated_at": utcnow()}},
        )

    return {"items": items}


class BulkShoppingIn(BaseModel):
    # One insert per name, so an uncapped list is one request that drives
    # arbitrarily many writes. A household shopping list is dozens of items,
    # not thousands; 200 is well clear of any real list and a firm ceiling.
    names: list[str] = Field(max_length=200)
    # Aisle per name, same order. The multilingual matching lives in the app,
    # so the client classifies and the server stores. Short, missing or
    # unrecognised entries fall back to "Other".
    categories: list[str] = Field(default=[], max_length=200)


@app.post("/api/shopping/bulk")
async def bulk_add_shopping(body: BulkShoppingIn, user=Depends(require_user)):
    """Add several items at once — used to restore selected items from a past
    list. Skips blanks and anything already on the current (unchecked) list."""
    database = get_db()
    existing = await database["shopping_list"].find(
        {"family_id": user["family_id"], "checked": False}, {"_id": 0, "name": 1}
    ).to_list(500)
    have = {(e.get("name") or "").strip().lower() for e in existing}
    added = 0
    for index, raw in enumerate(body.names):
        name = (raw or "").strip()
        if not name or name.lower() in have:
            continue
        have.add(name.lower())
        supplied = body.categories[index] if index < len(body.categories) else None
        await database["shopping_list"].insert_one({
            "item_id": new_id("shop"),
            "family_id": user["family_id"],
            "name": name,
            "category": supplied if supplied in SHOPPING_CATEGORIES else "Other",
            "checked": False,
            "added_by": user.get("name", ""),
            "created_at": utcnow(),
        })
        added += 1
    return {"ok": True, "added": added}


# -----------------------------------------------------------------------------
# Expense Tracking
# -----------------------------------------------------------------------------

@app.get("/api/expenses")
async def list_expenses(
    user=Depends(require_full_member),
    days: int = Query(default=30, ge=1, le=365),
):
    database = get_db()
    since = utcnow() - timedelta(days=days)
    rows = []
    async for exp in database["expenses"].find(
        {"family_id": user["family_id"], "created_at": {"$gte": since}},
        {"_id": 0},
    ).sort("created_at", -1):
        rows.append(public_expense(exp))
    return rows


@app.get("/api/expenses/summary")
async def expense_summary(
    user=Depends(require_full_member),
    days: int = Query(default=30, ge=1, le=365),
):
    database = get_db()
    since = utcnow() - timedelta(days=days)
    by_user: dict[str, float] = {}
    by_category: dict[str, float] = {}
    total = 0.0
    async for exp in database["expenses"].find(
        {"family_id": user["family_id"], "created_at": {"$gte": since}},
        {"_id": 0},
    ):
        amt = exp.get("amount", 0)
        total += amt
        name = exp.get("paid_by_name", "Unknown")
        by_user[name] = by_user.get(name, 0) + amt
        cat = exp.get("category", "General")
        by_category[cat] = by_category.get(cat, 0) + amt
    return {
        "total": round(total, 2),
        "by_person": {k: round(v, 2) for k, v in by_user.items()},
        "by_category": {k: round(v, 2) for k, v in by_category.items()},
        "days": days,
    }


@app.get("/api/expenses/overview")
async def expense_overview(
    user=Depends(require_full_member),
    months: int = Query(default=6, ge=1, le=24),
    category: Optional[str] = Query(default=None),
):
    """Spending by calendar month and by shop — the House expenses view.

    Three judgements are baked in here rather than left to the screen, so every
    client tells the family the same true story:

    1. Months are counted from the date on the receipt, not the day the expense
       was typed in. People photograph a week of tickets on a Sunday evening.
       The date is date-only, so a month total cannot shift because the family
       travelled to another timezone.

    2. Every total carries the number of receipts behind it. A month where four
       shops were forgotten looks like a month of admirable restraint, and a
       number without its coverage is how an app lies without meaning to.

    3. A month still running is never compared to a finished one, and "usual"
       means the average of the three previous COMPLETE months — not last month
       alone. One month against one month is noise: a birthday, a visitor, a
       bulk buy of nappies.
    """
    database = get_db()
    today = utcnow().replace(hour=0, minute=0, second=0, microsecond=0)

    def month_key(when: datetime) -> str:
        return f"{when.year:04d}-{when.month:02d}"

    # The window of month keys we care about, oldest first, ending this month.
    keys: list = []
    year, month = today.year, today.month
    for _ in range(months):
        keys.append(f"{year:04d}-{month:02d}")
        month -= 1
        if month == 0:
            year, month = year - 1, 12
    keys.reverse()
    this_month = keys[-1]
    wanted = set(keys)

    # Filtered in Python rather than in the query: a household has hundreds of
    # expenses, not millions, and doing the date arithmetic here keeps the
    # spent_on fallback (older rows have no such field) in one place.
    buckets: dict = {k: {"total": 0.0, "count": 0, "shops": {}, "people": {}} for k in keys}
    range_shops: dict = {}
    range_total, range_count = 0.0, 0

    async for exp in database["expenses"].find({"family_id": user["family_id"]}, {"_id": 0}):
        if category and (exp.get("category") or "General") != category:
            continue
        when = exp.get("spent_on") or exp.get("created_at")
        if not isinstance(when, datetime):
            continue
        key = month_key(when)
        if key not in wanted:
            continue
        amount = float(exp.get("amount") or 0)
        shop = exp.get("merchant") or exp.get("description") or "Unknown"

        bucket = buckets[key]
        bucket["total"] += amount
        bucket["count"] += 1
        # Who paid — the older job this screen has always done, kept alongside
        # the new one so co-parents can still settle up without a second screen.
        payer = exp.get("paid_by_name") or "—"
        bucket["people"][payer] = bucket["people"].get(payer, 0.0) + amount
        seen = bucket["shops"].setdefault(shop, {"total": 0.0, "visits": 0})
        seen["total"] += amount
        seen["visits"] += 1

        entry = range_shops.setdefault(shop, {"total": 0.0, "visits": 0})
        entry["total"] += amount
        entry["visits"] += 1
        range_total += amount
        range_count += 1

    def shop_rows(shops: dict) -> list:
        rows = [{"merchant": name,
                 "total": round(v["total"], 2),
                 "visits": v["visits"],
                 "average": round(v["total"] / v["visits"], 2) if v["visits"] else 0.0}
                for name, v in shops.items()]
        rows.sort(key=lambda r: r["total"], reverse=True)
        return rows

    month_rows = [{"month": k,
                   "total": round(buckets[k]["total"], 2),
                   "count": buckets[k]["count"],
                   "complete": k != this_month,
                   "by_merchant": shop_rows(buckets[k]["shops"]),
                   "by_person": {n: round(v, 2) for n, v in
                                 sorted(buckets[k]["people"].items(),
                                        key=lambda kv: kv[1], reverse=True)}}
                  for k in keys]

    # The comparison, offered only when it can be honest: the newest finished
    # month, against the three finished months before it.
    comparison = None
    finished = [m for m in month_rows if m["complete"] and m["count"] > 0]
    if finished:
        latest = finished[-1]
        prior = [m for m in month_rows if m["complete"] and m["count"] > 0
                 and m["month"] < latest["month"]][-3:]
        if len(prior) == 3:
            usual = sum(m["total"] for m in prior) / 3
            comparison = {
                "month": latest["month"],
                "total": latest["total"],
                "usual": round(usual, 2),
                "difference": round(latest["total"] - usual, 2),
                "basis_months": [m["month"] for m in prior],
            }

    current = next(m for m in month_rows if m["month"] == this_month)
    return {
        "category": category,
        "months": month_rows,
        "current": current,
        "days_into_month": today.day,
        "comparison": comparison,
        "range": {"months": months,
                  "total": round(range_total, 2),
                  "count": range_count,
                  "by_merchant": shop_rows(range_shops)},
    }


@app.post("/api/expenses")
async def add_expense(payload: ExpenseIn, user=Depends(require_full_member)):
    database = get_db()
    child_name = None
    if payload.child_member_id:
        member = await database["family_members"].find_one(
            {"family_id": user["family_id"], "member_id": payload.child_member_id},
            {"_id": 0, "name": 1},
        )
        if member:
            child_name = member["name"]
    merchant = tidy_merchant(payload.merchant)
    # A shop IS a description for a grocery run; making someone type "Aldi" twice
    # is the kind of small friction that stops a habit forming.
    description = sanitize_message_text(payload.description or "", 200).replace("\n", " ").strip()
    if not description:
        description = merchant or "Expense"

    now = utcnow()
    doc = {
        "expense_id": new_id("exp"),
        "family_id": user["family_id"],
        "description": description,
        "amount": round(payload.amount, 2),
        "category": payload.category,
        "merchant": merchant,
        "spent_on": parse_spent_on(payload.spent_on, now),
        "child_member_id": payload.child_member_id,
        "child_name": child_name,
        "paid_by_name": user.get("name", ""),
        "paid_by_user_id": user["user_id"],
        "created_at": now,
    }
    await database["expenses"].insert_one(doc)
    return public_expense(doc)


@app.delete("/api/expenses/{expense_id}")
async def delete_expense(expense_id: str, user=Depends(require_full_member)):
    database = get_db()
    result = await database["expenses"].delete_one(
        {"expense_id": expense_id, "family_id": user["family_id"]}
    )
    if result.deleted_count == 0:
        raise HTTPException(404, "Expense not found")
    return {"ok": True}


# -----------------------------------------------------------------------------
# Recurring Templates
# -----------------------------------------------------------------------------

@app.get("/api/templates")
async def list_templates(user=Depends(require_user)):
    database = get_db()
    rows = []
    async for tmpl in database["templates"].find(
        {"family_id": user["family_id"]},
        {"_id": 0},
    ).sort("created_at", -1):
        rows.append(public_template(tmpl))
    return rows


@app.post("/api/templates")
async def create_template(payload: TemplateIn, user=Depends(require_user)):
    database = get_db()
    doc = {
        "template_id": new_id("tmpl"),
        "family_id": user["family_id"],
        "title": payload.title.strip(),
        "description": (payload.description or "").strip() or None,
        "recurrence": payload.recurrence if payload.recurrence in ("daily", "weekly", "monthly") else "daily",
        "time_of_day": payload.time_of_day,
        "assignee": payload.assignee,
        "enabled": True,
        "created_at": utcnow(),
    }
    await database["templates"].insert_one(doc)
    return public_template(doc)


@app.patch("/api/templates/{template_id}")
async def update_template(template_id: str, user=Depends(require_user)):
    database = get_db()
    tmpl = await database["templates"].find_one(
        {"template_id": template_id, "family_id": user["family_id"]},
        {"_id": 0},
    )
    if not tmpl:
        raise HTTPException(404, "Template not found")
    new_enabled = not tmpl.get("enabled", True)
    await database["templates"].update_one(
        {"template_id": template_id},
        {"$set": {"enabled": new_enabled}},
    )
    tmpl["enabled"] = new_enabled
    return public_template(tmpl)


@app.delete("/api/templates/{template_id}")
async def delete_template(template_id: str, user=Depends(require_user)):
    database = get_db()
    result = await database["templates"].delete_one(
        {"template_id": template_id, "family_id": user["family_id"]}
    )
    if result.deleted_count == 0:
        raise HTTPException(404, "Template not found")
    return {"ok": True}


@app.post("/api/templates/{template_id}/generate")
async def generate_from_template(template_id: str, user=Depends(require_user)):
    database = get_db()
    tmpl = await database["templates"].find_one(
        {"template_id": template_id, "family_id": user["family_id"]},
        {"_id": 0},
    )
    if not tmpl:
        raise HTTPException(404, "Template not found")
    due = utcnow()
    if tmpl.get("time_of_day"):
        try:
            parts = tmpl["time_of_day"].split(":")
            due = due.replace(hour=int(parts[0]), minute=int(parts[1]) if len(parts) > 1 else 0, second=0, microsecond=0)
        except (ValueError, IndexError):
            pass
    card = {
        "card_id": new_id("card"),
        "family_id": user["family_id"],
        "type": "TASK",
        "title": tmpl["title"],
        "description": tmpl.get("description"),
        "assignee": tmpl.get("assignee"),
        "due_date": due,
        "status": "OPEN",
        "source": "MANUAL",
        "image_base64": None,
        "recurrence": tmpl.get("recurrence", "none"),
        "reminder_minutes": 60,
        "created_at": utcnow(),
        "completed_at": None,
        "created_by_user_id": user["user_id"],
        "shared": False,
    }
    await database["cards"].insert_one(card)
    return public_card(card)


# -----------------------------------------------------------------------------
# Morning Routines
# -----------------------------------------------------------------------------
@app.get("/api/routines")
async def list_routines(user: dict = Depends(require_user), database=Depends(get_db)):
    rows = await database["routines"].find(
        {"family_id": user["family_id"]}, {"_id": 0}
    ).sort("created_at", -1).to_list(100)
    return [public_routine(r) for r in rows]


@app.post("/api/routines")
async def create_routine(body: RoutineIn, user: dict = Depends(require_user), database=Depends(get_db)):
    routine = {
        "routine_id": new_id("rtn"),
        "family_id": user["family_id"],
        "name": body.name,
        "steps": body.steps,
        "member_id": body.member_id,
        "star_reward": max(0, int(body.star_reward or 0)),
        "created_at": utcnow(),
    }
    await database["routines"].insert_one(routine)
    return public_routine(routine)


@app.patch("/api/routines/{routine_id}")
async def update_routine(routine_id: str, body: RoutinePatchIn, user: dict = Depends(require_user), database=Depends(get_db)):
    updates = {k: v for k, v in body.dict(exclude_unset=True).items() if v is not None}
    if not updates:
        raise HTTPException(400, "No updates provided")
    await database["routines"].update_one(
        {"routine_id": routine_id, "family_id": user["family_id"]},
        {"$set": updates},
    )
    row = await database["routines"].find_one(
        {"routine_id": routine_id, "family_id": user["family_id"]}, {"_id": 0}
    )
    if not row:
        raise HTTPException(404, "Routine not found")
    return public_routine(row)


@app.delete("/api/routines/{routine_id}")
async def delete_routine(routine_id: str, user: dict = Depends(require_user), database=Depends(get_db)):
    result = await database["routines"].delete_one(
        {"routine_id": routine_id, "family_id": user["family_id"]}
    )
    if result.deleted_count == 0:
        raise HTTPException(404, "Routine not found")
    return {"ok": True}


@app.post("/api/routines/{routine_id}/log")
async def log_routine_completion(routine_id: str, user: dict = Depends(require_user), database=Depends(get_db)):
    routine = await database["routines"].find_one(
        {"routine_id": routine_id, "family_id": user["family_id"]}, {"_id": 0}
    )
    if not routine:
        raise HTTPException(404, "Routine not found")
    log_entry = {
        "log_id": new_id("rlog"),
        "routine_id": routine_id,
        "family_id": user["family_id"],
        "member_id": routine.get("member_id"),
        "completed_at": utcnow(),
        "steps_count": len(routine.get("steps", [])),
    }
    await database["routine_logs"].insert_one(log_entry)

    # Finishing a routine earns stars. Without this the child got a toast and
    # nothing else, so routines sat outside the economy they were meant to feed.
    txn = await award_stars_to_member(
        database,
        user["family_id"],
        routine.get("member_id"),
        int(routine.get("star_reward", 0) or 0),
        routine.get("name") or "Routine complete",
        user,
    )
    return {
        "ok": True,
        "log_id": log_entry["log_id"],
        "stars_awarded": txn["delta"] if txn else 0,
        "member_id": routine.get("member_id"),
    }


# -----------------------------------------------------------------------------
# Meal Planner
# -----------------------------------------------------------------------------
DAYS_OF_WEEK = ["monday", "tuesday", "wednesday", "thursday", "friday", "saturday", "sunday"]
MEAL_TYPES = ["breakfast", "lunch", "dinner", "snack"]


@app.get("/api/meals")
async def list_meals(user: dict = Depends(require_user), database=Depends(get_db)):
    rows = await database["meals"].find(
        {"family_id": user["family_id"]}, {"_id": 0}
    ).to_list(200)
    return [public_meal(m) for m in rows]


@app.post("/api/meals")
async def create_meal(body: MealPlanIn, user: dict = Depends(require_user), database=Depends(get_db)):
    await require_feature(user, "meal_planner")
    day = body.day.lower()
    if day not in DAYS_OF_WEEK:
        raise HTTPException(400, f"Day must be one of {DAYS_OF_WEEK}")
    meal = {
        "meal_id": new_id("meal"),
        "family_id": user["family_id"],
        "day": day,
        "meal_type": body.meal_type.lower(),
        "title": body.title,
        "ingredients": body.ingredients,
        "notes": body.notes,
        "recipe_id": body.recipe_id,
        "created_at": utcnow(),
    }
    await database["meals"].insert_one(meal)
    return public_meal(meal)


# Registered before /{meal_id} so "all" isn't captured as a meal id.
@app.delete("/api/meals/all")
async def clear_all_meals(user: dict = Depends(require_user), database=Depends(get_db)):
    result = await database["meals"].delete_many({"family_id": user["family_id"]})
    return {"deleted": result.deleted_count}


@app.delete("/api/meals/{meal_id}")
async def delete_meal(meal_id: str, user: dict = Depends(require_user), database=Depends(get_db)):
    result = await database["meals"].delete_one(
        {"meal_id": meal_id, "family_id": user["family_id"]}
    )
    if result.deleted_count == 0:
        raise HTTPException(404, "Meal not found")
    return {"ok": True}


class SavedPlanIn(BaseModel):
    name: str


def public_saved_plan(pl: dict) -> dict:
    return {"plan_id": pl["plan_id"], "name": pl.get("name", ""), "meals": pl.get("meals", []), "created_at": iso(pl["created_at"])}


@app.post("/api/meals/save")
async def save_meal_plan(body: SavedPlanIn, user: dict = Depends(require_user), database=Depends(get_db)):
    await require_feature(user, "meal_planner")
    meals = await database["meals"].find(
        {"family_id": user["family_id"]}, {"_id": 0, "day": 1, "title": 1, "ingredients": 1, "recipe_id": 1}
    ).to_list(200)
    snapshot = [
        {
            "day": m.get("day"),
            "title": m.get("title", ""),
            "ingredients": m.get("ingredients", []),
            "recipe_id": m.get("recipe_id"),
        }
        for m in meals
    ]
    plan = {
        "plan_id": new_id("mplan"),
        "family_id": user["family_id"],
        "name": (body.name or "Saved plan").strip()[:60] or "Saved plan",
        "meals": snapshot,
        "created_at": utcnow(),
    }
    await database["meal_plans_saved"].insert_one(plan)
    return public_saved_plan(plan)


@app.get("/api/meals/saved")
async def list_saved_plans(user: dict = Depends(require_user), database=Depends(get_db)):
    rows = await database["meal_plans_saved"].find(
        {"family_id": user["family_id"]}, {"_id": 0}
    ).sort("created_at", -1).to_list(30)
    return [public_saved_plan(p) for p in rows]


@app.post("/api/meals/saved/{plan_id}/reuse")
async def reuse_saved_plan(plan_id: str, user: dict = Depends(require_user), database=Depends(get_db)):
    await require_feature(user, "meal_planner")
    pl = await database["meal_plans_saved"].find_one(
        {"plan_id": plan_id, "family_id": user["family_id"]}, {"_id": 0}
    )
    if not pl:
        raise HTTPException(404, "Not found")
    added = 0
    for m in pl.get("meals", []):
        day = (m.get("day") or "monday").lower()
        if day not in DAYS_OF_WEEK:
            continue
        await database["meals"].insert_one({
            "meal_id": new_id("meal"),
            "family_id": user["family_id"],
            "day": day,
            "meal_type": "dinner",
            "title": m.get("title", ""),
            "ingredients": m.get("ingredients", []),
            "notes": None,
            "recipe_id": m.get("recipe_id"),
            "created_at": utcnow(),
        })
        added += 1
    return {"ok": True, "added": added}


@app.delete("/api/meals/saved/{plan_id}")
async def delete_saved_plan(plan_id: str, user: dict = Depends(require_user), database=Depends(get_db)):
    await database["meal_plans_saved"].delete_one(
        {"plan_id": plan_id, "family_id": user["family_id"]}
    )
    return {"ok": True}


RECIPE_LANGUAGE_NAMES = {
    "en": "English",
    "es": "Spanish",
    "fr": "French",
    "de": "German",
}


def normalize_diet(value) -> str:
    """Only "vegetarian" is a real diet today; everything else is "no diet".

    A closed set rather than free text — the diet only ever picks a fixed prompt
    clause we wrote, never carries user words into the model.
    """
    return VEGETARIAN if str(value or "").strip().lower() == VEGETARIAN else ""


def recipe_slot(diet: str) -> str:
    """Where a generated recipe caches on the meal doc.

    Vegetarian recipes live in their own slot so a veg rewrite never returns the
    cached omnivore version, and a plain "Cook it" never returns the veg one.
    """
    return f"ai_recipe_{diet}" if diet else "ai_recipe"


@app.post("/api/meals/{meal_id}/recipe")
async def generate_meal_recipe(
    meal_id: str,
    lang: str = "en",
    diet: str = "",
    variant: int = 0,
    user: dict = Depends(require_user),
    database=Depends(get_db),
):
    """Write a cooking method for a meal the family typed themselves.

    The curated library covers ~50 dishes and ships in the app; this covers
    everything else. Results are cached on the meal per language (and per diet),
    so opening the same recipe again costs nothing and returns instantly.

    `diet` ("vegetarian") rewrites the ingredients and steps into that diet
    rather than commenting on it, and caches in its own slot so it never clashes
    with the plain version. `variant` (>0) asks for a different take on the same
    dish and deliberately skips the cache so "Different recipe" is always fresh.
    """
    await require_feature(user, "meal_planner")

    language = lang if lang in RECIPE_LANGUAGE_NAMES else "en"

    meal = await database["meals"].find_one(
        {"meal_id": meal_id, "family_id": user["family_id"]}, {"_id": 0}
    )
    if not meal:
        raise HTTPException(404, "Not found")

    family = await get_family_doc(user["family_id"])
    # An explicit per-recipe diet wins; otherwise a vegetarian household gets
    # vegetarian recipes without having to ask each time.
    diet = normalize_diet(diet) or normalize_diet(family.get("diet"))
    variant = max(0, min(int(variant or 0), 20))
    slot = recipe_slot(diet)

    cached = (meal.get(slot) or {}).get(language)
    # A "different recipe" (variant>0) is a deliberate ask for something fresh,
    # so it skips the cache; everything else serves the cached copy free.
    if cached and not variant:
        return {"recipe": cached, "cached": True, "diet": diet}

    if not GOOGLE_API_KEY:
        raise HTTPException(503, "Recipe suggestions are unavailable right now.")

    sub = await build_subscription(user["family_id"])
    # Metered against the same monthly AI allowance as document scanning, so a
    # family has one number to understand rather than two.
    if not is_admin_user(user) and family.get("ai_scans_used", 0) >= sub["limits"]["ai_scans_per_month"]:
        plan_limit_error(
            feature="ai_scans",
            current_plan=sub["plan"],
            limit=sub["limits"]["ai_scans_per_month"],
            used=family.get("ai_scans_used", 0),
            message="AI limit reached for this billing period.",
        )

    title = sanitize_user_text(meal.get("title", ""))
    if len(title) < 2:
        raise HTTPException(400, "This meal needs a name first.")
    ingredients = sanitize_ingredients(meal.get("ingredients", []))

    try:
        text = await _gemini_text(
            build_recipe_prompt(title, ingredients, RECIPE_LANGUAGE_NAMES[language],
                                diet=diet, variant=variant),
            system=RECIPE_SYSTEM_PROMPT,
            # The lite model turns a dish name into a structured recipe quickly,
            # and the validator guards the shape either way — so "Cook it"
            # returns sooner. A "different recipe" also runs a little hotter so
            # the variety is real, not sampling noise.
            fast=True,
            temperature=0.9 if variant else None,
        )
        parsed = extract_json(text)
        if parsed is None:
            raise UnsafeRecipe("unparseable")
        recipe = validate_recipe(parsed)
        if "ingredients" not in recipe:
            # Still a usable recipe, but the model ignored half the prompt —
            # loud in the logs so a quiet downgrade cannot become the norm.
            log.warning("recipe for %r came back without ingredients", title)
    except UnsafeRecipe as exc:
        # The specific check that failed is useful to us and meaningless to the
        # user, so it is logged and not returned.
        log.info("recipe rejected by safety gate: %s", exc.reason)
        raise HTTPException(422, "We could not write a recipe for this one.")
    except Exception as exc:
        log.warning("recipe generation failed: %s", exc)
        raise HTTPException(502, "We could not write a recipe for this one.")

    await database["meals"].update_one(
        {"meal_id": meal_id, "family_id": user["family_id"]},
        {"$set": {f"{slot}.{language}": recipe}},
    )
    if not is_admin_user(user):
        # Guarded so two concurrent scans cannot both push the counter past the
        # limit: the increment only lands while the family is still under it.
        # A request that raced past the last slot did its work already — it goes
        # through uncharged rather than over-counting.
        await database["families"].update_one(
            {"family_id": user["family_id"],
             "ai_scans_used": {"$lt": sub["limits"]["ai_scans_per_month"]}},
            {"$inc": {"ai_scans_used": 1}, "$set": {"updated_at": utcnow()}},
        )

    return {"recipe": recipe, "cached": False, "diet": diet}


class DietIn(BaseModel):
    diet: str = ""


@app.get("/api/meals/diet")
async def get_meal_diet(user=Depends(require_user), database=Depends(get_db)):
    """The household's cooking diet, read on the Kitchen tab."""
    family = await get_family_doc(user["family_id"])
    return {"diet": normalize_diet(family.get("diet"))}


@app.put("/api/meals/diet")
async def set_meal_diet(payload: DietIn, user=Depends(require_user), database=Depends(get_db)):
    """Set the household's cooking diet. Vegetarian makes new recipes and the
    weekly suggestions come out vegetarian without asking each time. A closed
    set, so only "vegetarian" or "" (no diet) is ever stored."""
    diet = normalize_diet(payload.diet)
    await database["families"].update_one(
        {"family_id": user["family_id"]},
        {"$set": {"diet": diet, "updated_at": utcnow()}},
    )
    return {"diet": diet}


@app.post("/api/recipes/chef")
async def ask_the_chef(
    payload: dict = Body(...),
    lang: str = "en",
    user: dict = Depends(require_user),
    database=Depends(get_db),
):
    """Answer one short cooking question about a dish — a substitution, a
    variation, a timing query.

    Not tied to a meal document: the recipe page also shows curated recipes
    that were never added to the plan. Not cached: the question is free text,
    so each ask is metered against the family's AI allowance like a scan.
    """
    await require_feature(user, "meal_planner")

    language = lang if lang in RECIPE_LANGUAGE_NAMES else "en"
    title = sanitize_user_text(str(payload.get("title") or ""))
    question = sanitize_user_text(str(payload.get("question") or ""), MAX_QUESTION_LEN)
    if len(title) < 2 or len(question) < 5:
        raise HTTPException(400, "Ask a question about the dish.")

    if not GOOGLE_API_KEY:
        raise HTTPException(503, "The chef is unavailable right now.")

    sub = await build_subscription(user["family_id"])
    family = await get_family_doc(user["family_id"])
    if not is_admin_user(user) and family.get("ai_scans_used", 0) >= sub["limits"]["ai_scans_per_month"]:
        plan_limit_error(
            feature="ai_scans",
            current_plan=sub["plan"],
            limit=sub["limits"]["ai_scans_per_month"],
            used=family.get("ai_scans_used", 0),
            message="AI limit reached for this billing period.",
        )

    try:
        text = await _gemini_text(
            build_chef_prompt(title, question, RECIPE_LANGUAGE_NAMES[language]),
            system=CHEF_SYSTEM_PROMPT,
            fast=True,
        )
        parsed = extract_json(text)
        if parsed is None:
            raise UnsafeRecipe("unparseable")
        answer = validate_chef_answer(parsed)
    except UnsafeRecipe as exc:
        log.info("chef answer rejected by safety gate: %s", exc.reason)
        raise HTTPException(422, "The chef could not answer that one.")
    except Exception as exc:
        log.warning("chef answer failed: %s", exc)
        raise HTTPException(502, "The chef could not answer that one.")

    if not is_admin_user(user):
        # Guarded so two concurrent scans cannot both push the counter past the
        # limit: the increment only lands while the family is still under it.
        # A request that raced past the last slot did its work already — it goes
        # through uncharged rather than over-counting.
        await database["families"].update_one(
            {"family_id": user["family_id"],
             "ai_scans_used": {"$lt": sub["limits"]["ai_scans_per_month"]}},
            {"$inc": {"ai_scans_used": 1}, "$set": {"updated_at": utcnow()}},
        )

    return {"answer": answer}


@app.post("/api/recipes/generate")
async def generate_recipe_from_name(
    payload: dict = Body(...),
    lang: str = "en",
    user: dict = Depends(require_user),
    database=Depends(get_db),
):
    """Write a full recipe for any dish the family types, without a plan entry.

    The meal-planner recipe route needs a saved meal to hang its cache on; this
    is the same generator with the meal document removed, so "ask for fluffy
    pancakes" returns a recipe straight from the Kitchen. Not cached (there is
    nothing to cache it on), so each ask is metered against the family's AI
    allowance exactly like a scan or a chef question. A vegetarian household
    gets a vegetarian recipe without asking, same rule as the planner.
    """
    await require_feature(user, "meal_planner")

    language = lang if lang in RECIPE_LANGUAGE_NAMES else "en"
    title = sanitize_user_text(str(payload.get("title") or ""))
    if len(title) < 2:
        raise HTTPException(400, "Name a dish to get a recipe.")

    if not GOOGLE_API_KEY:
        raise HTTPException(503, "Recipe suggestions are unavailable right now.")

    family = await get_family_doc(user["family_id"])
    diet = normalize_diet(payload.get("diet")) or normalize_diet(family.get("diet"))
    variant = max(0, min(int(payload.get("variant") or 0), 20))

    sub = await build_subscription(user["family_id"])
    # Metered against the same monthly AI allowance as scans and the meal
    # recipes, so a family has one number to understand rather than three.
    if not is_admin_user(user) and family.get("ai_scans_used", 0) >= sub["limits"]["ai_scans_per_month"]:
        plan_limit_error(
            feature="ai_scans",
            current_plan=sub["plan"],
            limit=sub["limits"]["ai_scans_per_month"],
            used=family.get("ai_scans_used", 0),
            message="AI limit reached for this billing period.",
        )

    try:
        text = await _gemini_text(
            build_recipe_prompt(title, [], RECIPE_LANGUAGE_NAMES[language],
                                diet=diet, variant=variant),
            system=RECIPE_SYSTEM_PROMPT,
            fast=True,
            temperature=0.9 if variant else None,
        )
        parsed = extract_json(text)
        if parsed is None:
            raise UnsafeRecipe("unparseable")
        recipe = validate_recipe(parsed)
    except UnsafeRecipe as exc:
        log.info("recipe rejected by safety gate: %s", exc.reason)
        raise HTTPException(422, "We could not write a recipe for this one.")
    except Exception as exc:
        log.warning("recipe generation failed: %s", exc)
        raise HTTPException(502, "We could not write a recipe for this one.")

    if not is_admin_user(user):
        # Guarded so two concurrent asks cannot both push the counter past the
        # limit: the increment only lands while the family is still under it.
        await database["families"].update_one(
            {"family_id": user["family_id"],
             "ai_scans_used": {"$lt": sub["limits"]["ai_scans_per_month"]}},
            {"$inc": {"ai_scans_used": 1}, "$set": {"updated_at": utcnow()}},
        )

    return {"recipe": recipe, "diet": diet}


@app.post("/api/recipes/capture")
async def capture_recipe(
    payload: dict = Body(...),
    user: dict = Depends(require_user),
    database=Depends(get_db),
):
    """Read a photo of a printed or handwritten recipe into a structured one.

    Returns the recipe for review only — committing it to the planner goes
    through /api/meals/from-capture, which re-validates everything.
    """
    await require_feature(user, "meal_planner")

    image_b64 = str(payload.get("image_base64") or "")
    if "," in image_b64:
        image_b64 = image_b64.split(",")[-1]
    if len(image_b64) < 100:
        raise HTTPException(400, "No photo attached.")
    if len(image_b64) > MAX_IMAGE_B64_CHARS:
        raise HTTPException(413, "That photo is too large.")

    if not GOOGLE_API_KEY:
        raise HTTPException(503, "Photo capture is unavailable right now.")

    sub = await build_subscription(user["family_id"])
    family = await get_family_doc(user["family_id"])
    if not is_admin_user(user) and family.get("ai_scans_used", 0) >= sub["limits"]["ai_scans_per_month"]:
        plan_limit_error(
            feature="ai_scans",
            current_plan=sub["plan"],
            limit=sub["limits"]["ai_scans_per_month"],
            used=family.get("ai_scans_used", 0),
            message="AI limit reached for this billing period.",
        )

    try:
        text = await _gemini_vision(
            "Read the recipe in this photo.",
            image_b64,
            system=RECIPE_PHOTO_SYSTEM_PROMPT,
        )
        parsed = extract_json(text)
        if parsed is None:
            raise UnsafeRecipe("unparseable")
        captured = validate_captured_recipe(parsed)
    except UnsafeRecipe as exc:
        log.info("recipe capture rejected by safety gate: %s", exc.reason)
        raise HTTPException(422, "We could not read a recipe in that photo.")
    except Exception as exc:
        log.warning("recipe capture failed: %s", exc)
        raise HTTPException(502, "We could not read a recipe in that photo.")

    if not is_admin_user(user):
        # Guarded so two concurrent scans cannot both push the counter past the
        # limit: the increment only lands while the family is still under it.
        # A request that raced past the last slot did its work already — it goes
        # through uncharged rather than over-counting.
        await database["families"].update_one(
            {"family_id": user["family_id"],
             "ai_scans_used": {"$lt": sub["limits"]["ai_scans_per_month"]}},
            {"$inc": {"ai_scans_used": 1}, "$set": {"updated_at": utcnow()}},
        )

    return {"captured": captured}


@app.post("/api/meals/from-capture")
async def add_meal_from_capture(
    payload: dict = Body(...),
    lang: str = "en",
    user: dict = Depends(require_user),
    database=Depends(get_db),
):
    """Commit a captured recipe to a day of the week.

    The recipe arrives back from the client, so it passes the same gate
    again before anything is stored — the round trip through the app is
    not trusted.
    """
    await require_feature(user, "meal_planner")

    day = str(payload.get("day") or "").lower()
    if day not in DAYS_OF_WEEK:
        raise HTTPException(400, f"Day must be one of {DAYS_OF_WEEK}")
    language = lang if lang in RECIPE_LANGUAGE_NAMES else "en"

    try:
        captured = validate_captured_recipe(payload.get("recipe") or {})
    except UnsafeRecipe as exc:
        log.info("captured recipe rejected on commit: %s", exc.reason)
        raise HTTPException(422, "This recipe did not pass our checks.")

    title = captured.pop("title")
    meal = {
        "meal_id": new_id("meal"),
        "family_id": user["family_id"],
        "day": day,
        "meal_type": "dinner",
        "title": title,
        "ingredients": [i["name"] for i in captured["ingredients"]],
        "notes": None,
        "recipe_id": None,
        # The structured recipe rides on the meal exactly like an AI-written
        # one, so the recipe page shows amounts, stepper and steps for free.
        "ai_recipe": {language: captured},
        "created_at": utcnow(),
    }
    await database["meals"].insert_one(meal)
    return public_meal(meal)


@app.post("/api/meals/suggest-ai")
async def suggest_meals_ai(
    lang: str = "en",
    variant: int = 0,
    user: dict = Depends(require_user),
    database=Depends(get_db),
):
    """Propose a week of dinners from the family's actual shopping list.

    The offline engine can only rank the 38-plus dishes we ship, so the same
    list always produced the same week — and never a dish we had not thought
    to include. This asks for real meals built from what the family bought.

    The app falls back to the offline engine whenever this fails, so a bad
    response, a missing key or a spent quota degrades to the old behaviour
    rather than to nothing.
    """
    await require_feature(user, "meal_planner")

    language = lang if lang in RECIPE_LANGUAGE_NAMES else "en"

    if not GOOGLE_API_KEY:
        raise HTTPException(503, "Meal ideas are unavailable right now.")

    # The current list only — checked items included, since something ticked
    # off is a grocery bought this week, not an absent one.
    items = []
    async for row in database["shopping_list"].find(
        {"family_id": user["family_id"]}, {"_id": 0, "name": 1}
    ):
        name = sanitize_user_text(row.get("name") or "", MAX_INGREDIENT_LEN)
        if name:
            items.append(name)

    # The minimum is judged on the CURRENT list alone. History used to top the
    # list up before this check, which meant an empty list with old trips still
    # produced a week of meals — suggestions from nothing. No list, no meals.
    if len(items) < 3:
        raise HTTPException(422, "Add a few things to your shopping list first.")

    # Recent trips only ever enrich a real list; they cannot substitute for one.
    if len(items) < 6:
        async for row in database["shopping_history"].find(
            {"family_id": user["family_id"]}, {"_id": 0, "items": 1}
        ).sort("created_at", -1).limit(3):
            for raw in (row.get("items") or [])[:20]:
                name = sanitize_user_text(raw or "", MAX_INGREDIENT_LEN)
                if name and name not in items:
                    items.append(name)

    items = items[:40]

    sub = await build_subscription(user["family_id"])
    family = await get_family_doc(user["family_id"])
    if not is_admin_user(user) and family.get("ai_scans_used", 0) >= sub["limits"]["ai_scans_per_month"]:
        plan_limit_error(
            feature="ai_scans",
            current_plan=sub["plan"],
            limit=sub["limits"]["ai_scans_per_month"],
            used=family.get("ai_scans_used", 0),
            message="AI limit reached for this billing period.",
        )

    # Anything already on the plan is excluded, so asking twice in a week gives
    # a different week rather than the same one again.
    planned = []
    async for row in database["meals"].find(
        {"family_id": user["family_id"]}, {"_id": 0, "title": 1}
    ).limit(20):
        title = sanitize_user_text(row.get("title") or "")
        if title:
            planned.append(title)

    # What the last ask(s) proposed, remembered on the family. This is what
    # makes reopening the sheet produce a fresh week: the client resets its
    # "different ideas" counter every open, so without memory here every open
    # sent an identical prompt at low temperature and got the same week back.
    seen = []
    fam_doc = await database["families"].find_one(
        {"family_id": user["family_id"]}, {"_id": 0, "last_meal_suggestions": 1}
    ) or {}
    for raw in (fam_doc.get("last_meal_suggestions") or [])[:14]:
        title = sanitize_user_text(raw or "")
        if title:
            seen.append(title)

    # variant is the "different ideas" counter from the client. Bounded so a
    # hostile value cannot blow up the prompt, and used both to change the
    # prompt bytes and to raise sampling temperature after the first ask.
    variant = max(0, min(int(variant or 0), 20))

    try:
        text = await _gemini_text(
            build_suggest_prompt(
                items, RECIPE_LANGUAGE_NAMES[language], planned,
                variant=variant, seen_titles=seen,
                diet=normalize_diet(family.get("diet")),
            ),
            system=SUGGEST_SYSTEM_PROMPT,
            # First ask stays focused; repeat asks lean into variety.
            temperature=0.4 if variant == 0 else 1.0,
        )
        parsed = extract_json(text)
        if parsed is None:
            raise UnsafeRecipe("unparseable")
        meals = validate_suggestions(parsed, items)
    except UnsafeRecipe as exc:
        log.info("meal suggestions rejected by safety gate: %s", exc.reason)
        raise HTTPException(422, "We could not plan a week from this list.")
    except Exception as exc:
        log.warning("meal suggestion generation failed: %s", exc)
        raise HTTPException(502, "We could not plan a week from this list.")

    # Remember what was proposed (about two weeks' worth, most recent last)
    # so the next ask — however the client counts — avoids repeating it.
    # Memory, not metering: stored for admins too.
    remembered = list(dict.fromkeys(seen + [m["title"] for m in meals]))[-14:]
    await database["families"].update_one(
        {"family_id": user["family_id"]},
        {"$set": {"last_meal_suggestions": remembered}},
    )

    if not is_admin_user(user):
        # Guarded so two concurrent scans cannot both push the counter past the
        # limit: the increment only lands while the family is still under it.
        # A request that raced past the last slot did its work already — it goes
        # through uncharged rather than over-counting.
        await database["families"].update_one(
            {"family_id": user["family_id"],
             "ai_scans_used": {"$lt": sub["limits"]["ai_scans_per_month"]}},
            {"$inc": {"ai_scans_used": 1}, "$set": {"updated_at": utcnow()}},
        )

    return {"meals": meals}


@app.post("/api/meals/sync-shopping")
async def sync_meals_to_shopping(user: dict = Depends(require_user), database=Depends(get_db)):
    await require_feature(user, "meal_planner")
    meals = await database["meals"].find(
        {"family_id": user["family_id"]}, {"_id": 0}
    ).to_list(200)
    all_ingredients = []
    for meal in meals:
        all_ingredients.extend(meal.get("ingredients", []))
    unique = list(set(i.strip() for i in all_ingredients if i.strip()))
    added = 0
    for name in unique:
        existing = await database["shopping_list"].find_one(
            {"family_id": user["family_id"], "name": {"$regex": f"^{re.escape(name)}$", "$options": "i"}}
        )
        if not existing:
            item = {
                "item_id": new_id("shop"),
                "family_id": user["family_id"],
                "name": name,
                # "Other" (not the non-existent "Groceries") so the app's
                # name-based aisle derivation kicks in on display.
                "category": "Other",
                "checked": False,
                "added_by": user.get("name", ""),
                "created_at": utcnow(),
            }
            await database["shopping_list"].insert_one(item)
            added += 1
    return {"ok": True, "added": added, "total_ingredients": len(unique)}


# -----------------------------------------------------------------------------
# Carpool Coordinator
# -----------------------------------------------------------------------------
@app.get("/api/carpools")
async def list_carpools(user: dict = Depends(require_user), database=Depends(get_db)):
    rows = await database["carpools"].find(
        {"family_id": user["family_id"]}, {"_id": 0}
    ).sort("created_at", -1).to_list(100)
    return [public_carpool(c) for c in rows]


@app.post("/api/carpools")
async def create_carpool(body: CarpoolIn, user: dict = Depends(require_user), database=Depends(get_db)):
    await require_feature(user, "carpool")
    carpool = {
        "carpool_id": new_id("cpool"),
        "family_id": user["family_id"],
        "title": body.title,
        "day_of_week": body.day_of_week.lower(),
        "time": body.time,
        "driver_name": body.driver_name,
        "pickup_kids": body.pickup_kids,
        "notes": body.notes,
        "created_at": utcnow(),
    }
    await database["carpools"].insert_one(carpool)
    return public_carpool(carpool)


@app.delete("/api/carpools/{carpool_id}")
async def delete_carpool(carpool_id: str, user: dict = Depends(require_user), database=Depends(get_db)):
    result = await database["carpools"].delete_one(
        {"carpool_id": carpool_id, "family_id": user["family_id"]}
    )
    if result.deleted_count == 0:
        raise HTTPException(404, "Carpool not found")
    return {"ok": True}


# -----------------------------------------------------------------------------
# Allowance Tracker
# -----------------------------------------------------------------------------
@app.get("/api/allowances")
async def list_allowances(user: dict = Depends(require_user), database=Depends(get_db)):
    rows = await database["allowances"].find(
        {"family_id": user["family_id"]}, {"_id": 0}
    ).to_list(50)
    return [public_allowance_config(a) for a in rows]


@app.post("/api/allowances")
async def set_allowance(body: AllowanceIn, user: dict = Depends(require_user), database=Depends(get_db)):
    await require_feature(user, "allowance")
    existing = await database["allowances"].find_one(
        {"family_id": user["family_id"], "member_id": body.member_id}
    )
    if existing:
        await database["allowances"].update_one(
            {"allowance_id": existing["allowance_id"]},
            {"$set": {"amount": body.amount, "frequency": body.frequency}},
        )
        existing["amount"] = body.amount
        existing["frequency"] = body.frequency
        return public_allowance_config(existing)
    allowance = {
        "allowance_id": new_id("alw"),
        "family_id": user["family_id"],
        "member_id": body.member_id,
        "amount": body.amount,
        "frequency": body.frequency,
        "created_at": utcnow(),
    }
    await database["allowances"].insert_one(allowance)
    return public_allowance_config(allowance)


@app.delete("/api/allowances/{member_id}")
async def delete_allowance(member_id: str, user: dict = Depends(require_user), database=Depends(get_db)):
    result = await database["allowances"].delete_one(
        {"family_id": user["family_id"], "member_id": member_id}
    )
    if result.deleted_count == 0:
        raise HTTPException(404, "Allowance not found")
    return {"ok": True}


@app.get("/api/allowances/{member_id}/transactions")
async def list_allowance_transactions(member_id: str, user: dict = Depends(require_user), database=Depends(get_db)):
    rows = await database["allowance_txns"].find(
        {"family_id": user["family_id"], "member_id": member_id}, {"_id": 0}
    ).sort("created_at", -1).to_list(100)
    return [public_allowance_txn(t) for t in rows]


@app.post("/api/allowances/transaction")
async def add_allowance_transaction(body: AllowanceTxnIn, user: dict = Depends(require_user), database=Depends(get_db)):
    await require_feature(user, "allowance")
    txn = {
        "txn_id": new_id("atxn"),
        "family_id": user["family_id"],
        "member_id": body.member_id,
        "amount": body.amount,
        "description": body.description,
        "txn_type": body.txn_type,
        "created_at": utcnow(),
    }
    await database["allowance_txns"].insert_one(txn)
    return public_allowance_txn(txn)


@app.post("/api/allowances/{member_id}/pay")
async def pay_allowance(member_id: str, user: dict = Depends(require_user), database=Depends(get_db)):
    """Record this period's pocket money in one tap.

    Deliberately not automatic. An accrual on a timer would credit money that
    may never have physically changed hands, and a child's balance that says
    €20 when the tin holds €5 is worse than no tracker at all. A parent presses
    this when they actually hand it over.
    """
    await require_feature(user, "allowance")
    config = await database["allowances"].find_one(
        {"family_id": user["family_id"], "member_id": member_id}, {"_id": 0}
    )
    if not config:
        raise HTTPException(status_code=404, detail="No allowance set for this child")

    due = allowance_next_due(config)
    if (_coerce_dt(due) or utcnow()) > utcnow():
        raise HTTPException(status_code=400, detail="Not due yet")

    # Claim the period before writing the money, and carry the previous
    # last_paid_at in the filter so two taps — or two parents — settle it once.
    previous = config.get("last_paid_at")
    claim = {"allowance_id": config["allowance_id"], "family_id": user["family_id"]}
    claim["last_paid_at"] = previous if previous else {"$exists": False}
    result = await database["allowances"].update_one(claim, {"$set": {"last_paid_at": utcnow()}})
    if result.matched_count == 0:
        raise HTTPException(status_code=409, detail="Already paid for this period")

    txn = {
        "txn_id": new_id("atxn"),
        "family_id": user["family_id"],
        "member_id": member_id,
        "amount": config["amount"],
        "description": "Pocket money",
        "txn_type": "deposit",
        "created_at": utcnow(),
    }
    await database["allowance_txns"].insert_one(txn)

    updated = await database["allowances"].find_one(
        {"allowance_id": config["allowance_id"], "family_id": user["family_id"]}, {"_id": 0}
    )
    return {"ok": True, "transaction": public_allowance_txn(txn),
            "allowance": public_allowance_config(updated)}


@app.get("/api/allowances/{member_id}/balance")
async def get_allowance_balance(member_id: str, user: dict = Depends(require_user), database=Depends(get_db)):
    txns = await database["allowance_txns"].find(
        {"family_id": user["family_id"], "member_id": member_id}, {"_id": 0}
    ).to_list(1000)
    balance = 0.0
    for t in txns:
        if t.get("txn_type") == "withdrawal":
            balance -= t.get("amount", 0)
        else:
            balance += t.get("amount", 0)
    return {"member_id": member_id, "balance": round(balance, 2)}


# -----------------------------------------------------------------------------
# Family Announcements / Chat
# -----------------------------------------------------------------------------
@app.get("/api/announcements")
async def list_announcements(user: dict = Depends(require_user), database=Depends(get_db)):
    rows = await database["announcements"].find(
        {"family_id": user["family_id"]}, {"_id": 0}
    ).sort("created_at", -1).to_list(50)
    return [public_announcement(a) for a in rows]


@app.post("/api/announcements")
async def create_announcement(body: AnnouncementIn, user: dict = Depends(require_user), database=Depends(get_db)):
    announcement = {
        "announcement_id": new_id("ann"),
        "family_id": user["family_id"],
        "text": body.text,
        "author_name": user.get("name", ""),
        "priority": body.priority,
        "created_at": utcnow(),
    }
    await database["announcements"].insert_one(announcement)
    try:
        who = user.get("name") or "A co-parent"
        await send_coparent_alert(
            user["family_id"], f"{who} posted an announcement", body.text, "announcement",
            created_by_user_id=user["user_id"],
        )
    except Exception as e:
        log.warning("announcement alert failed: %s", e)
    return public_announcement(announcement)


@app.delete("/api/announcements/{announcement_id}")
async def delete_announcement(announcement_id: str, user: dict = Depends(require_user), database=Depends(get_db)):
    result = await database["announcements"].delete_one(
        {"announcement_id": announcement_id, "family_id": user["family_id"]}
    )
    if result.deleted_count == 0:
        raise HTTPException(404, "Announcement not found")
    return {"ok": True}


# -----------------------------------------------------------------------------
# Document Expiry Alerts
# -----------------------------------------------------------------------------
@app.get("/api/vault/expiry-alerts")
async def vault_expiry_alerts(user: dict = Depends(require_full_member), database=Depends(get_db)):
    docs = await database["vault"].find(
        {"family_id": user["family_id"]}, {"_id": 0}
    ).to_list(500)
    alerts = []
    for doc in docs:
        # An expiry reminder leaks a document's title too — same rule as the list.
        if not _may_see_vault_doc(doc, user):
            continue
        exp = doc.get("expiry_date")
        if exp:
            exp_dt = ensure_aware_utc(exp)
            if exp_dt:
                days_left = (exp_dt - utcnow()).days
                alerts.append({
                    "doc_id": doc["doc_id"],
                    "title": doc.get("title", ""),
                    "category": doc.get("category", ""),
                    "expiry_date": iso(exp_dt),
                    "days_left": days_left,
                    "status": "expired" if days_left < 0 else "urgent" if days_left <= 30 else "upcoming",
                })
    alerts.sort(key=lambda a: a["days_left"])
    return alerts


@app.patch("/api/vault/{doc_id}/expiry")
async def set_vault_expiry(doc_id: str, expiry_date: str = Query(...), user: dict = Depends(require_full_member), database=Depends(get_db)):
    exp_dt = parse_dt(expiry_date)
    if not exp_dt:
        raise HTTPException(400, "Invalid date format")
    existing = await database["vault"].find_one(
        {"doc_id": doc_id, "family_id": user["family_id"]}, {"_id": 0}
    )
    if existing and not _may_see_vault_doc(existing, user):
        raise HTTPException(404, "Document not found")
    result = await database["vault"].update_one(
        {"doc_id": doc_id, "family_id": user["family_id"]},
        {"$set": {"expiry_date": exp_dt}},
    )
    if result.matched_count == 0:
        raise HTTPException(404, "Document not found")
    return {"ok": True, "doc_id": doc_id, "expiry_date": iso(exp_dt)}


# -----------------------------------------------------------------------------
# Weekly Report Card
# -----------------------------------------------------------------------------
@app.get("/api/report/weekly")
async def weekly_report(user: dict = Depends(require_user), database=Depends(get_db)):
    await require_feature(user, "weekly_report")
    now = utcnow()
    week_ago = now - timedelta(days=7)
    fid = user["family_id"]

    cards_done = await database["cards"].count_documents(
        {"family_id": fid, "status": "DONE", "completed_at": {"$gte": week_ago}}
    )
    cards_created = await database["cards"].count_documents(
        {"family_id": fid, "created_at": {"$gte": week_ago}}
    )
    cards_overdue = await database["cards"].count_documents(
        {"family_id": fid, "status": "OPEN", "due_date": {"$lt": now}}
    )

    star_txns = await database["star_transactions"].find(
        {"family_id": fid, "created_at": {"$gte": week_ago}}, {"_id": 0}
    ).to_list(500)
    stars_given = sum(t.get("delta", 0) for t in star_txns if t.get("delta", 0) > 0)

    expenses = await database["expenses"].find(
        {"family_id": fid, "created_at": {"$gte": week_ago}}, {"_id": 0}
    ).to_list(500)
    total_spent = sum(e.get("amount", 0) for e in expenses)
    expense_categories = {}
    for e in expenses:
        cat = e.get("category", "Other")
        expense_categories[cat] = expense_categories.get(cat, 0) + e.get("amount", 0)

    upcoming_cards = await database["cards"].find(
        {
            "family_id": fid,
            "status": "OPEN",
            "due_date": {"$gte": now, "$lte": now + timedelta(days=7)},
            # Don't list a co-parent's private item titles in this report.
            "$or": [
                {"shared": True},
                {"created_by_user_id": user["user_id"]},
                {"created_by_user_id": {"$exists": False}},
            ],
        },
        {"_id": 0, "title": 1, "due_date": 1, "type": 1, "assignee": 1},
    ).sort("due_date", 1).to_list(10)

    routine_logs = await database["routine_logs"].count_documents(
        {"family_id": fid, "completed_at": {"$gte": week_ago}}
    )

    return {
        "period_start": iso(week_ago),
        "period_end": iso(now),
        "tasks_completed": cards_done,
        "tasks_created": cards_created,
        "tasks_overdue": cards_overdue,
        "stars_earned": stars_given,
        "total_spent": round(total_spent, 2),
        "expense_by_category": expense_categories,
        "routines_completed": routine_logs,
        "upcoming_deadlines": [
            {
                "title": c.get("title", ""),
                "due_date": iso(ensure_aware_utc(c.get("due_date"))),
                "type": c.get("type", "TASK"),
                "assignee": c.get("assignee"),
            }
            for c in upcoming_cards
        ],
    }


@app.get("/api/report/lite")
async def report_lite(user: dict = Depends(require_user), database=Depends(get_db)):
    """A minimal, un-gated weekly recap (tasks done + stars earned in the last
    7 days) for the Sunday recap — available to every plan, unlike the full
    premium weekly report."""
    now = utcnow()
    week_ago = now - timedelta(days=7)
    fid = user["family_id"]
    tasks_done = await database["cards"].count_documents(
        {"family_id": fid, "status": "DONE", "completed_at": {"$gte": week_ago}}
    )
    star_txns = await database["star_transactions"].find(
        {"family_id": fid, "created_at": {"$gte": week_ago}}, {"_id": 0}
    ).to_list(500)
    stars_earned = sum(t.get("delta", 0) for t in star_txns if t.get("delta", 0) > 0)
    return {"tasks_done": tasks_done, "stars_earned": stars_earned}


# -----------------------------------------------------------------------------
# Chore Wheel
# -----------------------------------------------------------------------------
@app.get("/api/chores")
async def list_chores(user: dict = Depends(require_user), database=Depends(get_db)):
    rows = await database["chores"].find(
        {"family_id": user["family_id"]}, {"_id": 0}
    ).sort("created_at", -1).to_list(100)
    return [public_chore(c) for c in rows]


@app.post("/api/chores")
async def create_chore(body: ChoreIn, user: dict = Depends(require_user), database=Depends(get_db)):
    chore = {
        "chore_id": new_id("chore"),
        "family_id": user["family_id"],
        "title": body.title,
        "frequency": body.frequency,
        "assigned_members": body.assigned_members,
        "current_assignee": body.assigned_members[0] if body.assigned_members else None,
        "rotate": body.rotate,
        "star_reward": max(0, int(body.star_reward or 0)),
        "last_rotated": utcnow(),
        "created_at": utcnow(),
    }
    await database["chores"].insert_one(chore)
    return public_chore(chore)


@app.post("/api/chores/{chore_id}/complete")
async def complete_chore(chore_id: str, user: dict = Depends(require_user), database=Depends(get_db)):
    """Mark a chore done: award the person who did it, then pass it on.

    Rotate alone only moves the chore to the next child — it says nothing about
    the work being done, so chores earned nothing. This is the "done" action:
    the current assignee is paid first, and only then does the wheel turn, so
    the stars go to whoever actually did it rather than the next person up.
    """
    chore = await database["chores"].find_one(
        {"chore_id": chore_id, "family_id": user["family_id"]}, {"_id": 0}
    )
    if not chore:
        raise HTTPException(404, "Chore not found")

    doer = chore.get("current_assignee")

    # Anti-double-pay claim: every other star path guards the award atomically,
    # but this one read-then-awarded unconditionally, so a rapid double-tap or a
    # retried request paid twice (and the second star went to the just-rotated
    # next child). Claim a coarse 2-second completion bucket with the same
    # $ne-on-a-marker pattern the weekly-claim guards use: a duplicate within the
    # same bucket sees the marker already set and loses the claim (returns
    # idempotently, no payment), while a genuine later re-completion — minutes or
    # days apart — falls in a new bucket and pays normally.
    now = utcnow()
    bucket = int(now.timestamp()) // 2
    claim = await database["chores"].update_one(
        {"chore_id": chore_id, "family_id": user["family_id"],
         "last_completed_bucket": {"$ne": bucket}},
        {"$set": {"last_completed_bucket": bucket, "last_completed_at": now}},
    )
    if claim.matched_count == 0:
        return {"ok": True, "chore": public_chore(chore), "stars_awarded": 0, "member_id": doer}

    txn = await award_stars_to_member(
        database,
        user["family_id"],
        doer,
        int(chore.get("star_reward", 0) or 0),
        chore.get("title") or "Chore done",
        user,
    )

    await database["chore_logs"].insert_one({
        "log_id": new_id("clog"),
        "chore_id": chore_id,
        "family_id": user["family_id"],
        "member_id": doer,
        "completed_at": utcnow(),
        "stars_awarded": txn["delta"] if txn else 0,
    })

    members = chore.get("assigned_members", [])
    if chore.get("rotate", True) and len(members) >= 2:
        try:
            next_idx = (members.index(doer) + 1) % len(members)
        except ValueError:
            next_idx = 0
        await database["chores"].update_one(
            {"chore_id": chore_id},
            {"$set": {"current_assignee": members[next_idx], "last_rotated": utcnow()}},
        )
        chore["current_assignee"] = members[next_idx]
        chore["last_rotated"] = utcnow()

    return {
        "ok": True,
        "chore": public_chore(chore),
        "stars_awarded": txn["delta"] if txn else 0,
        "member_id": doer,
    }


@app.post("/api/chores/{chore_id}/rotate")
async def rotate_chore(chore_id: str, user: dict = Depends(require_user), database=Depends(get_db)):
    chore = await database["chores"].find_one(
        {"chore_id": chore_id, "family_id": user["family_id"]}, {"_id": 0}
    )
    if not chore:
        raise HTTPException(404, "Chore not found")
    members = chore.get("assigned_members", [])
    if len(members) < 2:
        return public_chore(chore)
    current = chore.get("current_assignee")
    try:
        idx = members.index(current)
        next_idx = (idx + 1) % len(members)
    except ValueError:
        next_idx = 0
    await database["chores"].update_one(
        {"chore_id": chore_id},
        {"$set": {"current_assignee": members[next_idx], "last_rotated": utcnow()}},
    )
    chore["current_assignee"] = members[next_idx]
    chore["last_rotated"] = utcnow()
    return public_chore(chore)


@app.delete("/api/chores/{chore_id}")
async def delete_chore(chore_id: str, user: dict = Depends(require_user), database=Depends(get_db)):
    result = await database["chores"].delete_one(
        {"chore_id": chore_id, "family_id": user["family_id"]}
    )
    if result.deleted_count == 0:
        raise HTTPException(404, "Chore not found")
    return {"ok": True}


# -----------------------------------------------------------------------------
# Support Contact
# -----------------------------------------------------------------------------

class SupportContactIn(BaseModel):
    subject: str
    message: str

# -----------------------------------------------------------------------------
# Metrics (first-party, count-only — no payloads, no third-party SDKs)
# -----------------------------------------------------------------------------
ALLOWED_EVENTS = {
    "feed_open", "scan_used", "card_created", "vault_added", "vault_shared",
    "kids_open", "calendar_open", "onboarding_done", "calendar_import_cancelled",
    # AI reliability: bumped server-side from the central Gemini path so the
    # Metrics screen can show a real success rate, not just a live probe.
    "ai_call_ok", "ai_call_error",
}


async def _bump_metric(name: str) -> None:
    """Fire-and-forget daily counter for `name` in metrics_daily. Never raises —
    telemetry must not break the feature it measures."""
    try:
        database = get_db()
        today = utcnow().strftime("%Y-%m-%d")
        await database["metrics_daily"].update_one(
            {"date": today, "name": name},
            {"$inc": {"count": 1}},
            upsert=True,
        )
    except Exception:
        pass


class MetricEventIn(BaseModel):
    name: str


@app.post("/api/metrics/event")
async def log_metric_event(payload: MetricEventIn, user=Depends(require_user)):
    name = (payload.name or "").strip()
    if name not in ALLOWED_EVENTS:
        return {"ok": False}
    database = get_db()
    today = utcnow().strftime("%Y-%m-%d")
    try:
        await database["metrics_daily"].update_one(
            {"date": today, "name": name},
            {"$inc": {"count": 1}},
            upsert=True,
        )
    except Exception:
        pass
    return {"ok": True}


@app.get("/api/metrics/summary")
async def metrics_summary(days: int = 14, user=Depends(require_user)):
    if not is_admin_user(user):
        raise HTTPException(status_code=403, detail="Admin only")
    database = get_db()
    days = max(1, min(days, 90))
    cutoff = (utcnow() - timedelta(days=days)).strftime("%Y-%m-%d")
    rows = []
    cursor = database["metrics_daily"].find(
        {"date": {"$gte": cutoff}}, {"_id": 0}
    ).sort("date", -1)
    async for row in cursor:
        rows.append(row)
    return {"days": days, "rows": rows}


@app.get("/api/metrics/funnel")
async def metrics_funnel(days: int = 30, user=Depends(require_user), database=Depends(get_db)):
    """The activation + growth funnel, computed from existing collections (no
    per-user event pipeline needed): who signs up, finishes onboarding, invites
    a co-parent, has one join, shares an item — plus daily/weekly active users.
    Admin only. This is what turns "make the launch stick" from vibes to data.
    """
    if not is_admin_user(user):
        raise HTTPException(status_code=403, detail="Admin only")
    days = max(1, min(days, 365))
    now = utcnow()
    cutoff = now - timedelta(days=days)
    d1 = now - timedelta(days=1)
    d7 = now - timedelta(days=7)

    users_col = database["users"]
    invites_col = database["family_invites"]
    members_col = database["family_members"]
    cards_col = database["cards"]

    async def count_group(col, match):
        """Distinct family_id count matching `match`, via aggregation."""
        n = 0
        pipeline = [{"$match": match}, {"$group": {"_id": "$family_id"}}, {"$count": "c"}]
        async for row in col.aggregate(pipeline):
            n = row.get("c", 0)
        return n

    multi_member = 0
    async for row in members_col.aggregate([
        {"$group": {"_id": "$family_id", "n": {"$sum": 1}}},
        {"$match": {"n": {"$gt": 1}}},
        {"$count": "c"},
    ]):
        multi_member = row.get("c", 0)

    return {
        "window_days": days,
        "total_users": await users_col.count_documents({}),
        # New in the window
        "signups": await users_col.count_documents({"created_at": {"$gte": cutoff}}),
        "onboarded": await users_col.count_documents(
            {"created_at": {"$gte": cutoff}, "onboarding_completed": True}),
        "invites_sent": await invites_col.count_documents({"created_at": {"$gte": cutoff}}),
        "invites_accepted": await invites_col.count_documents(
            {"status": "accepted", "accepted_at": {"$gte": cutoff}}),
        # Activation state (all-time): the household actually became shared
        "multi_member_households": multi_member,
        "sharing_households": await count_group(cards_col, {"shared": True}),
        # Retention
        "active_1d": await users_col.count_documents({"last_active_at": {"$gte": d1}}),
        "active_7d": await users_col.count_documents({"last_active_at": {"$gte": d7}}),
    }


@app.post("/api/support/contact")
async def submit_support_contact(
    body: SupportContactIn,
    user: dict = Depends(require_user),
    database=Depends(get_db),
):
    subject = body.subject.strip()[:200]
    message = body.message.strip()[:5000]
    if not subject or not message:
        raise HTTPException(400, "Subject and message are required")
    ticket = {
        "ticket_id": new_id("tkt"),
        "family_id": user["family_id"],
        "user_id": user["user_id"],
        "user_email": user.get("email", ""),
        "user_name": user.get("name", ""),
        "subject": subject,
        "message": message,
        "status": "open",
        "created_at": utcnow(),
    }
    await database["support_tickets"].insert_one(ticket)
    return {"ok": True, "ticket_id": ticket["ticket_id"]}


# -----------------------------------------------------------------------------
# Railway entrypoint
# -----------------------------------------------------------------------------
if __name__ == "__main__":
    import uvicorn

    port = int(os.environ.get("PORT", "8080"))
    uvicorn.run(app, host="0.0.0.0", port=port)
